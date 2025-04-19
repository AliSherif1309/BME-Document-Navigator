# START OF FULL SCRIPT (v4 - File Tree Browser, Session, Notes Edit/Del, Outline+, Rank)
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, PanedWindow, Text, Scrollbar, Canvas, Frame, Label, Toplevel, Entry, simpledialog, Menu, Listbox
import sqlite3
import os
import sys
import platform
import subprocess
import threading
import queue
import fitz  # PyMuPDF
import time # For timestamps
import configparser # For session state
import ast # For evaluating stored tuples/dicts safely
try:
    from docx import Document # python-docx
    DOCX_ENABLED = True
except ImportError:
    print("WARNING: python-docx not found. DOCX support will be disabled.")
    DOCX_ENABLED = False
try:
    from PIL import Image, ImageTk
    PIL_ENABLED = True
except ImportError:
    print("WARNING: Pillow not found. PDF/Icon support might be affected.")
    PIL_ENABLED = False

import sys
import re # For metadata extraction
from collections import defaultdict # For managing tab state

# --- Configuration ---
DATABASE_FILE = 'bme_doc_index.db'
CONFIG_FILE = 'bme_navigator.ini' # For saving state

# --- Constants ---
ZOOM_STEP = 0.2
MIN_ZOOM = 0.3
MAX_ZOOM = 5.0
DEFAULT_ZOOM = 1.0
# --- UPDATE SUPPORTED EXTENSIONS ---
SUPPORTED_EXTENSIONS = (
    # Documents
    '.pdf', '.docx', '.doc', '.txt', '.rtf',
    '.html', '.htm',
    # Spreadsheets
    '.xlsx', '.xls', '.csv',
    # Presentations
    '.pptx', '.ppt',
    # Images (Common raster formats)
    '.png', '.jpg', '.jpeg', '.bmp', '.gif', '.tif', '.tiff',
    # --- ADD Compressed Formats ---
    '.zip', '.7z', '.rar', '.gz', '.tar', # Add common archive types
    '.exe',
)

# --- Global Variables ---
add_favorite_menu_index = 1
tab_states = {}
open_files_map = {}
root = None
viewer_notebook = None
details_notebook = None
file_tree = None
search_entry = None
status_bar_label = None
scan_progress_bar = None
metadata_widgets = {}
links_listbox = None
links_map = {}
notes_text_widget = None
outline_tree = None
outline_search_entry = None # For filtering outline
main_paned_window = None
config = configparser.ConfigParser() # For saving state
selected_note_id = None # Tracks the currently selected note_id in the Text widget
left_toggle_button = None
right_toggle_button = None
search_results_queue = queue.Queue() # Queue for search results
search_button_ref = None # Store reference to the search button
scan_status_queue = queue.Queue() # Queue for scan thread communication
scan_button_ref = None # Store reference to scan menu/button

# State for collapsible panes
is_left_pane_collapsed = False
is_right_pane_collapsed = False
left_sash_expanded_pos = 350
right_sash_expanded_pos = None

# --- Configuration & Session Functions ---
def load_config():
    """Loads settings from the config file."""
    global left_sash_expanded_pos # Allow modification
    config.read(CONFIG_FILE)
    # Load window geometry if present
    if 'Window' in config and 'geometry' in config['Window']:
        try:
            root.geometry(config['Window']['geometry'])
        except Exception as e:
            print(f"Error applying window geometry: {e}")
    # Load sash positions if present
    if 'Panes' in config:
        left_sash_expanded_pos = config['Panes'].getint('left_sash', 350) # Load stored expanded pos
        # Right sash pos is relative, load it later after window shown
        # is_left_collapsed = config['Panes'].getboolean('left_collapsed', False) # Optional: restore collapse state
        # is_right_collapsed = config['Panes'].getboolean('right_collapsed', False)

def save_config():
    """Saves current settings (window, panes, open tabs) to the config file."""
    global root, main_paned_window, is_left_pane_collapsed, is_right_pane_collapsed
    global left_sash_expanded_pos, right_sash_expanded_pos
    global viewer_notebook, tab_states # <<< Need notebook and states

    # --- Save Window Geometry ---
    if root and root.winfo_exists(): # Check if root window still exists
         if not config.has_section('Window'): config.add_section('Window')
         config['Window']['geometry'] = root.geometry()
    else: # Don't save geometry if window doesn't exist (e.g., error during init)
         if config.has_section('Window'): config.remove_section('Window')

    # --- Save Pane State ---
    if not config.has_section('Panes'): config.add_section('Panes')
    try:
        if main_paned_window and main_paned_window.winfo_exists():
            # Left Pane Sash
            current_left_pos = main_paned_window.sash_coord(0)[0]
            if not is_left_pane_collapsed and current_left_pos > COLLAPSED_PANE_WIDTH * 2:
                 config['Panes']['left_sash'] = str(current_left_pos)
            elif is_left_pane_collapsed and left_sash_expanded_pos > COLLAPSED_PANE_WIDTH * 2:
                 config['Panes']['left_sash'] = str(left_sash_expanded_pos)
            # Right Pane Sash
            current_right_pos = main_paned_window.sash_coord(1)[0]
            window_width = root.winfo_width() # Need window width here
            if not is_right_pane_collapsed and current_right_pos < window_width - COLLAPSED_PANE_WIDTH * 2:
                 config['Panes']['right_sash'] = str(current_right_pos)
            elif is_right_pane_collapsed and right_sash_expanded_pos is not None and right_sash_expanded_pos < window_width - COLLAPSED_PANE_WIDTH * 2:
                 config['Panes']['right_sash'] = str(right_sash_expanded_pos)
        # else: Don't save sash if panedwindow doesn't exist
    except Exception as e:
        print(f"Error getting sash positions for saving: {e}")

        # --- ADD THIS: Save Open Tabs ---
    if not config.has_section('Session'): config.add_section('Session')
    open_tab_data = [] # Store list of tuples: [(doc_id, page, zoom), ...]
    if viewer_notebook and viewer_notebook.winfo_exists():
        current_tabs = viewer_notebook.tabs()
        for tab_id in current_tabs:
            state = tab_states.get(tab_id)
            # Only save tabs that have a valid doc_id and state
            if state and 'doc_id' in state and state['doc_id'] is not None:
                 doc_id = state['doc_id']
                 page_num = state.get('page_num', 0) # Default to 0 if not found
                 zoom = state.get('zoom', DEFAULT_ZOOM) # Default to standard zoom
                 # Use repr to save the tuple as a string that literal_eval can parse
                 open_tab_data.append(repr((doc_id, page_num, zoom)))

    if open_tab_data:
        # Join the string representations with a specific delimiter (e.g., pipe |)
        # Using comma inside repr tuples might conflict with comma used for splitting later if not careful
        config['Session']['open_tabs_data'] = "|".join(open_tab_data) # Use '|' as delimiter
        print(f"Saving open tabs data: {config['Session']['open_tabs_data']}")
        # Remove the old 'open_tabs' key if it exists
        if config.has_option('Session', 'open_tabs'):
             config.remove_option('Session', 'open_tabs')
    elif config.has_option('Session', 'open_tabs_data'): # If no tabs open, remove the entry
        config.remove_option('Session', 'open_tabs_data')
        print("Clearing saved open tabs data.")
    # --- END MODIFICATION ---

    # Write to file
    try:
        with open(CONFIG_FILE, 'w') as configfile:
            config.write(configfile)
        print("Configuration saved.")
    except Exception as e:
        print(f"Error saving configuration: {e}")

def apply_saved_sash_positions():
     """Applies saved sash positions after main window is drawn."""
     global main_paned_window, config, root
     if 'Panes' in config and main_paned_window:
          # Apply left position first
          saved_left = config['Panes'].getint('left_sash', 350)
          try:
              main_paned_window.sash_place(0, saved_left, 0)
              print(f"Restored left sash position: {saved_left}")
              # Update expanded pos state variable
              global left_sash_expanded_pos
              left_sash_expanded_pos = saved_left
          except Exception as e: print(f"Error restoring left sash: {e}")

          # Apply right position
          # Need window width to calculate relative position correctly if stored
          root.update_idletasks() # Ensure winfo_width is accurate
          window_width = root.winfo_width()
          # Default to window_width - 300 if not saved or invalid
          saved_right = config['Panes'].getint('right_sash', window_width - 300)
          # Sanity check the saved right position
          saved_right = min(saved_right, window_width - COLLAPSED_PANE_WIDTH - 50) # Ensure not too far right
          saved_right = max(saved_right, saved_left + 200) # Ensure not too far left
          try:
              main_paned_window.sash_place(1, saved_right, 0)
              print(f"Restored right sash position: {saved_right}")
              # Update expanded pos state variable
              global right_sash_expanded_pos
              right_sash_expanded_pos = saved_right
          except Exception as e: print(f"Error restoring right sash: {e}")

def restore_session_tabs():
    """Loads the list of open tabs from config (with page/zoom) and reopens them."""
    global config, root, tab_states, viewer_notebook, status_bar_label # Need key globals

    session_key = 'open_tabs_data'
    if not config.has_section('Session') or not config.has_option('Session', session_key):
        print("No saved tab session data found.")
        return

    saved_tabs_str_list = config['Session'][session_key]
    if not saved_tabs_str_list:
        print("Saved tab session data is empty.")
        return

    tab_data_strings = saved_tabs_str_list.split('|')
    print(f"Attempting to restore session tabs from data: {tab_data_strings}")

    if status_bar_label: status_bar_label.config(text="Restoring session tabs...")
    if root: root.update_idletasks()
    restored_count = 0
    failed_count = 0

    if not viewer_notebook:
         print("Error: Viewer notebook not available for restoring tabs.")
         return

    for tab_data_str in tab_data_strings:
        doc_id = None # Ensure doc_id is reset for each loop iteration
        try:
            # Use ast.literal_eval for safe parsing
            parsed_data = ast.literal_eval(tab_data_str)
            if not isinstance(parsed_data, tuple) or len(parsed_data) != 3:
                 raise ValueError("Parsed data is not a valid 3-element tuple")

            doc_id, saved_page_num, saved_zoom = parsed_data
            doc_id = int(doc_id)
            saved_page_num = int(saved_page_num)
            saved_zoom = float(saved_zoom)

            # Verify doc_id still exists and file path is valid before opening
            details = get_document_details(doc_id)
            if details and os.path.exists(details[2]): # Check DB record and file existence
                filepath_to_open = details[2]
                print(f" - Restoring Doc ID {doc_id} ({os.path.basename(filepath_to_open)})")

                # --- Open the tab (or select if already opened somehow) ---
                # open_document_in_tab handles selecting if already open map entry exists
                restored_tab_id = open_document_in_tab(doc_id)

                if restored_tab_id:
                    # --- Find state and apply saved page/zoom ---
                    found_state = tab_states.get(restored_tab_id)

                    if found_state:
                        print(f"   Applying state to restored tab {restored_tab_id}: Target Page={saved_page_num+1}, Zoom={saved_zoom:.2f}")
                        # Check if document object is loaded (needed for clamping page num)
                        if found_state.get('doc_obj'): # Check if doc_obj exists first
                            doc_length = 0
                            try:
                                 # Check if it's a valid fitz doc before getting len
                                 if isinstance(found_state['doc_obj'], fitz.Document):
                                      doc_length = len(found_state['doc_obj'])
                                 else: # Handle non-pdf case if necessary, maybe just don't set page?
                                     print(f"   Warning: Restored tab {restored_tab_id} doc_obj is not a PDF, cannot set page/zoom accurately.")
                                     initial_load_successful = True # Mark as 'loaded' enough for restore count

                            except Exception as len_e:
                                 print(f"   Error getting document length for tab {restored_tab_id}: {len_e}")

                            # Apply state *after* document object is confirmed/length known
                            found_state['zoom'] = max(MIN_ZOOM, min(MAX_ZOOM, saved_zoom)) # Apply clamped zoom

                            # Apply clamped page number only if doc has length > 0
                            if doc_length > 0:
                                 target_page = max(0, min(saved_page_num, doc_length - 1))
                                 found_state['page_num'] = target_page

                                 # Load the specific page with the restored zoom
                                 # Make sure load_pdf_page also checks if doc_obj is valid fitz.Document
                                 if isinstance(found_state['doc_obj'], fitz.Document):
                                      load_pdf_page(restored_tab_id)
                                      initial_load_successful = True # Now mark as fully loaded
                                 else:
                                     # If not PDF, page num might not apply, but tab is open
                                      initial_load_successful = True # Still count as restored

                            else: # Document has 0 pages or length couldn't be determined
                                print(f"   Warning: Document in tab {restored_tab_id} has 0 pages or length error.")
                                initial_load_successful = True # Count as restored

                            # Ensure Back/Forward buttons are updated after setting state
                            update_back_forward_buttons(restored_tab_id)

                        else: # Doc obj wasn't loaded successfully inside open_document_in_tab
                             print(f"   Warning: Document object not ready/valid for tab {restored_tab_id} when restoring state.")
                             initial_load_successful = True # Count as restored, even if page/zoom not set

                        restored_count += 1
                    else:
                        print(f"   Error: Could not find state dictionary for newly opened/restored tab {restored_tab_id}.")
                        failed_count += 1
                else:
                    # open_document_in_tab returned None (e.g., file check failed inside it)
                    print(f"   Skipping restore for Doc ID {doc_id}: open_document_in_tab failed.")
                    failed_count += 1

            else:
                 print(f"Skipping restore for Doc ID {doc_id}: Document details not found or file path missing.")
                 failed_count += 1
        except (SyntaxError, ValueError, TypeError) as parse_e: # Catch eval/int/float errors
            print(f"Skipping restore: Error parsing saved tab data '{tab_data_str}': {parse_e}")
            failed_count += 1
        except Exception as e:
            print(f"Unexpected error restoring tab for Doc ID {doc_id if 'doc_id' in locals() else 'UNKNOWN'} from data '{tab_data_str}': {e}")
            failed_count += 1

    final_msg = f"Session restore attempt complete. Restored: {restored_count}"
    if failed_count > 0: final_msg += f", Failed/Skipped: {failed_count}"
    final_msg += ". Ready."
    if status_bar_label: status_bar_label.config(text=final_msg)
    print(final_msg)
    # After restoring, update details panel based on the *last* restored/selected tab
    root.after(50, lambda: on_viewer_tab_changed()) # Short delay before updating details
    
# --- Database Functions (Keep As Is - No changes needed) ---
# Rename this function
def perform_search_worker(query, results_queue):
    """Worker function to perform search in a separate thread."""
    print(f"Search worker started for query: '{query}'")
    results_data = [] # Default to empty list
    error_occurred = False
    error_message = ""
    try:
        # Perform the combined search (includes FTS+metadata and rank)
        # This call might take time
        results_data = search_documents(query)

        # Also get snippets (this might also take time)
        # We need both ranked results and snippets for the final display
        print("Search worker fetching snippets...")
        snippet_results = search_content_with_snippets(query)
        snippet_map = {res[0]: (res[3], res[4]) for res in snippet_results} # doc_id -> (snippet, page)
        print(f"Search worker finished. Found {len(results_data)} ranked docs, {len(snippet_map)} snippets.")

    except Exception as e:
        print(f"!!! Error in search worker thread for query '{query}': {e}")
        error_occurred = True
        error_message = str(e)
        # results_data remains empty or potentially partially filled

    # Put results (or error indicator) into the queue
    # Include the original query for context when processing results
    results_queue.put({
        "query": query,
        "results_data": results_data,
        "snippet_map": snippet_map if not error_occurred else {}, # Pass snippet map too
        "error": error_occurred,
        "error_message": error_message
    })
    print("Search worker finished and put results in queue.")

def search_content_with_snippets(query):
    """
    Performs an FTS search for the query text and retrieves snippets
    and associated page number for the best match per document.
    Returns list of tuples: (doc_id, filename, filepath, best_snippet, page_number)
    """
    if not query: return []

    conn = sqlite3.connect(DATABASE_FILE)
    cursor = conn.cursor()
    results = []
    fts_query = query # Use the raw query

    try:
        # Use ROW_NUMBER() to get the highest ranked snippet/page per doc_id.
        cursor.execute(f"""
            WITH RankedMatches AS (
                SELECT
                    doc_id, page_number,
                    snippet(documents_fts, 2, '[', ']', '...', 15) as snippet_text,
                    rank, ROW_NUMBER() OVER(PARTITION BY doc_id ORDER BY rank) as rn
                FROM documents_fts WHERE documents_fts MATCH ?
            )
            SELECT rm.doc_id, d.filename, d.filepath, rm.snippet_text, rm.page_number
            FROM RankedMatches rm JOIN documents d ON rm.doc_id = d.id
            WHERE rm.rn = 1 ORDER BY rm.rank;
        """, (fts_query,))
        results = cursor.fetchall()
        print(f"Content search for '{fts_query}' found {len(results)} top snippets.")

    except sqlite3.OperationalError as e:
        print(f"FTS Error during content search for '{fts_query}': {e}")
        print("Falling back to simpler FTS query...")
        try: # Fallback without ROW_NUMBER
             cursor.execute(f"""
                 SELECT f.doc_id, d.filename, d.filepath,
                        snippet(documents_fts, 2, '[', ']', '...', 15), f.page_number
                 FROM documents_fts f JOIN documents d ON f.doc_id = d.id
                 WHERE f.documents_fts MATCH ? ORDER BY rank, f.doc_id, f.page_number
             """, (fts_query,))
             results = cursor.fetchall() # Might contain multiple results per doc
             print(f"Fallback FTS search found {len(results)} total snippets.")
             # Optional: Filter results in Python to keep only one per doc_id if needed
             # unique_results = {}
             # for r in results:
             #    if r[0] not in unique_results: unique_results[r[0]] = r
             # results = list(unique_results.values())

        except Exception as fallback_e:
             print(f"Fallback FTS query also failed: {fallback_e}")
             messagebox.showerror("Search Error", f"Full-text search failed.\nError: {e}")
    except sqlite3.Error as e:
        print(f"Database error during content search: {e}")
        messagebox.showerror("Database Error", f"Error searching content:\n{e}")
    finally:
        conn.close()
    return results

def init_db():
    """Initializes the SQLite database and tables if they don't exist."""
    conn = sqlite3.connect(DATABASE_FILE)
    cursor = conn.cursor()

    # --- Scan Paths Table ---
    # Replace placeholder with correct definition
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS scan_paths (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            path TEXT NOT NULL UNIQUE
        )
    ''')

    # --- Documents Table ---
    # Ensure this is the FULL definition with BME enhancements
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS documents (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            filename TEXT NOT NULL,
            filepath TEXT NOT NULL UNIQUE,
            manufacturer TEXT,
            device_model TEXT,
            document_type TEXT,
            keywords TEXT,
            last_modified REAL NOT NULL,
            revision_number TEXT,
            revision_date TEXT,
            status TEXT,
            applicable_models TEXT,
            associated_test_equipment TEXT
        )
    ''')

    # --- FTS5 Table for Full-Text Search ---
    cursor.execute('''
        CREATE VIRTUAL TABLE IF NOT EXISTS documents_fts USING fts5(
            doc_id UNINDEXED,
            page_number UNINDEXED,
            content
            -- Optionally add tokenize='porter'
        )
    ''')
    # Optional: Trigger to keep FTS table synced on documents delete
    cursor.execute('''
        CREATE TRIGGER IF NOT EXISTS documents_ad_trigger AFTER DELETE ON documents BEGIN
            DELETE FROM documents_fts WHERE doc_id=old.id;
        END;
    ''')

    # --- Links Table ---
    # Ensure this is the correct definition with description
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS links (
            link_id INTEGER PRIMARY KEY AUTOINCREMENT,
            source_doc_id INTEGER NOT NULL,
            target_doc_id INTEGER NOT NULL,
            description TEXT, -- Added description
            FOREIGN KEY (source_doc_id) REFERENCES documents (id) ON DELETE CASCADE,
            FOREIGN KEY (target_doc_id) REFERENCES documents (id) ON DELETE CASCADE,
            UNIQUE (source_doc_id, target_doc_id)
        )
    ''')

    # --- Notes Table ---
    # Ensure this is the correct definition with page_number
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS notes (
            note_id INTEGER PRIMARY KEY AUTOINCREMENT,
            doc_id INTEGER NOT NULL,
            page_number INTEGER,  -- Added page context
            note_text TEXT NOT NULL,
            created_timestamp REAL NOT NULL,
            FOREIGN KEY (doc_id) REFERENCES documents (id) ON DELETE CASCADE
        )
    ''')

    # --- Favorites Table ---
    # Ensure this is the correct definition
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS favorites (
            fav_id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL UNIQUE,
            doc_id INTEGER NOT NULL,
            page_number INTEGER NOT NULL,
            FOREIGN KEY (doc_id) REFERENCES documents (id) ON DELETE CASCADE
        )
    ''')

    # --- Indexes ---
    # Ensure all necessary indexes are created
    cursor.execute('CREATE INDEX IF NOT EXISTS idx_doc_filepath ON documents (filepath)')
    cursor.execute('CREATE INDEX IF NOT EXISTS idx_doc_filename ON documents (filename)')
    cursor.execute('CREATE INDEX IF NOT EXISTS idx_doc_manufacturer ON documents (manufacturer)')
    # ... (all other indexes from previous working versions) ...
    cursor.execute('CREATE INDEX IF NOT EXISTS idx_doc_model ON documents (device_model)')
    cursor.execute('CREATE INDEX IF NOT EXISTS idx_doc_type ON documents (document_type)')
    cursor.execute('CREATE INDEX IF NOT EXISTS idx_doc_revision ON documents (revision_number)')
    cursor.execute('CREATE INDEX IF NOT EXISTS idx_doc_status ON documents (status)')
    cursor.execute('CREATE INDEX IF NOT EXISTS idx_link_source ON links (source_doc_id)')
    cursor.execute('CREATE INDEX IF NOT EXISTS idx_link_target ON links (target_doc_id)')
    cursor.execute('CREATE INDEX IF NOT EXISTS idx_note_doc ON notes (doc_id)')
    cursor.execute('CREATE INDEX IF NOT EXISTS idx_fav_name ON favorites (name)')
    cursor.execute('CREATE INDEX IF NOT EXISTS idx_fav_doc ON favorites (doc_id)')


    conn.commit()
    conn.close()
    print("Database initialized/verified (All Tables).") # Updated print message

def add_favorite(name, doc_id, page_number):
    """Adds a new favorite bookmark to the database."""
    if not name:
        messagebox.showerror("Error", "Favorite name cannot be empty.")
        return False
    conn = sqlite3.connect(DATABASE_FILE)
    cursor = conn.cursor()
    try:
        cursor.execute("""
            INSERT INTO favorites (name, doc_id, page_number) VALUES (?, ?, ?)
        """, (name, doc_id, page_number))
        conn.commit()
        print(f"Favorite added: '{name}' -> Doc {doc_id}, Page {page_number+1}")
        return True
    except sqlite3.IntegrityError:
        messagebox.showerror("Error", f"A favorite with the name '{name}' already exists.")
        return False
    except sqlite3.Error as e:
        messagebox.showerror("Database Error", f"Could not add favorite:\n{e}")
        conn.rollback()
        return False
    finally:
        conn.close()

def get_favorites():
    """Retrieves all favorites ordered by name."""
    favorites = []
    conn = sqlite3.connect(DATABASE_FILE)
    cursor = conn.cursor()
    try:
        cursor.execute("""
            SELECT fav_id, name, doc_id, page_number FROM favorites ORDER BY name COLLATE NOCASE
        """)
        favorites = cursor.fetchall()
    except sqlite3.Error as e:
        print(f"Database error getting favorites: {e}")
    finally:
        conn.close()
    return favorites # Returns list of (fav_id, name, doc_id, page_number) tuples

def delete_favorite(fav_id):
    """Deletes a favorite by its ID."""
    conn = sqlite3.connect(DATABASE_FILE)
    cursor = conn.cursor()
    try:
        cursor.execute("DELETE FROM favorites WHERE fav_id = ?", (fav_id,))
        conn.commit()
        return cursor.rowcount > 0
    except sqlite3.Error as e:
        messagebox.showerror("Database Error", f"Could not delete favorite ID {fav_id}:\n{e}")
        conn.rollback()
        return False
    finally:
        conn.close()

def rename_favorite(fav_id, new_name):
    """Renames a favorite."""
    if not new_name:
        messagebox.showerror("Error", "New favorite name cannot be empty.")
        return False
    conn = sqlite3.connect(DATABASE_FILE)
    cursor = conn.cursor()
    try:
        cursor.execute("UPDATE favorites SET name = ? WHERE fav_id = ?", (new_name, fav_id))
        conn.commit()
        return True
    except sqlite3.IntegrityError:
        messagebox.showerror("Error", f"A favorite with the name '{new_name}' already exists.")
        return False
    except sqlite3.Error as e:
        messagebox.showerror("Database Error", f"Could not rename favorite ID {fav_id}:\n{e}")
        conn.rollback()
        return False
    finally:
        conn.close()

def get_scan_paths():
    """Retrieves the list of scan paths from the database."""
    paths = []
    conn = sqlite3.connect(DATABASE_FILE)
    cursor = conn.cursor()
    try:
        cursor.execute("SELECT path FROM scan_paths ORDER BY path")
        paths = [row[0] for row in cursor.fetchall()]
    except sqlite3.Error as e:
        print(f"Database error getting scan paths: {e}")
    finally:
        conn.close()
    return paths

def add_scan_path(path):
    """Adds a new scan path to the database."""
    conn = sqlite3.connect(DATABASE_FILE)
    cursor = conn.cursor()
    try:
        cursor.execute("INSERT INTO scan_paths (path) VALUES (?)", (path,))
        conn.commit()
        return True
    except sqlite3.IntegrityError:
        messagebox.showwarning("Path Exists", f"The path '{path}' is already in the list.")
        return False
    except sqlite3.Error as e:
        messagebox.showerror("Database Error", f"Failed to add path '{path}':\n{e}")
        return False
    finally:
        conn.close()

def remove_scan_path(path):
    """Removes a scan path from the database."""
    conn = sqlite3.connect(DATABASE_FILE)
    cursor = conn.cursor()
    try:
        cursor.execute("DELETE FROM scan_paths WHERE path = ?", (path,))
        conn.commit()
        if cursor.rowcount > 0:
             return True
        else:
             # Path might have been removed already
             return False
    except sqlite3.Error as e:
        messagebox.showerror("Database Error", f"Failed to remove path '{path}':\n{e}")
        return False
    finally:
        conn.close()


def extract_metadata_from_path(filepath):
    """ BASIC heuristic metadata extraction (same as before). """
    # ... (Keep the previous implementation) ...
    metadata = {'manufacturer': None, 'device_model': None, 'document_type': None}
    path_lower = filepath.lower()
    filename_lower = os.path.basename(path_lower)
    known_manufacturers = ['siemens', 'ge', 'philips', 'draeger', 'medtronic'] # Fallback list
    known_doc_types = ['manual', 'sop', 'datasheet', 'service', 'user', 'quick guide', 'pm', 'calibration', 'protocol']
    for manuf in known_manufacturers:
        if manuf in path_lower:
            metadata['manufacturer'] = manuf.title()
            break
    for dtype in known_doc_types:
        if re.search(r'\b' + re.escape(dtype) + r'\b', path_lower):
            metadata['document_type'] = dtype.title()
            match_model = re.search(r'([a-zA-Z0-9]+(?:[-_][a-zA-Z0-9]+)*)_{0}|{0}_([a-zA-Z0-9]+(?:[-_][a-zA-Z0-9]+)*)'.format(dtype), filename_lower)
            if match_model:
                 potential_model = match_model.group(1) or match_model.group(2)
                 if potential_model and len(potential_model) > 2: metadata['device_model'] = potential_model.upper()
            break
    if not metadata['device_model']:
         match_generic_model = re.search(r'\b([a-zA-Z]{2,6}[-_][a-zA-Z0-9]{2,8})\b', filename_lower)
         if match_generic_model: metadata['device_model'] = match_generic_model.group(1).upper()
    return metadata
def scan_and_update_worker(status_queue):
    """
    Worker function to perform scan/index/FTS in a background thread.
    Communicates status and results via the queue.
    DO NOT interact directly with Tkinter widgets from here.
    """
    scan_paths = []
    try:
        scan_paths = get_scan_paths() # Get paths within the thread
    except Exception as e:
         status_queue.put({'type': 'error', 'message': f"Failed to get scan paths: {e}"})
         return

    if not scan_paths:
        status_queue.put({'type': 'info', 'message': "No scan paths configured."})
        status_queue.put({'type': 'finished', 'added': 0, 'updated': 0, 'reindexed': 0, 'removed': 0, 'errors': 0, 'duration': 0})
        return

    conn = None
    added_count = 0; updated_count = 0; fts_reindexed_count = 0
    skipped_page_errors = 0; removed_count = 0
    scan_start_time = time.time()

    try:
        conn = sqlite3.connect(DATABASE_FILE)
        cursor = conn.cursor()
        existing_files = {row[0]: (row[1], row[2]) for row in cursor.execute('SELECT filepath, id, last_modified FROM documents')}
        found_paths_ids = {}
        status_queue.put({'type': 'status', 'message': "Starting incremental scan..."})
        print("[Worker] Starting incremental scan...") # Worker console log

        files_processed_total = 0

        for directory in scan_paths:
            if not os.path.isdir(directory):
                print(f"[Worker] Skipping inaccessible path: {directory}")
                status_queue.put({'type': 'status', 'message': f"Skipping: {directory[:50]}..."})
                continue

            # --- Prompting for Manufacturer CANNOT easily be done in worker thread ---
            # Option 1: Prompt *before* starting the thread (simpler).
            # Option 2: Use queue to ask main thread to prompt (more complex).
            # Let's stick to Option 1: Modify the trigger function to prompt first.
            # We need to pass the folder_manufacturer map to the worker.
            # For now, let's assume folder_manufacturer is handled *before* calling worker
            # Or remove the prompting feature from the threaded version for simplicity first.
            # ---> Let's REMOVE folder prompting for this threaded version first <---
            folder_manufacturer = None # Remove prompting from worker

            print(f"[Worker] Scanning: {directory}...")
            status_queue.put({'type': 'status', 'message': f"Scanning: {directory[:50]}..."})

            for root_dir, _, files in os.walk(directory):
                # Send periodic status update maybe based on folder?
                status_queue.put({'type': 'status', 'message': f"Scanning: ...{os.path.basename(root_dir)}"})

                for filename in files:
                    files_processed_total += 1
                    # Send occasional progress update (e.g., every 100 files)
                    if files_processed_total % 100 == 0:
                         status_queue.put({'type': 'progress', 'count': files_processed_total})


                    if filename.lower().endswith(SUPPORTED_EXTENSIONS) and not filename.startswith('~$'):
                        filepath = os.path.join(root_dir, filename)
                        doc_id = None
                        try:
                            current_last_modified = os.path.getmtime(filepath)
                            needs_db_processing = False; is_new_file = False

                            if filepath not in existing_files: needs_db_processing = True; is_new_file = True
                            else:
                                doc_id, db_last_modified = existing_files[filepath]
                                if current_last_modified > db_last_modified: needs_db_processing = True
                                else: found_paths_ids[filepath] = doc_id; continue # Unchanged

                            if needs_db_processing:
                                extracted_metadata = extract_metadata_from_path(filepath)
                                final_manufacturer = folder_manufacturer or extracted_metadata.get('manufacturer') # Using None for folder_manuf now
                                device_model = extracted_metadata.get('device_model')
                                document_type = extracted_metadata.get('document_type')

                                if is_new_file:
                                    cursor.execute('INSERT INTO documents (...) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)', (...)) # Full INSERT
                                    doc_id = cursor.lastrowid; added_count += 1
                                else: # Update
                                    cursor.execute('UPDATE documents SET ... WHERE id=?', (...)) # Full UPDATE
                                    updated_count += 1
                                existing_files[filepath] = (doc_id, current_last_modified)
                                found_paths_ids[filepath] = doc_id

                                # --- FTS Processing ---
                                print(f"[Worker] Processing FTS for Doc ID {doc_id} ({filename[:20]}...)")
                                # Put indexing status update
                                status_queue.put({'type': 'status', 'message': f"Indexing: {filename[:40]}..."})

                                cursor.execute("DELETE FROM documents_fts WHERE doc_id = ?", (doc_id,))
                                pages_indexed_this_file = 0; file_ext = os.path.splitext(filename)[1].lower()
                                try: # Page-by-page extraction/indexing
                                    if file_ext == '.pdf':
                                         # ... PDF FTS logic ...
                                         pass
                                    elif file_ext == '.docx' and DOCX_ENABLED:
                                         # ... DOCX FTS logic ...
                                         pass
                                    elif file_ext == '.txt':
                                         # ... TXT FTS logic ...
                                         pass
                                    elif file_ext in ['.html', '.htm']:
                                         # ... HTML FTS logic ...
                                         pass

                                    if pages_indexed_this_file > 0: fts_reindexed_count += 1
                                    # No need to print here, status sent via queue

                                except Exception as text_ex:
                                    print(f"[Worker] !!! Text/FTS error for {filename}: {text_ex}")
                                    skipped_page_errors += 1

                        except Exception as e: print(f"[Worker] Error processing file {filepath}: {e}")

        # --- End Scan Loop ---
        scan_duration = time.time() - scan_start_time
        print(f"[Worker] Scan loop finished in {scan_duration:.2f}s.")
        status_queue.put({'type': 'status', 'message': "Removing obsolete entries..."})

        # --- Remove obsolete entries ---
        paths_in_db = set(existing_files.keys())
        paths_found_on_disk_this_scan = set(found_paths_ids.keys())
        paths_to_remove = paths_in_db - paths_found_on_disk_this_scan
        if paths_to_remove:
            ids_to_remove = [existing_files[p][0] for p in paths_to_remove if p in existing_files]
            if ids_to_remove:
                 cursor.executemany('DELETE FROM documents WHERE id = ?', [(id,) for id in ids_to_remove])
                 removed_count = cursor.rowcount
                 print(f"[Worker] Removed {removed_count} obsolete documents.")

        conn.commit()
        print("[Worker] DB commit successful.")

        # --- Optimize FTS ---
        try:
             status_queue.put({'type': 'status', 'message': "Optimizing index..."})
             print("[Worker] Optimizing FTS index...")
             cursor.execute("INSERT INTO documents_fts(documents_fts) VALUES('optimize');")
             conn.commit()
             print("[Worker] Optimization complete.")
        except Exception as opt_e: print(f"[Worker] DB optimize error: {opt_e}")

        # --- Put final result on queue ---
        status_queue.put({
            'type': 'finished', 'added': added_count, 'updated': updated_count,
            'reindexed': fts_reindexed_count, 'removed': removed_count,
            'errors': skipped_page_errors, 'duration': scan_duration
        })

    except sqlite3.Error as e:
         print(f"[Worker] !!! DB error during scan: {e}")
         if conn: conn.rollback()
         status_queue.put({'type': 'error', 'message': f"Scan DB error:\n{e}"})
    except Exception as e:
         print(f"[Worker] !!! Unexpected scan error: {e}")
         if conn: conn.rollback()
         status_queue.put({'type': 'error', 'message': f"Unexpected scan error:\n{e}"})
    finally:
        if conn: conn.close()
        print("[Worker] Thread finished.")

def scan_and_update_index():
    """
    Scans configured directories, prompts for manufacturer, extracts text PAGE BY PAGE,
    and incrementally updates the database and FTS index based on modification times.
    Shows progress.
    """
    global status_bar_label, root, scan_progress_bar # Ensure progress bar is global

    scan_paths = get_scan_paths()
    if not scan_paths:
        messagebox.showinfo("Scan", "No scan paths configured. Please add paths via 'File -> Manage Scan Paths...'.")
        return

    conn = None
    final_msg = "Scan Aborted. Ready." # Default message if errors occur early
    try:
        conn = sqlite3.connect(DATABASE_FILE)
        cursor = conn.cursor()
        added_count = 0
        updated_count = 0 # Counts metadata updates
        fts_reindexed_count = 0 # Counts files whose FTS was fully re-indexed
        skipped_page_errors = 0

        # Get existing files: filepath -> (id, last_modified)
        existing_files = {row[0]: (row[1], row[2]) for row in cursor.execute('SELECT filepath, id, last_modified FROM documents')}
        print(f"Loaded {len(existing_files)} existing file records from DB.")
        # Keep track of files found on disk this scan to detect deletions
        # Map filepath -> id (for both existing and newly added files)
        found_paths_ids = {}

        print("Starting incremental scan...")
        status_bar_label.config(text="Starting incremental scan...")
        # --- Show and Start Progress Bar ---
        if scan_progress_bar:
             scan_progress_bar.pack(side=tk.BOTTOM, fill=tk.X, padx=1, pady=0, before=status_bar_label)
             scan_progress_bar.start(10) # Start animation (interval in ms)
        root.update_idletasks()
        # --- End Progress Bar Start ---
        scan_start_time = time.time()
        files_processed_since_update = 0

        # --- Loop through each configured scan path ---
        for directory in scan_paths:
            if not os.path.isdir(directory):
                print(f"Warning: Configured scan path not found or inaccessible: {directory}")
                status_bar_label.config(text=f"Skipping inaccessible path: {directory[:50]}...")
                root.update_idletasks()
                continue

            # --- Prompt for Folder-Level Manufacturer ---
            status_bar_label.config(text=f"Prompting for: {directory[:50]}...")
            root.update_idletasks()
            prompt_title = "Assign Manufacturer"
            prompt_text = f"Enter default manufacturer for files found within:\n'{directory}'\n\n(Leave blank or Cancel to skip)"
            folder_manufacturer = simpledialog.askstring(prompt_title, prompt_text, parent=root)
            if folder_manufacturer is not None: # Check if Cancel was pressed
                folder_manufacturer = folder_manufacturer.strip() or None # Treat empty as None
                if folder_manufacturer: print(f"Assigning Manufacturer '{folder_manufacturer}' to files in {directory}")
                else: print(f"No default manufacturer assigned for {directory}.")
            else: # User pressed Cancel
                folder_manufacturer = None
                print(f"Manufacturer assignment cancelled for {directory}.")


            status_bar_label.config(text=f"Scanning: {directory[:50]}...")
            root.update_idletasks()
            print(f"Scanning: {directory}...")

            # --- Walk through the current directory ---
            for root_dir, _, files in os.walk(directory):
                status_bar_label.config(text=f"Scanning: ...{os.path.basename(root_dir)}")
                for filename in files:
                    # --- Periodic GUI Update ---
                    files_processed_since_update += 1
                    if files_processed_since_update >= 50: root.update_idletasks(); files_processed_since_update = 0

                    if filename.lower().endswith(SUPPORTED_EXTENSIONS) and not filename.startswith('~$'):
                        filepath = os.path.join(root_dir, filename)
                        # found_paths_on_disk.add(filepath) # We use found_paths_ids instead now
                        doc_id = None

                        try:
                            current_last_modified = os.path.getmtime(filepath)
                            needs_db_processing = False # Flag if DB needs any write operation

                            # --- Determine File Status ---
                            if filepath not in existing_files:
                                # Case 1: New File
                                print(f" + Found new file: {filename}")
                                needs_db_processing = True
                                is_new_file = True
                                # doc_id will be assigned after insert
                            else:
                                # Case 2: Existing File - Check modification time
                                doc_id, db_last_modified = existing_files[filepath]
                                is_new_file = False
                                if current_last_modified > db_last_modified:
                                    print(f" * Found updated file: {filename} (ID: {doc_id})")
                                    needs_db_processing = True
                                else:
                                    # File exists and hasn't changed, store its ID and skip DB processing
                                    found_paths_ids[filepath] = doc_id # <<< RECORD ID FOR UNCHANGED FILE
                                    continue # Skip to next file in loop

                            # --- Process if New or Updated ---
                            if needs_db_processing:
                                # Prepare metadata
                                extracted_metadata = extract_metadata_from_path(filepath)
                                final_manufacturer = folder_manufacturer or extracted_metadata.get('manufacturer')
                                device_model = extracted_metadata.get('device_model')
                                document_type = extracted_metadata.get('document_type')

                                if is_new_file:
                                    cursor.execute('''INSERT INTO documents (filename, filepath, manufacturer, device_model, document_type, keywords, last_modified, revision_number, revision_date, status, applicable_models, associated_test_equipment) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)''',
                                                   (filename, filepath, final_manufacturer, device_model, document_type, None, current_last_modified, None,None,None,None,None))
                                    doc_id = cursor.lastrowid; added_count += 1
                                else: # File updated
                                    cursor.execute('UPDATE documents SET filename=?, manufacturer = CASE WHEN ? IS NOT NULL THEN ? ELSE COALESCE(documents.manufacturer, ?) END, device_model=COALESCE(documents.device_model, ?), document_type=COALESCE(documents.document_type, ?), last_modified=? WHERE id=?',
                                                   (filename, folder_manufacturer, folder_manufacturer, final_manufacturer, device_model, document_type, current_last_modified, doc_id))
                                    updated_count += 1
                                # Update cache/tracking dictionary immediately after insert/update
                                existing_files[filepath] = (doc_id, current_last_modified) # Update cache with new time/id
                                found_paths_ids[filepath] = doc_id # Store ID for this processed file


                                # --- Process FTS Indexing/Re-indexing ---
                                print(f"   - Processing FTS for Doc ID {doc_id}...")
                                status_bar_label.config(text=f"Indexing: {filename[:40]}...")
                                root.update_idletasks()
                                cursor.execute("DELETE FROM documents_fts WHERE doc_id = ?", (doc_id,))
                                pages_indexed_this_file = 0; file_ext = os.path.splitext(filename)[1].lower()
                                try:
                                    if file_ext == '.pdf':
                                        with fitz.open(filepath) as doc:
                                            for page_num, page in enumerate(doc):
                                                page_text = page.get_text("text", sort=True)
                                                if page_text and page_text.strip():
                                                    cursor.execute("INSERT INTO documents_fts (doc_id, page_number, content) VALUES (?, ?, ?)",(doc_id, page_num, page_text))
                                                    pages_indexed_this_file += 1
                                    elif file_ext == '.docx' and DOCX_ENABLED:
                                         doc = Document(filepath); full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                                         if full_text: cursor.execute("INSERT INTO documents_fts (doc_id, page_number, content) VALUES (?, ?, ?)", (doc_id, 0, full_text)); pages_indexed_this_file = 1
                                    elif file_ext == '.txt':
                                         content = None;
                                         for enc in ['utf-8', 'cp1252', 'latin-1']:
                                             try:
                                                  with open(filepath, 'r', encoding=enc) as f: content = f.read(); break
                                             except: continue
                                         if content and content.strip(): cursor.execute("INSERT INTO documents_fts (doc_id, page_number, content) VALUES (?, ?, ?)", (doc_id, 0, content)); pages_indexed_this_file = 1
                                    elif file_ext in ['.html', '.htm']:
                                        html_content = ""; extracted_text = ""
                                        try:
                                             for enc in ['utf-8', 'cp1252', 'latin-1']:
                                                  try:
                                                       with open(filepath, 'r', encoding=enc) as f: html_content = f.read(); break
                                                  except UnicodeDecodeError: continue
                                             if not html_content: raise ValueError("Could not decode HTML")
                                             html_content = re.sub(r'<script.*?<\/script>','',html_content,flags=re.I|re.S); html_content = re.sub(r'<style.*?<\/style>','',html_content,flags=re.I|re.S)
                                             extracted_text = re.sub(r'<.*?>',' ',html_content); extracted_text = re.sub(r'\s+',' ',extracted_text).strip()
                                             if extracted_text: cursor.execute("INSERT INTO documents_fts (doc_id, page_number, content) VALUES (?, ?, ?)",(doc_id, 0, extracted_text)); pages_indexed_this_file = 1
                                        except Exception as html_ex: print(f"HTML Err: {html_ex}"); skipped_page_errors += 1

                                    if pages_indexed_this_file > 0: print(f"     - Re-Indexed {pages_indexed_this_file} pages/blocks."); fts_reindexed_count += 1
                                    else: print(f"     - No indexable text found.")
                                except Exception as text_ex: print(f"   - !!! Error during text extraction/FTS indexing for {filename}: {text_ex}"); skipped_page_errors += 1

                        except OSError as e: print(f"OS Error processing file {filepath}: {e}")
                        except Exception as e: print(f"General Error processing file {filepath}: {e}")

        # --- End Main Scan Loop ---
        scan_duration = time.time() - scan_start_time
        print(f"Directory scanning finished in {scan_duration:.2f} seconds.")
        status_bar_label.config(text="Removing obsolete entries...")
        root.update_idletasks()

        # --- Remove obsolete entries ---
        removed_count = 0
        paths_in_db = set(existing_files.keys())
        # Use the keys from found_paths_ids which includes new, updated, AND unchanged files found
        paths_found_on_disk_this_scan = set(found_paths_ids.keys())
        paths_to_remove = paths_in_db - paths_found_on_disk_this_scan
        if paths_to_remove:
            ids_to_remove = [existing_files[p][0] for p in paths_to_remove if p in existing_files]
            if ids_to_remove:
                 cursor.executemany('DELETE FROM documents WHERE id = ?', [(id,) for id in ids_to_remove])
                 removed_count = cursor.rowcount
                 print(f"Removed {removed_count} obsolete document entries.")

        conn.commit()
        print("Database transaction committed.")
        final_msg = f"Scan Complete ({scan_duration:.1f}s). Added: {added_count}, Updated: {updated_count}, Re-Indexed: {fts_reindexed_count}, Removed: {removed_count}."
        if skipped_page_errors > 0: final_msg += f" Text Errors: {skipped_page_errors}."
        final_msg += " Ready."
        messagebox.showinfo("Scan Complete", f"Scan finished in {scan_duration:.1f} seconds.\nDocs Added: {added_count}\nDocs Updated: {updated_count}\nFiles Re-Indexed(FTS): {fts_reindexed_count}\nDocs Removed: {removed_count}\nText Extraction Errors: {skipped_page_errors}")

    except sqlite3.Error as e: print(f"DB error during scan: {e}"); messagebox.showerror("DB Error", f"Scan DB error:\n{e}"); conn.rollback() if conn else None; final_msg="Scan failed! DB error."
    except Exception as e: print(f"Unexpected scan error: {e}"); messagebox.showerror("Scan Error", f"Unexpected scan error:\n{e}"); conn.rollback() if conn else None; final_msg="Scan failed! Unexpected error."
    finally:
        if scan_progress_bar: scan_progress_bar.stop(); scan_progress_bar.pack_forget()
        if conn:
             try:
                  print("Optimizing FTS index..."); status_bar_label.config(text="Optimizing index..."); root.update_idletasks()
                  opt_cursor = conn.cursor(); opt_cursor.execute("INSERT INTO documents_fts(documents_fts) VALUES('optimize');"); conn.commit(); print("Optimization complete.")
             except Exception as opt_e: print(f"DB optimize error: {opt_e}")
             conn.close()
        build_file_tree(); clear_details_panel(); status_bar_label.config(text=final_msg)
        
def get_document_details(doc_id):
    """Retrieves all details for a single document by its ID."""
    if not doc_id: return None
    conn = sqlite3.connect(DATABASE_FILE)
    # Ensure connection has text_factory set to str if needed (usually default)
    # conn.text_factory = str
    cursor = conn.cursor()
    details = None
    try:
        cursor.execute("""
            SELECT id, filename, filepath, manufacturer, device_model,
                   document_type, keywords, revision_number, revision_date,
                   status, applicable_models, associated_test_equipment
            FROM documents WHERE id = ?
            """, (doc_id,))
        details = cursor.fetchone()
    except sqlite3.Error as e:
        print(f"Error getting document details for ID {doc_id}: {e}")
    finally:
        conn.close()
    return details


def get_linked_documents(doc_id):
    """Retrieves linked documents WITH description."""
    conn = sqlite3.connect(DATABASE_FILE)
    cursor = conn.cursor()
    linked_docs = []
    try:
        cursor.execute('''
            SELECT d.id, d.filename, d.filepath, l.description -- Select description
            FROM links l
            JOIN documents d ON l.target_doc_id = d.id
            WHERE l.source_doc_id = ?
            ORDER BY d.filename
        ''', (doc_id,))
        linked_docs = cursor.fetchall()
    except sqlite3.Error as e:
         print(f"Error getting linked documents for {doc_id}: {e}")
    finally:
        conn.close()
    return linked_docs


def update_document_metadata(doc_id, metadata):
    """Updates the metadata for a specific document ID."""
    conn = sqlite3.connect(DATABASE_FILE)
    cursor = conn.cursor()
    success = False
    try:
        cursor.execute('''
            UPDATE documents SET
                manufacturer = ?,
                device_model = ?,
                document_type = ?,
                keywords = ?,
                revision_number = ?,
                revision_date = ?,
                status = ?,
                applicable_models = ?,
                associated_test_equipment = ?
            WHERE id = ?
        ''', (
            metadata.get('manufacturer'), metadata.get('device_model'),
            metadata.get('document_type'), metadata.get('keywords'),
            metadata.get('revision_number'), metadata.get('revision_date'), # New
            metadata.get('status'), metadata.get('applicable_models'),    # New
            metadata.get('associated_test_equipment'),                     # New
            doc_id
        ))
        conn.commit()
        print(f"Metadata updated for doc ID: {doc_id}")
        success = True
    except sqlite3.Error as e:
        print(f"Database error updating metadata for doc ID {doc_id}: {e}")
        conn.rollback()
    finally:
        conn.close()
    return success


def add_document_link(source_id, target_id, description):
    """Adds a link WITH description."""
    if source_id == target_id:
         messagebox.showwarning("Link Error", "Cannot link a document to itself.")
         return False
    conn = sqlite3.connect(DATABASE_FILE)
    cursor = conn.cursor()
    try:
        cursor.execute("INSERT INTO links (source_doc_id, target_doc_id, description) VALUES (?, ?, ?)",
                       (source_id, target_id, description))
        conn.commit()
        print(f"Link added: {source_id} -> {target_id} ('{description}')")
        return True
    except sqlite3.IntegrityError:
        messagebox.showinfo("Link Exists", "This link already exists.")
        return False
    except sqlite3.Error as e:
        print(f"DB Error adding link {source_id}->{target_id}: {e}")
        conn.rollback()
        messagebox.showerror("Database Error", f"Could not add link:\n{e}")
        return False
    finally:
        conn.close()


def remove_document_link(source_id, target_id):
     """Removes a link (no change needed)."""
     # ... (Same as previous) ...
     conn = sqlite3.connect(DATABASE_FILE)
     cursor = conn.cursor()
     try:
         cursor.execute("DELETE FROM links WHERE source_doc_id = ? AND target_doc_id = ?", (source_id, target_id))
         conn.commit()
         return cursor.rowcount > 0
     except sqlite3.Error as e: print(f"DB Error removing link {source_id}->{target_id}: {e}"); conn.rollback(); return False
     finally: conn.close()


def get_notes_for_document(doc_id):
    """Retrieves notes for a specific document ID."""
    notes = []
    conn = sqlite3.connect(DATABASE_FILE)
    cursor = conn.cursor()
    try:
        cursor.execute("""
            SELECT note_id, note_text, created_timestamp FROM notes
            WHERE doc_id = ? ORDER BY created_timestamp DESC
        """, (doc_id,))
        notes = cursor.fetchall()
    except sqlite3.Error as e:
        print(f"Database error getting notes for doc {doc_id}: {e}")
    finally:
        conn.close()
    return notes


def add_note_for_document(doc_id, note_text):
    """Adds a new note for a document."""
    if not note_text: return False
    conn = sqlite3.connect(DATABASE_FILE)
    cursor = conn.cursor()
    try:
        timestamp = time.time()
        cursor.execute("INSERT INTO notes (doc_id, note_text, created_timestamp) VALUES (?, ?, ?)",
                       (doc_id, note_text, timestamp))
        conn.commit()
        return True
    except sqlite3.Error as e:
        print(f"Database error adding note for doc {doc_id}: {e}")
        conn.rollback()
        return False
    finally:
        conn.close()

def delete_note(note_id):
    """Deletes a specific note by its ID."""
    conn = sqlite3.connect(DATABASE_FILE)
    cursor = conn.cursor()
    try:
        cursor.execute("DELETE FROM notes WHERE note_id = ?", (note_id,))
        conn.commit()
        return cursor.rowcount > 0
    except sqlite3.Error as e:
        print(f"Database error deleting note {note_id}: {e}")
        conn.rollback()
        return False
    finally:
        conn.close()
# Placeholder comment - Ensure the full functions from the previous version are here
# --- End Database Functions Placeholder ---


# --- Modified Search Function (for Rank) ---
def search_documents(query):
    """
    Searches metadata AND full-text index.
    Returns list of document detail tuples ORDERED potentially by FTS Rank.
    If query is empty, returns ALL documents ordered by filename.
    """
    conn = sqlite3.connect(DATABASE_FILE)
    cursor = conn.cursor()

    if not query:
        # ... (return all documents logic - same as before) ...
        try:
            cursor.execute('SELECT id, filename, filepath, manufacturer, device_model, document_type FROM documents ORDER BY filename'); results = cursor.fetchall(); conn.close(); return results
        except sqlite3.Error as e: print(f"DB error fetching all docs: {e}"); conn.close(); return []

    # --- Logic for non-empty query ---
    matching_doc_ids_ranked = {} # Store ID -> rank (rank is lower for better match)
    search_term_meta = f"%{query}%"
    search_term_fts = query

    try:
        # 1. Metadata Search (doesn't provide rank, assign default low relevance rank)
        cursor.execute('''
            SELECT id FROM documents
            WHERE filename LIKE ? OR filepath LIKE ? OR manufacturer LIKE ?
               OR device_model LIKE ? OR document_type LIKE ? OR keywords LIKE ?
        ''', (search_term_meta, search_term_meta, search_term_meta,
              search_term_meta, search_term_meta, search_term_meta))
        for row in cursor.fetchall():
            if row[0] not in matching_doc_ids_ranked: # Avoid overwriting potential FTS rank
                matching_doc_ids_ranked[row[0]] = 9999 # Assign low relevance

        # 2. Full-Text Search
        try:
            # Select doc_id and rank. Lower rank values indicate better matches in SQLite FTS5.
            cursor.execute('''
                SELECT doc_id, rank FROM documents_fts WHERE documents_fts MATCH ? ORDER BY rank
            ''', (search_term_fts,))
            for doc_id, rank in cursor.fetchall():
                matching_doc_ids_ranked[doc_id] = rank # Overwrite/add with actual FTS rank
            print(f"FTS search updated ranks for {len(matching_doc_ids_ranked)} matches.")

        except sqlite3.OperationalError as fts_e:
             print(f"FTS search failed for '{query}': {fts_e}. Searching metadata only.")

    except sqlite3.Error as e:
        print(f"Database error during search for '{query}': {e}")
        conn.close(); return []

    # --- Fetch details for unique matching IDs, ordered by rank then filename ---
    if not matching_doc_ids_ranked:
        conn.close(); return []

    # Sort IDs based on rank (lower is better), then use filename as tie-breaker later
    sorted_ids = sorted(matching_doc_ids_ranked.keys(), key=lambda doc_id: matching_doc_ids_ranked[doc_id])

    try:
         placeholders = ','.join('?' * len(sorted_ids))
         # Preserve the rank-based order using a trick with INSTR in ORDER BY
         # This ensures the results are returned in the desired rank order.
         # Alternatively, fetch all then re-sort in Python using the ranks dict.
         sql = f'''
            SELECT id, filename, filepath, manufacturer, device_model, document_type
            FROM documents
            WHERE id IN ({placeholders})
            ORDER BY INSTR(?, ',' || id || ',') -- Order by position in sorted_ids list
         '''
         # Create the ordered ID string for INSTR: ",id1,id2,id3,"
         ordered_id_string = ',' + ','.join(map(str, sorted_ids)) + ','
         cursor.execute(sql, list(sorted_ids) + [ordered_id_string]) # Pass IDs twice for IN and INSTR
         results = cursor.fetchall()
         print(f"Fetched details for {len(results)} matching documents (ranked).")
         conn.close()
         return results
    except sqlite3.Error as e:
         print(f"Database error fetching final ranked results: {e}")
         conn.close(); return []


# --- Utility & Helper Functions (Keep As Is) ---
def open_file_externally_selected():
    """Opens the currently selected file externally."""
    # ... (Same as previous) ...
    global results_tree
    selected_items = results_tree.selection()
    if not selected_items: messagebox.showinfo("Open Externally", "Select a document first."); return
    doc_id = selected_items[0]
    details = get_document_details(doc_id)
    if details: open_file_externally(details[2])
    else: messagebox.showerror("Error", f"Could not get details for ID {doc_id}.")


def open_file_externally(filepath):
    """Opens the given file using the default system application."""
    # ... (Same as previous) ...
    try:
        if not os.path.exists(filepath): messagebox.showerror("Error", f"File not found:\n{filepath}"); return
        if platform.system() == "Windows": os.startfile(filepath)
        elif platform.system() == "Darwin": subprocess.run(['open', filepath], check=True)
        else: subprocess.run(['xdg-open', filepath], check=True)
    except Exception as e: messagebox.showerror("Error", f"Could not open file externally:\n{filepath}\n\nError: {e}")


def show_about():
     messagebox.showinfo("About BME Document Navigator", "BME Document Navigator v3.0\n\nFeatures: FTS, Notes, Outline, Links+, Config Paths.\nBuilt with Python & Tkinter.")

def do_nothing(): # Placeholder for menu items or disabled buttons
    print("Action not implemented yet.")

def set_theme(theme_name):
     """Sets the ttk theme."""
     # ... (Same as previous) ...
     try:
         style = ttk.Style()
         style.theme_use(theme_name)
         print(f"Theme set to: {theme_name}")
     except tk.TclError: messagebox.showerror("Theme Error", f"Could not apply theme: {theme_name}")

# Placeholder comment - Ensure the full functions from the previous version are here
# --- End Utility Functions Placeholder ---


# --- Viewer & Tab Functions (Keep mostly As Is) ---
def get_active_tab_id():
    """Returns the ID (widget name) of the currently selected tab."""
    global viewer_notebook
    if not viewer_notebook: return None
    try:
        return viewer_notebook.select()
    except tk.TclError:
        return None
def get_active_tab_state():
    """Returns the state dictionary for the currently active tab."""
    tab_id = get_active_tab_id()
    return tab_states.get(tab_id) if tab_id else None
def update_page_label_for_tab(tab_id):
    """Updates the page/zoom label for the controls within a specific tab."""
    state = tab_states.get(tab_id)
    if state and state.get('doc_obj') and 'widgets' in state:
        doc = state['doc_obj']
        page_num = state['page_num']
        zoom = state['zoom']
        label_widget = state['widgets'].get('page_label')
        if label_widget:
            if 0 <= page_num < len(doc):
                 page_info = f"Page {page_num + 1} of {len(doc)} (Zoom: {zoom:.2f}x)"
            else:
                 page_info = f"Page ? of {len(doc)} (Zoom: {zoom:.2f}x)" # Should not happen
            label_widget.config(text=page_info)
    elif state and 'widgets' in state:
         label_widget = state['widgets'].get('page_label')
         if label_widget: label_widget.config(text="") # Clear if no doc

def load_pdf_page(tab_id):
    """Renders and displays a specific PDF page in the specified tab."""
    state = tab_states.get(tab_id)
    if not state or not state.get('doc_obj'):
        print(f"Error: No valid state or document object for tab {tab_id}")
        return

    doc = state['doc_obj']
    page_number = state['page_num']
    zoom = state['zoom']
    canvas_widget = state['widgets'].get('canvas')
    text_widget = state['widgets'].get('text_viewer') # To display errors
    canvas_frame = state['widgets'].get('canvas_frame')
    text_frame = state['widgets'].get('text_frame')

    if not canvas_widget or not text_widget or not canvas_frame or not text_frame:
         print(f"Error: Missing widgets in state for tab {tab_id}")
         return

    # Ensure canvas is visible
    if not canvas_frame.winfo_ismapped(): # Check if it's packed
        canvas_frame.pack(expand=True, fill=tk.BOTH)
        text_frame.pack_forget()

    # Boundary check
    if not (0 <= page_number < len(doc)):
        print(f"Error: Invalid page number {page_number} for doc in tab {tab_id}")
        return

    update_page_label_for_tab(tab_id) # Update label before rendering

    try:
        page = doc.load_page(page_number)
        mat = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=mat, alpha=False)

        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        tk_img = ImageTk.PhotoImage(img)
        state['tk_img_ref'] = tk_img # Store reference in state

        canvas_widget.delete("all")
        canvas_widget.create_image(0, 0, anchor=tk.NW, image=tk_img)
        # Update scrollregion to match the new image size
        canvas_widget.config(scrollregion=(0, 0, pix.width, pix.height))
        update_back_forward_buttons(tab_id)

    except Exception as e:
        error_message = f"Error rendering page {page_number + 1}:\n{e}"
        print(f"!!! Tab {tab_id}: {error_message}")
        # Display error in this tab's text widget
        text_widget.config(state=tk.NORMAL)
        text_widget.delete('1.0', tk.END)
        text_widget.insert('1.0', error_message)
        text_widget.config(state=tk.DISABLED)
        # Switch view to text
        text_frame.pack(expand=True, fill=tk.BOTH)
        canvas_frame.pack_forget()
def display_text_in_tab(tab_id, text_content):
    """Displays plain text in the specified tab's text widget."""
    state = tab_states.get(tab_id)
    if not state or 'widgets' not in state: return

    text_widget = state['widgets'].get('text_viewer')
    canvas_frame = state['widgets'].get('canvas_frame')
    text_frame = state['widgets'].get('text_frame')
    pdf_nav_frame = state['widgets'].get('pdf_nav_frame')

    if not text_widget or not canvas_frame or not text_frame or not pdf_nav_frame: return

    text_widget.config(state=tk.NORMAL)
    text_widget.delete('1.0', tk.END)
    text_widget.insert('1.0', text_content)
    text_widget.config(state=tk.DISABLED)

    # Switch view to Text widget, hide PDF nav
    if not text_frame.winfo_ismapped(): # Check if it's packed
         text_frame.pack(expand=True, fill=tk.BOTH)
         canvas_frame.pack_forget()
    # Always ensure PDF nav is hidden for text
    pdf_nav_frame.pack_forget()
    update_page_label_for_tab(tab_id) # Clear label
def update_history_and_load(tab_id, new_page_num, is_history_navigation=False):
    """Loads a new page, updating history if it's not a back/forward action."""
    state = tab_states.get(tab_id)
    if not state or 'page_num' not in state or 'history_back' not in state or 'history_forward' not in state:
        print(f"Error: Invalid state for tab {tab_id} in update_history_and_load")
        return

    current_page = state['page_num']

    # Only update history if it's a new navigation action (not Back/Forward itself)
    # and if the page actually changes
    if not is_history_navigation and new_page_num != current_page:
        # Add current page to back history
        state['history_back'].append(current_page)
        # Clear forward history whenever a new navigation path is taken
        state['history_forward'].clear()
        print(f"Tab {tab_id}: History updated. Back stack: {state['history_back']}, Fwd stack cleared.")

    # Update current page number in state
    state['page_num'] = new_page_num

    # Load the requested page visually
    load_pdf_page(tab_id) # load_pdf_page just renders state['page_num']

    # Update the state of Back/Forward buttons after navigation
    update_back_forward_buttons(tab_id)
def go_back():
    """Navigates back using the history stack."""
    tab_id = get_active_tab_id()
    state = get_active_tab_state()
    if state and state['history_back']: # Check if back history exists
        # Pop previous page from back stack
        prev_page_num = state['history_back'].pop()
        # Add current page to forward stack
        state['history_forward'].append(state['page_num'])
        print(f"Tab {tab_id}: Going Back. Popped {prev_page_num} from back. Pushed {state['page_num']} to fwd.")
        # Navigate using helper, marking it as history nav
        update_history_and_load(tab_id, prev_page_num, is_history_navigation=True)
    else:
        print("Cannot go back, history empty.")

def go_forward():
    """Navigates forward using the history stack."""
    tab_id = get_active_tab_id()
    state = get_active_tab_state()
    if state and state['history_forward']: # Check if forward history exists
        # Pop next page from forward stack
        next_page_num = state['history_forward'].pop()
        # Add current page to back stack
        state['history_back'].append(state['page_num'])
        print(f"Tab {tab_id}: Going Forward. Popped {next_page_num} from fwd. Pushed {state['page_num']} to back.")
        # Navigate using helper, marking it as history nav
        update_history_and_load(tab_id, next_page_num, is_history_navigation=True)
    else:
        print("Cannot go forward, history empty.")

def update_back_forward_buttons(tab_id):
     """Enables/disables Back/Forward buttons based on history stack state."""
     state = tab_states.get(tab_id)
     if not state or 'widgets' not in state: return

     back_button = state['widgets'].get('back_button')
     forward_button = state['widgets'].get('forward_button')

     if back_button:
          back_button.config(state=tk.NORMAL if state.get('history_back') else tk.DISABLED)
     if forward_button:
          forward_button.config(state=tk.NORMAL if state.get('history_forward') else tk.DISABLED)

def next_page():
    """Goes to the next page, updating history."""
    tab_id = get_active_tab_id()
    state = get_active_tab_state()
    if state and state.get('doc_obj'):
        doc = state['doc_obj']
        current_page = state['page_num']
        if current_page < len(doc) - 1:
            # Call helper to handle history and load
            update_history_and_load(tab_id, current_page + 1)

def prev_page():
    """Goes to the previous page, updating history."""
    tab_id = get_active_tab_id()
    state = get_active_tab_state()
    if state and state.get('doc_obj'):
        current_page = state['page_num']
        if current_page > 0:
            # Call helper to handle history and load
            update_history_and_load(tab_id, current_page - 1)

def zoom_in():
    tab_id = get_active_tab_id()
    state = get_active_tab_state()
    if state and state.get('doc_obj'):
        new_zoom = min(MAX_ZOOM, state['zoom'] + ZOOM_STEP)
        if abs(new_zoom - state['zoom']) > 0.01:
            state['zoom'] = new_zoom
            load_pdf_page(tab_id)

def zoom_out():
    tab_id = get_active_tab_id()
    state = get_active_tab_state()
    if state and state.get('doc_obj'):
        new_zoom = max(MIN_ZOOM, state['zoom'] - ZOOM_STEP)
        if abs(new_zoom - state['zoom']) > 0.01:
            state['zoom'] = new_zoom
            load_pdf_page(tab_id)

def handle_scroll_or_zoom(event):
    """Handles mouse wheel scrolling OR zooming (if Ctrl is pressed) on the PDF canvas of the ACTIVE tab."""
    tab_id = get_active_tab_id()
    state = get_active_tab_state()
    if not state or not state.get('doc_obj') or 'widgets' not in state:
        return # Only act if a PDF is loaded in the active tab

    canvas_widget = state['widgets'].get('canvas')
    # Check if the event occurred directly on the target canvas widget
    if not canvas_widget or event.widget != canvas_widget:
         return

    ctrl_pressed = (event.state & 0x0004) != 0
    scroll_direction = 0
    zoom_in_flag = False
    zoom_out_flag = False

    # Platform-specific detection
    if sys.platform == 'win32' or sys.platform == 'darwin':
        if event.delta > 0: scroll_direction = -1; zoom_in_flag = True
        elif event.delta < 0: scroll_direction = 1; zoom_out_flag = True
    else: # Linux
        if event.num == 4: scroll_direction = -1; zoom_in_flag = True
        elif event.num == 5: scroll_direction = 1; zoom_out_flag = True

    # Perform Action
    if ctrl_pressed:
        # Zoom Logic
        new_zoom = state['zoom']
        if zoom_in_flag: new_zoom += ZOOM_STEP
        elif zoom_out_flag: new_zoom -= ZOOM_STEP
        new_zoom = max(MIN_ZOOM, min(MAX_ZOOM, new_zoom))
        if abs(new_zoom - state['zoom']) > 0.01:
            state['zoom'] = new_zoom
            load_pdf_page(tab_id)
    elif scroll_direction != 0:
        # Scroll Logic
        scroll_units = 2
        canvas_widget.yview_scroll(scroll_direction * scroll_units, 'units')

    return "break" # Prevent default scroll behavior

def open_selected_in_new_tab():
    """Opens the selected document from the FILE tree in a new viewer tab."""
    global file_tree
    selected_iid = file_tree.focus()
    if not selected_iid: return

    item_type = file_tree.set(selected_iid, "type")
    if item_type == "file":
        filepath = file_tree.set(selected_iid, "path")
        doc_id = None
        # --- LOOKUP doc_id from DB ---
        try:
            conn = sqlite3.connect(DATABASE_FILE)
            cursor = conn.cursor()
            cursor.execute("SELECT id FROM documents WHERE filepath = ?", (filepath,))
            result = cursor.fetchone()
            conn.close()
            if result:
                doc_id = result[0]
        except sqlite3.Error as e:
            print(f"Database error looking up doc_id for opening tab {filepath}: {e}")
        except Exception as e:
             print(f"Error during doc_id lookup for opening tab: {e}")
        # --- END LOOKUP ---

        if doc_id is not None: # Check if lookup was successful
            open_document_in_tab(doc_id) # Pass integer doc_id
        else:
             messagebox.showerror("Error", f"Could not find database entry for:\n{filepath}\n\nPlease re-scan the index.")
             print(f"Error: Could not find doc_id for file node: {filepath}")
    # else: Double-clicked a folder - do nothing (or maybe expand?)

def open_document_in_tab(doc_id):
    """Opens a document specified by doc_id in a new tab or selects existing tab.
       Handles PDF, DOCX, TXT, HTML (text view), and others (no preview).
       Returns the tab_id of the opened/selected tab, or None on failure."""
    global viewer_notebook, tab_states, open_files_map, root, file_icon, folder_icon # Need icons too

    details = get_document_details(doc_id)
    if not details:
        messagebox.showerror("Error", f"Could not retrieve details for document ID {doc_id}.")
        return None

    doc_id_chk, filename, filepath, _, _, _, _, _, _, _, _, _ = details
    if doc_id_chk != doc_id: print(f"Warning: Mismatch between requested doc_id ({doc_id}) and retrieved id ({doc_id_chk})")

    if not os.path.exists(filepath):
        messagebox.showerror("Error", f"File not found (it may have been moved or deleted):\n{filepath}\n\nConsider re-scanning.")
        return None

    # Check if file is already open
    if filepath in open_files_map:
        tab_id = open_files_map[filepath]
        try:
            if tab_id in viewer_notebook.tabs():
                viewer_notebook.select(tab_id)
                print(f"Selected existing tab {tab_id} for Doc ID {doc_id}")
                update_details_panel(doc_id) # Ensure details panel syncs on selection
                return tab_id
            else: # Cleanup stale mapping
                print(f"Cleaning up stale mapping for missing tab: {tab_id}")
                del open_files_map[filepath]
                if tab_id in tab_states: del tab_states[tab_id]
        except Exception as e:
            print(f"Error checking/selecting existing tab {tab_id}: {e}")
            if filepath in open_files_map: del open_files_map[filepath]
            if tab_id in tab_states: del tab_states[tab_id]


    # --- Create a New Tab ---
    tab_frame = ttk.Frame(viewer_notebook)
    tab_text = os.path.basename(filename)
    viewer_notebook.add(tab_frame, text=tab_text[:25] + ('...' if len(tab_text) > 25 else ''), padding=2)
    tab_id = viewer_notebook.tabs()[-1]
    print(f"Creating new tab {tab_id} for Doc ID {doc_id} ({filepath})")

    # --- Create widgets INSIDE the tab_frame ---
    tab_viewer_content_frame = Frame(tab_frame, bd=0); tab_viewer_content_frame.pack(expand=True, fill=tk.BOTH, padx=0, pady=0)
    tab_text_frame = Frame(tab_viewer_content_frame); tab_viewer_text = Text(tab_text_frame, wrap=tk.WORD, state=tk.DISABLED, bd=0, yscrollcommand=True); tab_scrollbar_text_y = ttk.Scrollbar(tab_text_frame, orient=tk.VERTICAL, command=tab_viewer_text.yview); tab_viewer_text.config(yscrollcommand=tab_scrollbar_text_y.set); tab_scrollbar_text_y.pack(side=tk.RIGHT, fill=tk.Y); tab_viewer_text.pack(side=tk.LEFT, expand=True, fill=tk.BOTH)
    tab_canvas_frame = Frame(tab_viewer_content_frame); tab_viewer_canvas = Canvas(tab_canvas_frame, bg='lightgrey', bd=0); tab_scrollbar_canvas_y = ttk.Scrollbar(tab_canvas_frame, orient=tk.VERTICAL, command=tab_viewer_canvas.yview); tab_scrollbar_canvas_x = ttk.Scrollbar(tab_canvas_frame, orient=tk.HORIZONTAL, command=tab_viewer_canvas.xview); tab_viewer_canvas.config(yscrollcommand=tab_scrollbar_canvas_y.set, xscrollcommand=tab_scrollbar_canvas_x.set); tab_scrollbar_canvas_y.pack(side=tk.RIGHT, fill=tk.Y); tab_scrollbar_canvas_x.pack(side=tk.BOTTOM, fill=tk.X); tab_viewer_canvas.pack(side=tk.LEFT, expand=True, fill=tk.BOTH)
    if sys.platform == 'win32' or sys.platform == 'darwin': tab_viewer_canvas.bind("<MouseWheel>", handle_scroll_or_zoom, add='+')
    else: tab_viewer_canvas.bind("<Button-4>", handle_scroll_or_zoom, add='+'); tab_viewer_canvas.bind("<Button-5>", handle_scroll_or_zoom, add='+')
    tab_pdf_nav_frame = ttk.Frame(tab_frame, padding=(5,2)); tab_pdf_nav_frame.pack(side=tk.BOTTOM, fill=tk.X, pady=0); tab_pdf_nav_frame.pack_forget()
    tab_back_button = ttk.Button(tab_pdf_nav_frame, text="< Back", command=go_back, width=6, state=tk.DISABLED); tab_back_button.pack(side=tk.LEFT, padx=(0,2)); tab_forward_button = ttk.Button(tab_pdf_nav_frame, text="Fwd >", command=go_forward, width=6, state=tk.DISABLED); tab_forward_button.pack(side=tk.LEFT, padx=(2,5)); tab_prev_button = ttk.Button(tab_pdf_nav_frame, text="< Prev", command=prev_page, width=6); tab_prev_button.pack(side=tk.LEFT, padx=(5,2)); tab_page_label = ttk.Label(tab_pdf_nav_frame, text="", width=25, anchor='c'); tab_page_label.pack(side=tk.LEFT, padx=2, expand=True, fill=tk.X); tab_zoom_out_button = ttk.Button(tab_pdf_nav_frame, text="-", command=zoom_out, width=3); tab_zoom_out_button.pack(side=tk.LEFT, padx=(0,2)); tab_zoom_in_button = ttk.Button(tab_pdf_nav_frame, text="+", command=zoom_in, width=3); tab_zoom_in_button.pack(side=tk.LEFT, padx=2); tab_next_button = ttk.Button(tab_pdf_nav_frame, text="Next >", command=next_page, width=6); tab_next_button.pack(side=tk.LEFT, padx=(2,0))

    # --- Store initial state ---
    initial_state = {
        'filepath': filepath, 'doc_id': doc_id, 'doc_obj': None,
        'page_num': 0, 'zoom': DEFAULT_ZOOM, 'tk_img_ref': None,
        'history_back': [], 'history_forward': [],
        'widgets': {
            'tab_frame': tab_frame, 'content_frame': tab_viewer_content_frame,
            'text_frame': tab_text_frame, 'text_viewer': tab_viewer_text,
            'canvas_frame': tab_canvas_frame, 'canvas': tab_viewer_canvas,
            'pdf_nav_frame': tab_pdf_nav_frame, 'page_label': tab_page_label,
            'back_button': tab_back_button, 'forward_button': tab_forward_button
        }
    }
    tab_states[tab_id] = initial_state
    open_files_map[filepath] = tab_id

    # --- Load content into the new tab ---
    ext = os.path.splitext(filepath)[1].lower()
    initial_load_successful = False
    state = tab_states[tab_id]

    try:
        if ext == '.pdf':
            pdf_doc_obj = None
            try:
                 pdf_doc_obj = fitz.open(filepath)
                 print(f"Tab {tab_id}: Opened fitz object for PDF. Type: {type(pdf_doc_obj)}, Pages: {len(pdf_doc_obj) if pdf_doc_obj else 'N/A'}")
                 state['doc_obj'] = pdf_doc_obj # Assign ONLY if successful
            except Exception as fitz_open_e:
                 error_msg = f"Error opening PDF file:\n{filepath}\n\n{fitz_open_e}"
                 print(f"!!! {error_msg}"); display_text_in_tab(tab_id, error_msg)
                 # state['doc_obj'] remains None

            if state.get('doc_obj') and len(state['doc_obj']) > 0:
                try:
                    nav_frame_widget = state['widgets'].get('pdf_nav_frame')
                    if nav_frame_widget: nav_frame_widget.pack(side=tk.BOTTOM, fill=tk.X, pady=0)
                    else: print("Warning: pdf_nav_frame widget not found in state.")
                    load_pdf_page(tab_id) # Load page 0
                    initial_load_successful = True
                except Exception as load_page_e:
                     error_msg = f"Error initially loading page 0 for PDF:\n{filepath}\n\n{load_page_e}"
                     print(f"!!! {error_msg}"); display_text_in_tab(tab_id, error_msg)
            elif state.get('doc_obj'): # 0 pages
                display_text_in_tab(tab_id, "Error: PDF has no pages.");
                if state['doc_obj']: state['doc_obj'].close(); state['doc_obj'] = None

        elif ext == '.docx' and DOCX_ENABLED:
             try:
                  doc = Document(filepath); full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                  display_text_in_tab(tab_id, full_text); initial_load_successful = True
             except Exception as docx_e: display_text_in_tab(tab_id, f"Error reading DOCX:\n{docx_e}")

        elif ext == '.txt':
             content = None; error_txt = None
             for enc in ['utf-8', 'cp1252', 'latin-1']:
                try:
                    with open(filepath, 'r', encoding=enc) as f: content = f.read(); break
                except UnicodeDecodeError: continue
                except Exception as txt_e: error_txt = f"Error reading TXT:\n{txt_e}"; break
             if content is not None: display_text_in_tab(tab_id, content); initial_load_successful = True
             elif error_txt: display_text_in_tab(tab_id, error_txt)
             else: display_text_in_tab(tab_id, f"Error decoding TXT.")

        elif ext in ['.html', '.htm']: # HTML Display Added
            print(f"Tab {tab_id}: Displaying HTML as extracted text.")
            html_content = ""; extracted_text = ""; error_html = None
            try:
                 for enc in ['utf-8', 'cp1252', 'latin-1']:
                      try:
                           with open(filepath, 'r', encoding=enc) as f: html_content = f.read(); break
                      except UnicodeDecodeError: continue
                 if not html_content: raise ValueError("Could not decode HTML")
                 html_content = re.sub(r'<script.*?<\/script>', '', html_content, flags=re.IGNORECASE | re.DOTALL)
                 html_content = re.sub(r'<style.*?<\/style>', '', html_content, flags=re.IGNORECASE | re.DOTALL)
                 extracted_text = re.sub(r'<.*?>', ' ', html_content)
                 extracted_text = re.sub(r'\s+', ' ', extracted_text).strip()
                 display_text_in_tab(tab_id, extracted_text or "(HTML content extracted as empty)")
                 initial_load_successful = True
            except Exception as html_ex:
                 error_html = f"Error reading/parsing HTML file:\n{html_ex}"
                 print(f"!!! {error_html}"); display_text_in_tab(tab_id, error_html)

        else: # Handle other non-previewable types
            display_text_in_tab(tab_id, f"'{filename}'\n\nPreview not available for this file type.\nUse 'File' -> 'Open Externally'.")
            initial_load_successful = True

    except Exception as e:
         error_msg = f"Unexpected error loading document content:\n{filepath}\n\n{e}"
         print(f"!!! {error_msg} (Type: {type(e)})")
         try: display_text_in_tab(tab_id, error_msg)
         except Exception as display_e: print(f"Further error displaying loading error: {display_e}"); status_bar_label.config(text=f"Error loading tab content for {filename[:30]}...")
         # --- Cleanup doc_obj only if PDF load failed initially ---
         current_state = tab_states.get(tab_id)
         if current_state and current_state.get('doc_obj') and not initial_load_successful and ext == '.pdf':
             print("Clearing PDF doc_obj due to loading error.")
             try: current_state['doc_obj'].close()
             except: pass
             current_state['doc_obj'] = None

    # --- Select tab and return ID ---
    try:
         if tab_id in viewer_notebook.tabs(): viewer_notebook.select(tab_id)
    except Exception as e: print(f"Error selecting new tab {tab_id}: {e}")
    root.update_idletasks()
    return tab_id
def close_current_tab():
    """Closes the currently active tab."""
    global viewer_notebook, tab_states, open_files_map
    try:
         tab_id = viewer_notebook.select() # Get ID of selected tab
         if not tab_id: return
    except tk.TclError: # No tabs selected / exist
         return

    print(f"Closing tab: {tab_id}")
    state = tab_states.get(tab_id)

    # Forget the tab visually BEFORE cleaning up state
    viewer_notebook.forget(tab_id)
    update_add_favorite_menu_state()

    # Clean up state
    if state:
        filepath = state.get('filepath')
        doc = state.get('doc_obj')

        # Close Fitz document object if open
        if doc:
            try:
                doc.close()
                print(f"Closed fitz object for {filepath}")
            except Exception as e:
                print(f"Error closing fitz object for {filepath}: {e}")

        # Remove from mappings
        if filepath and filepath in open_files_map:
            del open_files_map[filepath]
        if tab_id in tab_states:
            del tab_states[tab_id]

    # If no tabs left, clear metadata panel
    if not viewer_notebook.tabs():
         clear_details_panel()
# Placeholder comment - Ensure the full viewer functions from the *previous* version are present
# --- End Viewer Functions Placeholder ---

# --- Collapsible Pane Functions (Keep As Is) ---
COLLAPSED_PANE_WIDTH = 5
def toggle_left_pane():
    """Collapses or expands the left pane."""
    global is_left_pane_collapsed, left_sash_expanded_pos, root, main_paned_window

    if not main_paned_window: return # Ensure paned window exists

    current_sash_pos_x = main_paned_window.sash_coord(0)[0] # Get current X of first sash

    if is_left_pane_collapsed:
        # --- Expand Left Pane ---
        # Use stored position if valid, otherwise use default
        target_x = left_sash_expanded_pos if left_sash_expanded_pos > COLLAPSED_PANE_WIDTH else 350
        try:
            main_paned_window.sash_place(0, target_x, 0)
            root.update_idletasks()
            is_left_pane_collapsed = False
            # Update menu item text (optional)
            # view_menu.entryconfig("Toggle Left Pane", label="Hide Left Pane")
            if left_toggle_button: left_toggle_button.config(text="<")
            print(f"Expanded left pane, sash to {target_x}")
        except tk.TclError as e:
            print(f"Error expanding left pane: {e}")
    else:
        # --- Collapse Left Pane ---
        # Store the current position *before* collapsing, if it looks expanded
        if current_sash_pos_x > COLLAPSED_PANE_WIDTH * 2: # Heuristic to avoid storing collapsed pos
            left_sash_expanded_pos = current_sash_pos_x
            print(f"Stored left expanded pos: {left_sash_expanded_pos}")

        target_x = COLLAPSED_PANE_WIDTH
        try:
            main_paned_window.sash_place(0, target_x, 0)
            root.update_idletasks()
            is_left_pane_collapsed = True
            # Update menu item text (optional)
            # view_menu.entryconfig("Toggle Left Pane", label="Show Left Pane")
            if left_toggle_button: left_toggle_button.config(text=">")
            print(f"Collapsed left pane, sash to {target_x}")
        except tk.TclError as e:
            print(f"Error collapsing left pane: {e}")


def toggle_right_pane():
    """Collapses or expands the right pane."""
    global is_right_pane_collapsed, right_sash_expanded_pos, root, main_paned_window

    if not main_paned_window or not root: return # Ensure components exist

    sash_index = 1 # The second sash controls the right pane boundary
    current_sash_pos_x = main_paned_window.sash_coord(sash_index)[0]
    window_width = root.winfo_width() # Get current window width

    if is_right_pane_collapsed:
        # --- Expand Right Pane ---
        # Calculate a reasonable default if not stored
        if right_sash_expanded_pos is None or right_sash_expanded_pos >= window_width - COLLAPSED_PANE_WIDTH * 2:
             # If stored pos looks collapsed or is invalid, calculate default
             target_x = window_width - 300 # Default expanded width
        else:
             target_x = right_sash_expanded_pos

        # Ensure target isn't too close to left sash
        left_sash_x = main_paned_window.sash_coord(0)[0]
        target_x = max(target_x, left_sash_x + 200) # Ensure minimum center pane width

        try:
            main_paned_window.sash_place(sash_index, target_x, 0)
            is_right_pane_collapsed = False
            # Update menu item text (optional)
            # view_menu.entryconfig("Toggle Right Pane", label="Hide Right Pane")
            if right_toggle_button: right_toggle_button.config(text=">")
            print(f"Expanded right pane, sash to {target_x}")
        except tk.TclError as e:
            print(f"Error expanding right pane: {e}")
    else:
        # --- Collapse Right Pane ---
        # Store the current position *before* collapsing, if it looks expanded
        if current_sash_pos_x < window_width - COLLAPSED_PANE_WIDTH * 2: # Heuristic
            right_sash_expanded_pos = current_sash_pos_x
            print(f"Stored right expanded pos: {right_sash_expanded_pos}")

        # Place sash very close to the right edge
        target_x = window_width - COLLAPSED_PANE_WIDTH
        try:
            main_paned_window.sash_place(sash_index, target_x, 0)
            is_right_pane_collapsed = True
            # Update menu item text (optional)
            # view_menu.entryconfig("Toggle Right Pane", label="Show Right Pane")
            if right_toggle_button: right_toggle_button.config(text="<")
            print(f"Collapsed right pane, sash to {target_x}")
        except tk.TclError as e:
            print(f"Error collapsing right pane: {e}")

# Placeholder comment - Ensure the full functions from the previous version are here
# --- End Collapsible Pane Placeholder ---


# --- GUI Action Functions ---

# >>>>> NEW: File Tree Browser Functions <<<<<
# --- Context Menu for File Tree ---
file_tree_context_menu = None # Global reference
# --- Context Menu for Viewer Tabs ---
tab_context_menu = None # Global reference
# Global reference to the search results tab and its treeview
search_results_tab_id = None
search_results_tree = None
search_results_map = {} # map tree iid -> {'doc_id': id, 'page': page_num}
# --- Semi-Automatic Linking Functions ---

# --- Compile Regex Patterns (REVISED based on examples) ---

# 1. Filename Pattern (Keep as general fallback - might need adjustment)
#    This captures words/spaces/dots/hyphens ending in a common extension.
FILENAME_PATTERN = re.compile(r'([\w\-. ]+\.(?:pdf|docx|doc|xlsx|xls|txt))\b', re.IGNORECASE)

# 2. DiaSorin Corp Codes (e.g., Corp-GOP-000141, Corp-TN-000177)
#    Looks for 'Corp', hyphen, 2-3 letters (Type), hyphen, digits.
CORP_CODE_PATTERN = re.compile(r'\b(Corp-[A-Z]{2,3}-\d+)\b', re.IGNORECASE)

# 3. Part Numbers (e.g., PN A0227, PN 100001893)
#    Looks for 'PN', space, then alphanumeric code (letters/numbers).
#    This is quite broad, might match other things. Needs testing.
PART_NUMBER_PATTERN = re.compile(r'\b(PN\s+[A-Z0-9]+)\b', re.IGNORECASE)

# 4. Error Codes (Keep previous example, adjust if needed for DiaSorin specific errors)
ERROR_CODE_PATTERN = re.compile(r'\b(E[- ]?\d+|Err(?:or)? \d+|Code 0x[A-F0-9]+)\b', re.IGNORECASE)

# Combine patterns for easier iteration - PRIORITIZE specific patterns
REFERENCE_PATTERNS = [
    {'type': 'diasorin_code', 'regex': CORP_CODE_PATTERN},   # More specific first
    {'type': 'part_no',       'regex': PART_NUMBER_PATTERN},
    {'type': 'error_code',    'regex': ERROR_CODE_PATTERN},
    {'type': 'filename',      'regex': FILENAME_PATTERN},     # Filename as fallback
]
# --- End Patterns ---
def check_scan_queue():
    """Checks the scan status queue and updates the GUI."""
    global scan_status_queue, status_bar_label, scan_progress_bar, scan_button_ref, root

    try:
        # Get message from queue if available, non-blocking
        message = scan_status_queue.get_nowait()

        msg_type = message.get('type')

        if msg_type == 'status':
            if status_bar_label: status_bar_label.config(text=message.get('message', 'Scanning...'))
            # Keep checking queue if status update received
            root.after(50, check_scan_queue)
        elif msg_type == 'progress':
             # Optional: Update progress bar if using determinate mode later
             # count = message.get('count', 0)
             # If progress bar is indeterminate, just keep checking
             root.after(50, check_scan_queue)
        elif msg_type == 'error':
            # Error occurred in worker thread
            if scan_progress_bar: scan_progress_bar.stop(); scan_progress_bar.pack_forget()
            if scan_button_ref: scan_button_ref.config(state=tk.NORMAL) # Re-enable button
            messagebox.showerror("Scan Error", message.get('message', 'Unknown error during scan.'))
            if status_bar_label: status_bar_label.config(text="Scan failed! Ready.")
            build_file_tree(); clear_details_panel() # Refresh tree even on error
        elif msg_type == 'info':
             # Informational message (e.g., no paths)
             messagebox.showinfo("Scan Info", message.get('message', 'Scan information.'))
             # Assume scan finished cleanly in this case
             if scan_progress_bar: scan_progress_bar.stop(); scan_progress_bar.pack_forget()
             if scan_button_ref: scan_button_ref.config(state=tk.NORMAL)
             if status_bar_label: status_bar_label.config(text="Scan finished. Ready.")
             build_file_tree(); clear_details_panel()
        elif msg_type == 'finished':
            # Scan finished successfully
            if scan_progress_bar: scan_progress_bar.stop(); scan_progress_bar.pack_forget()
            if scan_button_ref: scan_button_ref.config(state=tk.NORMAL) # Re-enable

            # Format final message from received stats
            duration = message.get('duration', 0)
            final_msg = f"Scan Complete ({duration:.1f}s). Added: {message.get('added',0)}, Updated: {message.get('updated',0)}, Re-Indexed: {message.get('reindexed',0)}, Removed: {message.get('removed',0)}."
            if message.get('errors', 0) > 0: final_msg += f" Text Errors: {message.get('errors',0)}."
            final_msg += " Ready."
            if status_bar_label: status_bar_label.config(text=final_msg)
            messagebox.showinfo("Scan Complete", f"Scan finished in {duration:.1f} seconds.\nDocs Added: {message.get('added',0)}\nDocs Updated: {message.get('updated',0)}\nFiles Re-Indexed(FTS): {message.get('reindexed',0)}\nDocs Removed: {message.get('removed',0)}\nText Extraction Errors: {message.get('errors',0)}")
            build_file_tree(); clear_details_panel() # Refresh tree

    except queue.Empty:
        # Queue empty, check again later if scan is still running
        # Check button state to see if scan is considered 'running' by the UI
        if scan_button_ref and scan_button_ref['state'] == tk.DISABLED:
             root.after(100, check_scan_queue) # Check again in 100ms
    except Exception as e:
         print(f"Error processing scan status queue: {e}")
         if scan_progress_bar: scan_progress_bar.stop(); scan_progress_bar.pack_forget()
         if scan_button_ref: scan_button_ref.config(state=tk.NORMAL)
         if status_bar_label: status_bar_label.config(text="Error processing scan results. Ready.")
def start_scan_thread():
    """Disables UI, starts the scan worker thread, and initiates queue check."""
    global scan_button_ref, scan_status_queue, scan_progress_bar, root, status_bar_label

    # --- Check if scan is already running (optional) ---
    # if scan_button_ref and scan_button_ref['state'] == tk.DISABLED:
    #     messagebox.showinfo("Scan Info", "A scan is already in progress.")
    #     return

    # --- Disable Scan Button ---
    if scan_button_ref: scan_button_ref.config(state=tk.DISABLED)
    # Disable Scan menu item too? Requires storing menu item reference. Let's skip for now.

    # --- Show and Start Progress Bar ---
    if scan_progress_bar:
        scan_progress_bar.pack(side=tk.BOTTOM, fill=tk.X, padx=1, pady=0, before=status_bar_label)
        scan_progress_bar.start(10)
    status_bar_label.config(text="Starting scan thread...")
    root.update_idletasks()

    # --- Clear the queue (in case of previous aborted run) ---
    while not scan_status_queue.empty():
        try: scan_status_queue.get_nowait()
        except queue.Empty: break

    # --- Start Worker Thread ---
    scan_thread = threading.Thread(target=scan_and_update_worker, args=(scan_status_queue,), daemon=True)
    scan_thread.start()

    # --- Start Queue Check Loop ---
    root.after(100, check_scan_queue) # Start checking the queue

def check_search_queue():
    """Checks the results queue from the search thread without blocking and updates GUI."""
    print(f"--- check_search_queue called (Time: {time.time():.2f}) ---") # DEBUG - Top level entry
    global search_results_queue, status_bar_label, search_button_ref, search_entry # Added search_entry
    global search_results_tab_id, search_results_tree, search_results_map # Need these to update results tab
    global root, viewer_notebook # Need root and notebook

    try:
        # Get result from queue if available, non-blocking
        result_package = search_results_queue.get_nowait()

        print("Processing search results from queue...") # Debug
        query = result_package["query"]
        results_data = result_package["results_data"]
        snippet_map = result_package["snippet_map"]
        error_occurred = result_package["error"]
        error_message = result_package["error_message"]

        # --- Re-enable search button/entry ---
        if search_button_ref:
            try: search_button_ref.config(state=tk.NORMAL)
            except: pass # Ignore error if button destroyed
        if search_entry:
             try: search_entry.config(state=tk.NORMAL)
             except: pass # Ignore error if entry destroyed


        if error_occurred:
            messagebox.showerror("Search Error", f"An error occurred during search:\n{error_message}")
            if status_bar_label: status_bar_label.config(text=f"Search failed for '{query}'. Ready.")
            return # Stop processing this result

        # --- Update GUI with results ---

        # --- DEBUG: Check if tab exists or needs creation ---
        tab_exists = False
        print(f"Checking for Search Results Tab ID: {search_results_tab_id}") # Debug
                # --- Switch to the results tab (Deferred) ---
        if search_results_tab_id:
             try:
                  print(f"  Scheduling selection of Search Results tab: {search_results_tab_id}") # Debug
                  # Define a nested function or lambda for the deferred call
                  def _select_tab():
                       try: # Check again inside the deferred call
                            if viewer_notebook and search_results_tab_id in viewer_notebook.tabs():
                                 viewer_notebook.select(search_results_tab_id)
                                 print(f"  Deferred selection of tab {search_results_tab_id} successful.")
                            else:
                                 print(f"  Deferred selection failed: Tab {search_results_tab_id} no longer exists.")
                       except Exception as select_e_after:
                            print(f"  !!! ERROR during deferred selection of search results tab: {select_e_after}")

                  root.after(50, _select_tab) # Select after 50ms delay

             except Exception as select_e: # Catch error during scheduling itself
                  print(f"  !!! ERROR scheduling selection of search results tab: {select_e}")

        if status_bar_label: status_bar_label.config(text=f"Search for '{query}' complete. Found {len(results_data)} documents. Ready.")
        if not tab_exists:
            print("  Search Results Tab does not exist or notebook invalid, creating...") # Debug
            if not viewer_notebook or not viewer_notebook.winfo_exists():
                 print("  !!! Cannot create search results tab: viewer_notebook is invalid.")
                 status_bar_label.config(text="Error: Cannot display search results.")
                 return # Cannot proceed

            # --- CREATE TAB ---
            try:
                 results_tab_frame = ttk.Frame(viewer_notebook, padding=5)
                 viewer_notebook.add(results_tab_frame, text=" Search Results ", sticky="nsew")
                 search_results_tab_id = viewer_notebook.tabs()[-1] # Get new ID
                 print(f"  -> Created Search Results tab: {search_results_tab_id}") # Debug
                 root.update_idletasks()

                 cols = ('filename', 'page', 'snippet'); tree = ttk.Treeview(results_tab_frame, columns=cols, show='headings', selectmode='browse')
                 tree.heading('filename', text='Document'); tree.heading('page', text='Page'); tree.heading('snippet', text='Context Snippet')
                 tree.column('filename', width=250, stretch=tk.YES, anchor='w'); tree.column('page', width=50, stretch=tk.NO, anchor='e'); tree.column('snippet', width=450, stretch=tk.YES, anchor='w')
                 vsb = ttk.Scrollbar(results_tab_frame, orient="vertical", command=tree.yview); hsb = ttk.Scrollbar(results_tab_frame, orient="horizontal", command=tree.xview)
                 tree.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)
                 vsb.pack(side=tk.RIGHT, fill=tk.Y); hsb.pack(side=tk.BOTTOM, fill=tk.X); tree.pack(expand=True, fill=tk.BOTH)
                 search_results_tree = tree # Assign to global
                 print(f"  -> Assigned Treeview widget: {search_results_tree}") # Debug
                 tree.bind("<Double-1>", on_search_result_double_click) # Bind double-click
                 root.update_idletasks()
            except Exception as create_e:
                 print(f"  !!! ERROR CREATING SEARCH RESULTS TAB/TREE: {create_e}") # Debug
                 # Reset globals if creation failed
                 search_results_tab_id = None
                 search_results_tree = None
                 if status_bar_label: status_bar_label.config(text="Error creating search results view. Ready.")
                 return # Stop processing if UI failed


        # --- Populate Results Treeview ---
        if search_results_tree: # Check if treeview reference is valid
             search_results_tree.delete(*search_results_tree.get_children())
             search_results_map.clear()
             print(f"  Populating results tree. Number of results: {len(results_data)}") # Debug

             if results_data:
                  for i, row_data in enumerate(results_data):
                       doc_id, filename, _, _, _, _ = row_data # Get basic info
                       snippet_info = snippet_map.get(doc_id) # Get snippet/page from FTS results
                       if snippet_info: snippet, page_num = snippet_info; display_snippet = snippet.replace('\n', ' ').replace('\r', ''); display_page = str(page_num + 1) if page_num is not None else "N/A"
                       else: display_snippet = "(Metadata Match)"; display_page = "N/A"; page_num = None
                       try:
                            iid = search_results_tree.insert('', tk.END, values=(filename, display_page, display_snippet))
                            search_results_map[iid] = {'doc_id': doc_id, 'page': page_num}
                            # print(f"    Inserted row {i+1}: {filename}") # Very verbose
                       except Exception as insert_e:
                            print(f"  !!! ERROR Inserting row into search results tree: {insert_e}") # Debug
             else:
                  search_results_tree.insert('', tk.END, values=("No matches found.", "", ""))
             print("  Finished populating results tree.") # Debug
        else:
             print("  Error: search_results_tree reference is None, cannot populate.") # Debug


        # --- Switch to the results tab ---
        if search_results_tab_id:
             try:
                  print(f"  Selecting Search Results tab: {search_results_tab_id}") # Debug
                  if viewer_notebook and search_results_tab_id in viewer_notebook.tabs():
                       viewer_notebook.select(search_results_tab_id)
                  else:
                       print(f"  Cannot select tab {search_results_tab_id}, it doesn't exist.")
             except Exception as select_e:
                  print(f"  !!! ERROR selecting search results tab: {select_e}") # Debug

        if status_bar_label: status_bar_label.config(text=f"Search for '{query}' complete. Found {len(results_data)} documents. Ready.")

    except queue.Empty:
        # Queue empty, check again later if scan is still running
        is_search_running = False
        if search_button_ref:
             try: # Check state safely
                button_state = search_button_ref['state']
                print(f"  Queue empty. Search button state: {button_state}") # DEBUG
                if search_button_ref['state'] == tk.DISABLED:
                       is_search_running = True
             except Exception as e:
              print(f"  Error getting button state: {e}") # DEBUG

        if is_search_running:
             print(f"  Rescheduling check_search_queue...") # DEBUG
             if root: root.after(100, check_search_queue) # Check again in 100ms
        else:
            print(f"  Search button not disabled. Stopping queue check loop.") # DEBUG
    except Exception as e:
         print(f"Error processing search results queue: {e}")
         # Ensure UI unlocked on error
         if search_button_ref:
              try: search_button_ref.config(state=tk.NORMAL)
              except: pass
         if search_entry:
              try: search_entry.config(state=tk.NORMAL)
              except: pass
         if status_bar_label: status_bar_label.config(text="Error processing search results. Ready.")
         
def find_potential_references(text_content):
    """Scans text content for potential document references using regex."""
    potential_refs = defaultdict(set) # type -> {matched_string, ...}
    if not text_content: return potential_refs

    for pattern_info in REFERENCE_PATTERNS:
        try:
            matches = pattern_info['regex'].finditer(text_content)
            for match in matches:
                # Add the matched string (group 1 if pattern uses capturing group)
                ref_text = match.group(1) if pattern_info['regex'].groups > 0 else match.group(0)
                if ref_text: # Ensure not empty
                    potential_refs[pattern_info['type']].add(ref_text.strip())
        except Exception as e:
            print(f"Regex error for pattern type {pattern_info['type']}: {e}")
    return potential_refs


def find_matching_docs_in_db(references_dict, current_doc_id):
    """Looks up potential references in the database (filename, keywords, AND FTS),
       excluding self-references and existing links."""
    if not references_dict: return []

    conn = sqlite3.connect(DATABASE_FILE)
    cursor = conn.cursor()
    matches = [] # List of {'ref_type': type, 'ref_text': text, 'match_doc_id': id, 'match_filename': fname}
    checked_matches = set() # Keep track of (ref_text, match_doc_id) pairs

    for ref_type, ref_set in references_dict.items():
        for ref_text in ref_set:
            found_match_info = None # Store tuple (match_doc_id, match_filename)
            print(f"Looking up Ref: '{ref_text}' (Type: {ref_type})") # Debug

            try:
                if ref_type == 'filename':
                    # 1. Exact Filename Match
                    cursor.execute("SELECT id, filename FROM documents WHERE filename = ? AND id != ?", (ref_text, current_doc_id))
                    found_match_info = cursor.fetchone()
                    if not found_match_info: # 2. Partial Filename Match
                         cursor.execute("SELECT id, filename FROM documents WHERE filename LIKE ? AND id != ?", (f"%{ref_text}%", current_doc_id))
                         found_match_info = cursor.fetchone() # Take first partial

                # --- Handle Codes/PNs/Errors ---
                elif ref_type in ['diasorin_code', 'part_no', 'error_code']:
                    # 1. Check Filename/Keywords first (quick check)
                    cursor.execute("""SELECT id, filename FROM documents WHERE
                                      (filename LIKE ? OR keywords LIKE ?) AND id != ?""",
                                   (f"%{ref_text}%", f"%{ref_text}%", current_doc_id))
                    found_match_info = cursor.fetchone()

                    # 2. If no filename/keyword match, try FTS Content Search
                    if not found_match_info:
                        print(f"  -> No filename/keyword match for '{ref_text}', trying FTS...")
                        try:
                             # Find the highest-ranked document containing the ref_text via FTS
                             # Exclude the current document ID from FTS results as well
                             cursor.execute("""
                                 SELECT fts.doc_id, d.filename
                                 FROM documents_fts fts JOIN documents d ON fts.doc_id = d.id
                                 WHERE fts.documents_fts MATCH ? AND fts.doc_id != ?
                                 ORDER BY rank -- Get best match first
                                 LIMIT 1
                             """, (ref_text, current_doc_id))
                             found_match_info = cursor.fetchone()
                             if found_match_info: print(f"     FOUND FTS match: ID={found_match_info[0]}, File={found_match_info[1]}")
                             else: print(f"     NO FTS match found.")
                        except sqlite3.OperationalError as fts_e:
                             print(f"     FTS search failed for '{ref_text}': {fts_e}")
                        except Exception as e:
                             print(f"     Unexpected error during FTS lookup for '{ref_text}': {e}")

                # --- Process found match (if any) ---
                if found_match_info:
                    match_doc_id, match_filename = found_match_info
                    match_key = (ref_text, match_doc_id) # Use tuple as key

                    # Avoid suggesting links that already exist OR suggesting same match twice
                    if match_key not in checked_matches:
                        cursor.execute("SELECT 1 FROM links WHERE source_doc_id = ? AND target_doc_id = ?", (current_doc_id, match_doc_id))
                        if not cursor.fetchone():
                            matches.append({
                                'ref_type': ref_type, 'ref_text': ref_text,
                                'match_doc_id': match_doc_id, 'match_filename': match_filename
                            })
                            print(f"  --> Added potential link suggestion: '{ref_text}' -> '{match_filename}'")
                        else:
                             print(f"  --> Link already exists for '{ref_text}' -> '{match_filename}'")
                    # Always mark as checked to prevent duplicates from different patterns/methods
                    checked_matches.add(match_key)
                # else: No match found by any method for this ref_text

            except sqlite3.Error as e:
                print(f"DB Error looking up ref '{ref_text}' (Type: {ref_type}): {e}")

    conn.close()
    return matches

def suggest_links_for_current_doc():
    """Scans the active document, finds potential references, and suggests links."""
    global root, status_bar_label
    active_tab_id = get_active_tab_id()
    state = get_active_tab_state()

    if not state or not state.get('doc_id') or not state.get('filepath'):
        messagebox.showinfo("Suggest Links", "Please open a document tab first.")
        return

    current_doc_id = state['doc_id']
    current_filepath = state['filepath']
    status_bar_label.config(text=f"Analyzing '{os.path.basename(current_filepath)}' for references...")
    root.update_idletasks()

    # --- Extract text (use full text for better results) ---
    # NOTE: Re-extracting text here is inefficient. Ideally, use indexed FTS text?
    # But FTS text might lose context/formatting needed for regex.
    # Let's re-extract for now.
    full_text = ""
    print(f"Extracting full text for {current_filepath}...")
    ext = os.path.splitext(current_filepath)[1].lower()
    try:
        if ext == '.pdf':
             with fitz.open(current_filepath) as doc:
                  for page in doc: full_text += page.get_text("text", sort=True) + "\n"
        elif ext == '.docx' and DOCX_ENABLED:
             doc = Document(current_filepath); full_text = "\n".join(p.text for p in doc.paragraphs)
        elif ext == '.txt':
             # Simple read, might fail on encoding
              with open(current_filepath, 'r', encoding='utf-8', errors='ignore') as f: full_text = f.read()
        # Add other formats if necessary
    except Exception as e:
         messagebox.showerror("Error", f"Could not extract text from current document:\n{e}")
         status_bar_label.config(text="Error extracting text. Ready.")
         return

    if not full_text:
         messagebox.showinfo("Suggest Links", "No text content found or extracted from the current document.")
         status_bar_label.config(text="No text content found. Ready.")
         return

    # --- Find potential references in text ---
    print("Scanning text for references...")
    potential_refs = find_potential_references(full_text)
    if not potential_refs:
         messagebox.showinfo("Suggest Links", "No potential document references found in the text.")
         status_bar_label.config(text="No references found. Ready.")
         return

    # --- Look up references in database ---
    print(f"Looking up {sum(len(s) for s in potential_refs.values())} potential references in DB...")
    suggested_links = find_matching_docs_in_db(potential_refs, current_doc_id)

    if not suggested_links:
        messagebox.showinfo("Suggest Links", "Found potential references, but none matched existing documents in the index (or links already exist).")
        status_bar_label.config(text="No new linkable matches found. Ready.")
        return

    # --- Display Suggestions in a Dialog ---
    dialog = Toplevel(root)
    dialog.title("Suggested Links")
    dialog.geometry("600x400")
    dialog.transient(root); dialog.grab_set()

    frame = ttk.Frame(dialog, padding="10"); frame.pack(expand=True, fill=tk.BOTH)
    ttk.Label(frame, text=f"Found {len(suggested_links)} potential new links. Select links to create:").pack(anchor='w', pady=(0, 5))

    # Use Checkbuttons in a scrolled frame/canvas? Simpler: Listbox with multi-select
    link_list_frame = ttk.Frame(frame); link_list_frame.pack(expand=True, fill=tk.BOTH, pady=(0, 10))
    link_listbox = Listbox(link_list_frame, height=10, selectmode='multiple', exportselection=False) # Allow multiple selection
    link_scrollbar = ttk.Scrollbar(link_list_frame, orient=tk.VERTICAL, command=link_listbox.yview); link_listbox.config(yscrollcommand=link_scrollbar.set)
    link_scrollbar.pack(side=tk.RIGHT, fill=tk.Y); link_listbox.pack(side=tk.LEFT, expand=True, fill=tk.BOTH)

    # Populate listbox and store data
    suggestion_data = {} # map listbox index -> full suggestion dict
    for index, suggestion in enumerate(suggested_links):
        # {'ref_type': type, 'ref_text': text, 'match_doc_id': id, 'match_filename': fname}
        display_text = f"Ref '{suggestion['ref_text']}' -> Link to: {suggestion['match_filename']} (ID: {suggestion['match_doc_id']})"
        link_listbox.insert(tk.END, display_text)
        suggestion_data[index] = suggestion

    # Action buttons
    button_frame = ttk.Frame(frame); button_frame.pack(fill=tk.X)

    def create_selected_links():
        selected_indices = link_listbox.curselection()
        if not selected_indices:
             messagebox.showwarning("No Selection", "Please select one or more suggested links to create.", parent=dialog)
             return

        created_count = 0
        error_count = 0
        for index in selected_indices:
            suggestion = suggestion_data.get(index)
            if suggestion:
                 # Create default description
                 link_desc = f"Found ref '{suggestion['ref_text']}' ({suggestion['ref_type']})"
                 if add_document_link(current_doc_id, suggestion['match_doc_id'], link_desc):
                      created_count += 1
                 else: error_count += 1 # add_document_link shows its own errors
            else: error_count += 1

        messagebox.showinfo("Link Creation Complete", f"Successfully created {created_count} links.\nFailed attempts/errors: {error_count}.", parent=dialog)
        update_links_tab(current_doc_id) # Refresh main links tab
        dialog.destroy()


    create_button = ttk.Button(button_frame, text="Create Selected Links", command=create_selected_links)
    create_button.pack(side=tk.LEFT, padx=10)
    cancel_button = ttk.Button(button_frame, text="Cancel", command=dialog.destroy)
    cancel_button.pack(side=tk.RIGHT, padx=10)

    dialog.wait_window()
    status_bar_label.config(text="Ready.")

def on_search_result_double_click(event):
    """Handles double-click on the Search Results treeview item."""
    global search_results_tree, search_results_map
    if not search_results_tree: return

    selected_iid = search_results_tree.focus()
    if not selected_iid: return

    target_data = search_results_map.get(selected_iid)
    if target_data:
        target_doc_id = target_data.get('doc_id')
        target_page_num = target_data.get('page') # 0-based (can be None)

        if target_doc_id is not None:
            print(f"Opening/Navigating to Doc ID {target_doc_id}, Page {target_page_num} from search results.")
            # Use go_to_favorite helper for consistency in opening/navigating
            go_to_favorite(0, target_doc_id, target_page_num) # Pass dummy fav_id=0
        else:
            print(f"Error: Could not find doc_id for selected search result {selected_iid}")
    else:
        print(f"Error: Could not find data map entry for selected search result {selected_iid}")

def execute_combined_search(event=None):
    """Starts the combined search in a separate thread."""
    global search_entry, status_bar_label, root, search_results_queue, search_button_ref

    query = search_entry.get().strip()
    if not query:
        status_bar_label.config(text="Enter search query and press Enter or click Search.")
        return

    # --- Disable Search Button/Entry ---
    if search_button_ref:
        search_button_ref.config(state=tk.DISABLED)
    search_entry.config(state=tk.DISABLED) # Disable entry during search too
    status_bar_label.config(text=f"Starting search for '{query}'...")
    root.update_idletasks()

    # --- Start Worker Thread ---
    # Pass the query and the queue to the worker
    search_thread = threading.Thread(target=perform_search_worker, args=(query, search_results_queue), daemon=True)
    search_thread.start()
    print("--- Search thread started ---") # DEBUG

    # --- Start Queue Check Loop ---
    # Schedule the first check of the queue using root.after
    root.after(100, check_search_queue) # Check queue after 100ms
    print("--- Initial check_search_queue scheduled ---") # DEBUG

def create_tab_context_menu():
    """Creates the context menu for the viewer notebook tabs."""
    global tab_context_menu, root
    tab_context_menu = Menu(root, tearoff=0)
    # Add command to close the specific tab right-clicked on
    tab_context_menu.add_command(label="Close Tab", command=close_specific_tab)
    # Add command to close all other tabs
    tab_context_menu.add_command(label="Close Other Tabs", command=close_other_tabs)
    # Add command to close all tabs
    tab_context_menu.add_command(label="Close All Tabs", command=close_all_tabs)
    tab_context_menu.add_separator()
    # Add command to open the file of the right-clicked tab externally
    tab_context_menu.add_command(label="Open Externally", command=open_external_from_tab)
    # Add command to reveal in file tree? More complex.

# Store the ID of the tab that was right-clicked
right_clicked_tab_id = None

def show_tab_context_menu(event):
     """Display context menu on right-click on a notebook tab."""
     global tab_context_menu, viewer_notebook, right_clicked_tab_id
     if not tab_context_menu: return

     # Identify the tab element under the cursor
     try:
          element = viewer_notebook.identify(event.x, event.y)
          # Only show menu if clicking on a tab element itself
          # Heuristic: tab elements often contain 'tab' or 'label' in their Tk name
          if element and ('tab' in element or 'label' in element):
               tab_index = viewer_notebook.index(f"@{event.x},{event.y}")
               # Store the actual widget ID of the right-clicked tab
               right_clicked_tab_id = viewer_notebook.tabs()[tab_index]
               print(f"Right-clicked on Tab ID: {right_clicked_tab_id}")

               # Enable/disable options based on context if needed (e.g., disable Close Other if only one tab)
               num_tabs = len(viewer_notebook.tabs())
               tab_context_menu.entryconfig("Close Other Tabs", state=tk.NORMAL if num_tabs > 1 else tk.DISABLED)

               # Show the menu
               tab_context_menu.tk_popup(event.x_root, event.y_root)
          else:
               # Clicked elsewhere on the notebook (client area, border, etc.)
               right_clicked_tab_id = None
               print(f"Right-click on notebook element '{element}', not a tab.")
     except tk.TclError:
          # Error identifying (e.g., no tabs exist)
          right_clicked_tab_id = None
     except Exception as e:
         print(f"Error showing tab context menu: {e}")
         right_clicked_tab_id = None
def close_specific_tab():
    """Closes the tab that was right-clicked."""
    global right_clicked_tab_id, viewer_notebook, tab_states, open_files_map
    if right_clicked_tab_id and right_clicked_tab_id in viewer_notebook.tabs():
        print(f"Closing specific tab via context menu: {right_clicked_tab_id}")
        tab_id_to_close = right_clicked_tab_id # Use the stored ID

        # --- Reuse cleanup logic from close_current_tab ---
        state = tab_states.get(tab_id_to_close)
        viewer_notebook.forget(tab_id_to_close) # Forget visual tab first
        if state: # Cleanup state
            filepath = state.get('filepath')
            doc = state.get('doc_obj')
            if doc:
                 try: doc.close()
                 except: pass # Ignore close errors
            if filepath and filepath in open_files_map: del open_files_map[filepath]
            if tab_id_to_close in tab_states: del tab_states[tab_id_to_close]
        # --- End Cleanup ---

        # Update details/menu if needed (e.g., if last tab closed)
        if not viewer_notebook.tabs(): clear_details_panel()
        update_add_favorite_menu_state()
    else:
        print("No valid right-clicked tab ID to close.")
    right_clicked_tab_id = None # Reset after action

def close_other_tabs():
     """Closes all tabs EXCEPT the one that was right-clicked."""
     global right_clicked_tab_id, viewer_notebook
     if right_clicked_tab_id is None: return

     tabs_to_close = [tab for tab in viewer_notebook.tabs() if tab != right_clicked_tab_id]
     print(f"Closing other tabs. Keeping: {right_clicked_tab_id}. Closing: {tabs_to_close}")
     for tab_id in tabs_to_close:
          # Reuse cleanup logic (similar to close_specific_tab/close_current_tab)
          state = tab_states.get(tab_id)
          viewer_notebook.forget(tab_id)
          if state:
               filepath = state.get('filepath'); doc = state.get('doc_obj')
               if doc:
                    try: doc.close()
                    except: pass
               if filepath and filepath in open_files_map: del open_files_map[filepath]
               if tab_id in tab_states: del tab_states[tab_id]
     update_add_favorite_menu_state() # State might have changed

def close_all_tabs():
     """Closes all tabs in the viewer notebook."""
     global viewer_notebook
     tabs_to_close = list(viewer_notebook.tabs()) # Get copy before modifying
     print(f"Closing all tabs: {tabs_to_close}")
     for tab_id in tabs_to_close:
          # Reuse cleanup logic
          state = tab_states.get(tab_id)
          viewer_notebook.forget(tab_id)
          if state:
               filepath = state.get('filepath'); doc = state.get('doc_obj')
               if doc:
                    try: doc.close()
                    except: pass
               if filepath and filepath in open_files_map: del open_files_map[filepath]
               if tab_id in tab_states: del tab_states[tab_id]
     clear_details_panel() # Clear details as no tabs left
     update_add_favorite_menu_state()

def open_external_from_tab():
    """Opens the file associated with the right-clicked tab externally."""
    global right_clicked_tab_id, tab_states
    if right_clicked_tab_id:
        state = tab_states.get(right_clicked_tab_id)
        if state and 'filepath' in state:
            filepath = state['filepath']
            print(f"Opening externally from tab {right_clicked_tab_id}: {filepath}")
            open_file_externally(filepath) # Use existing helper
        else:
            print(f"Could not find filepath for right-clicked tab {right_clicked_tab_id}")
    else:
        print("No valid right-clicked tab ID for external open.")
    right_clicked_tab_id = None # Reset after action

def on_notebook_click(event):
    """Handles left clicks on the viewer notebook, potentially forcing detail update
       even if the selected tab doesn't change (for single-tab scenario)."""
    global viewer_notebook, tab_states

    # Identify the element under the click
    try:
        element = viewer_notebook.identify(event.x, event.y)
        # Check if the click was on a tab element (usually contains 'label' or 'tab')
        # Note: Exact element names might vary slightly across themes/platforms
        # We check if it's NOT the main 'client' area or empty space
        if element and 'client' not in element and 'padding' not in element and 'border' not in element:
             # It's likely a click on or near a tab label
             current_selection_id = viewer_notebook.select() # Get the currently selected tab ID
             clicked_tab_index = viewer_notebook.index(f"@{event.x},{event.y}") # Get index of tab under cursor
             clicked_tab_id = viewer_notebook.tabs()[clicked_tab_index]

             print(f"Notebook click detected on element '{element}'. Clicked Tab ID: {clicked_tab_id}, Current Selection: {current_selection_id}") # Debug

             # If the clicked tab IS the currently selected one (the single-tab case, or re-clicking active tab)
             # OR if identify worked but select failed (e.g., notebook is empty)
             if clicked_tab_id == current_selection_id or not current_selection_id:
                  # Force details update based on the clicked/selected tab's state
                  state = tab_states.get(clicked_tab_id)
                  doc_id_to_display = state.get('doc_id') if state else None

                  if doc_id_to_display is not None:
                       print(f"Notebook click forcing details update for Doc ID: {doc_id_to_display}")
                       update_details_panel(doc_id_to_display)
                  else:
                       print("Notebook clicked tab has no associated document state. Clearing details.")
                       clear_details_panel()
                  # No need to call update_add_favorite_menu_state here, as <<NotebookTabChanged>> handles it
                  # if the selection *actually* changes. If it doesn't change, the state shouldn't change either.

    except tk.TclError:
        # Error identifying element (e.g., notebook is empty)
        print("Notebook click - TclError identifying element.")
        pass
    except Exception as e:
        print(f"Error in on_notebook_click: {e}")

def update_add_favorite_menu_state():
    """Enables/disables the 'Add Current View to Favorites' menu item."""
    global favorites_menu, add_favorite_menu_index
    print("--- Updating Add Favorite Menu State ---")
    if not favorites_menu: 
        print("  Favorites menu not yet created.")
        return

    state = get_active_tab_state()
    # --- DEBUG PRINTS ---
    active_tab_id_debug = get_active_tab_id()
    print(f"  Active Tab ID: {active_tab_id_debug}")
    print(f"  Retrieved State: {type(state)}")
    if state:
        print(f"  State Keys: {state.keys()}")
        print(f"  State doc_id: {state.get('doc_id')}")
    # --- END DEBUG ---
    # Enable only if a tab is active AND it has a valid doc_id
    can_add = state and 'doc_id' in state and state['doc_id'] is not None
    print(f"  Can Add?: {can_add}")

    try:
        target_state = tk.NORMAL if can_add else tk.DISABLED
        favorites_menu.entryconfig(add_favorite_menu_index, state=target_state)
        print(f"  Set menu item {add_favorite_menu_index} state to: {target_state}") # DEBUG
    except tk.TclError as e:
        print(f"  Error updating 'Add Favorite' menu state (TclError): {e}")
    except NameError:
        print("  Error updating 'Add Favorite' menu state: add_favorite_menu_index NameError.")
    except Exception as e:
         print(f"  Unexpected error updating 'Add Favorite' menu state: {e}")
# --- Favorite Action Functions ---

def add_current_view_to_favorites():
    """Adds the document/page in the active viewer tab to favorites."""
    global root
    active_tab_id = get_active_tab_id()
    state = get_active_tab_state()

    if not state or 'doc_id' not in state or state['doc_id'] is None:
        messagebox.showinfo("Add Favorite", "Please open and select a document tab first.")
        return

    doc_id = state['doc_id']
    page_num = state.get('page_num', 0) # Get current 0-based page

    # Get filename for default name suggestion
    details = get_document_details(doc_id)
    default_name = ""
    if details:
        default_name = f"{os.path.basename(details[1])} - Page {page_num + 1}"

    fav_name = simpledialog.askstring("Add Favorite", "Enter a name for this favorite:",
                                      initialvalue=default_name, parent=root)

    if fav_name: # User entered name and didn't cancel
        fav_name = fav_name.strip()
        if add_favorite(fav_name, doc_id, page_num):
            # Optionally refresh the Favorites menu immediately
            populate_favorites_menu()
            messagebox.showinfo("Favorite Added", f"'{fav_name}' added successfully.")
        # Else: Error shown by add_favorite

def go_to_favorite(fav_id, doc_id, page_number):
    """Opens/navigates to a specific favorite OR doc_id/page."""
    global root, tab_states
    print(f"Navigating via go_to_favorite: fav_id={fav_id}, Doc={doc_id}, Page={page_number}")

    target_tab_id = open_document_in_tab(doc_id)
    if target_tab_id is None: return

    # --- ADD CHECK FOR page_number ---
    if page_number is not None:
        state = tab_states.get(target_tab_id)
        if state and state.get('doc_obj'):
            doc_length = len(state['doc_obj'])
            target_page = max(0, min(page_number, doc_length - 1)) # Clamp page
            update_history_and_load(target_tab_id, target_page) # Use history nav
            print(f"Navigated tab {target_tab_id} to favorite/target page {target_page + 1}")
        else:
            print(f"Warning: Could not get document object for tab {target_tab_id} to navigate page.")
            messagebox.showwarning("Navigation Warning", "Opened document tab, but could not navigate to the specific page.")
    else:
        # No page number provided (e.g., from error search on TXT/DOCX), just ensure tab is open/selected
        print("No specific page number provided for navigation, tab opened/selected.")

def on_manage_fav_double_click(event, listbox_widget, data_map, dialog_window):
    """Handles double-click on the Manage Favorites listbox to navigate."""
    selected_indices = listbox_widget.curselection()
    if not selected_indices: return # Nothing selected

    index = selected_indices[0]
    if index in data_map:
        fav_id = data_map[index] # Get the fav_id stored for this list index

        # --- Get doc_id and page_number for this fav_id ---
        # Query the DB again to ensure we have the latest info
        fav_details = None
        conn = sqlite3.connect(DATABASE_FILE)
        cursor = conn.cursor()
        try:
            cursor.execute("SELECT doc_id, page_number FROM favorites WHERE fav_id = ?", (fav_id,))
            fav_details = cursor.fetchone()
        except sqlite3.Error as e:
            print(f"DB error fetching favorite details for ID {fav_id}: {e}")
        finally:
            conn.close()

        if fav_details:
            doc_id, page_number = fav_details
            print(f"Double-click navigate request: fav_id={fav_id}, doc_id={doc_id}, page={page_number}")
            # Close the dialog *before* navigating
            dialog_window.destroy()
            # Call the navigation function
            go_to_favorite(fav_id, doc_id, page_number)
        else:
            messagebox.showerror("Error", f"Could not retrieve details for Favorite ID {fav_id}.", parent=dialog_window)
    else:
        print(f"Error: Listbox index {index} not found in favorites_data map.")
        
def open_manage_favorites_dialog():
    """Opens a dialog to rename/delete existing favorites."""
    global root
    dialog = Toplevel(root) # <<< Store dialog reference
    dialog.title("Manage Favorites")
    dialog.geometry("500x350")
    dialog.transient(root); dialog.grab_set(); dialog.resizable(True, True)

    frame = ttk.Frame(dialog, padding="10"); frame.pack(expand=True, fill=tk.BOTH)

    # Listbox Frame
    list_frame = ttk.Frame(frame); list_frame.pack(expand=True, fill=tk.BOTH, pady=(0, 10))
    ttk.Label(list_frame, text="Saved Favorites:").pack(anchor='w')
    fav_listbox = Listbox(list_frame, height=10, exportselection=False)
    fav_scrollbar = ttk.Scrollbar(list_frame, orient=tk.VERTICAL, command=fav_listbox.yview)
    fav_listbox.config(yscrollcommand=fav_scrollbar.set)
    fav_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
    fav_listbox.pack(side=tk.LEFT, expand=True, fill=tk.BOTH)

    # --- Store data map locally ---
    favorites_data_map = {} # Map listbox index -> fav_id

    def refresh_fav_list():
        fav_listbox.delete(0, tk.END)
        favorites_data_map.clear() # Use local map
        favorites = get_favorites()
        for index, (fav_id, name, doc_id, page_num) in enumerate(favorites):
            details = get_document_details(doc_id)
            filename_part = f"({os.path.basename(details[1])[:20]}... - P{page_num+1})" if details else "(Doc Unknown)"
            fav_listbox.insert(tk.END, f"{name} {filename_part}")
            favorites_data_map[index] = fav_id # Store in local map
        update_manage_fav_buttons()

    def update_manage_fav_buttons():
        has_selection = bool(fav_listbox.curselection())
        rename_button.config(state=tk.NORMAL if has_selection else tk.DISABLED)
        delete_button.config(state=tk.NORMAL if has_selection else tk.DISABLED)

    fav_listbox.bind("<<ListboxSelect>>", lambda e: update_manage_fav_buttons())
    # --- ADD Double-Click Binding ---
    # Use lambda to pass necessary context to the handler
    fav_listbox.bind("<Double-Button-1>",
                     lambda event: on_manage_fav_double_click(event, fav_listbox, favorites_data_map, dialog))
    # --- END Binding ---


    # Button Frame
    button_frame = ttk.Frame(frame)
    button_frame.pack(fill=tk.X)

    def rename_selected():
        selected_indices = fav_listbox.curselection()
        if not selected_indices: return
        index = selected_indices[0]
        if index in favorites_data_map: # Use local map
            fav_id = favorites_data_map[index] # Use local map
            current_name = ""; favorites = get_favorites()
            for f_id, name, _, _ in favorites:
                 if f_id == fav_id: current_name = name; break
            new_name = simpledialog.askstring("Rename Favorite", "Enter new name:", initialvalue=current_name, parent=dialog)
            if new_name:
                 new_name = new_name.strip()
                 if new_name and rename_favorite(fav_id, new_name):
                      refresh_fav_list(); populate_favorites_menu()

    def delete_selected():
        selected_indices = fav_listbox.curselection()
        if not selected_indices: return
        index = selected_indices[0]
        if index in favorites_data_map: # Use local map
            fav_id = favorites_data_map[index] # Use local map
            name_to_delete = fav_listbox.get(index).split(" (")[0]
            confirm = messagebox.askyesno("Confirm Delete", f"Delete favorite '{name_to_delete}'?", parent=dialog)
            if confirm:
                if delete_favorite(fav_id):
                    refresh_fav_list(); populate_favorites_menu()

    rename_button = ttk.Button(button_frame, text="Rename Selected", command=rename_selected, state=tk.DISABLED)
    rename_button.pack(side=tk.LEFT, padx=5)
    delete_button = ttk.Button(button_frame, text="Delete Selected", command=delete_selected, state=tk.DISABLED)
    delete_button.pack(side=tk.LEFT, padx=5)
    close_button = ttk.Button(button_frame, text="Close", command=dialog.destroy)
    close_button.pack(side=tk.RIGHT, padx=5)

    refresh_fav_list() # Initial population
    dialog.wait_window()

# --- Function to Build Favorites Menu Dynamically ---
favorites_menu = None # Global reference to the menu itself

def populate_favorites_menu():
    """Clears and rebuilds the Favorites menu based on DB."""
    global favorites_menu
    if not favorites_menu: return # Menu not created yet

    # Remove all existing favorite entries (index 2 onwards, assuming Manage is 0, Sep is 1)
    try:
        last_index = favorites_menu.index(tk.END)
        if last_index is not None and last_index >= 2:
             favorites_menu.delete(2, tk.END)
    except tk.TclError: # No entries exist beyond separator
         pass
    except Exception as e:
         print(f"Error clearing favorites menu: {e}")


    # Get current favorites from DB
    favorites = get_favorites()

    if favorites:
        for fav_id, name, doc_id, page_number in favorites:
            # Create menu command, use lambda to capture loop variables correctly
            cmd = lambda fid=fav_id, did=doc_id, pn=page_number: go_to_favorite(fid, did, pn)
            favorites_menu.add_command(label=name, command=cmd)
    else:
        favorites_menu.add_command(label="(No favorites yet)", state=tk.DISABLED)

def create_file_tree_context_menu():
    """Creates the context menu for the file tree."""
    global file_tree_context_menu, root
    file_tree_context_menu = Menu(root, tearoff=0)
    file_tree_context_menu.add_command(label="Open in New Tab", command=open_selected_in_new_tab) # Can open multiple? No, let's keep single for now.
    file_tree_context_menu.add_command(label="Open Externally", command=open_file_externally_selected) # Opens the focused item
    file_tree_context_menu.add_separator()
    file_tree_context_menu.add_command(label="Edit Metadata (Single)...", command=open_edit_metadata_dialog) # Keep single edit
    file_tree_context_menu.add_command(label="Batch Edit Metadata...", command=open_batch_edit_dialog) # New Batch Edit
    file_tree_context_menu.add_command(label="Add Link from this...", command=open_add_link_dialog)
    # Initially disable items that require a selection
    update_file_tree_context_menu_state()


def show_file_tree_context_menu(event):
     """Display context menu on right-click in the file tree."""
     global file_tree_context_menu, file_tree
     if not file_tree_context_menu: return

     # Identify item under cursor (might not be selected yet with multi-select)
     iid_under_cursor = file_tree.identify_row(event.y)

     # If right-clicking on an unselected item, select ONLY that item first
     if iid_under_cursor and iid_under_cursor not in file_tree.selection():
          file_tree.selection_set(iid_under_cursor)
          file_tree.focus(iid_under_cursor)

     update_file_tree_context_menu_state() # Update state based on current selection
     file_tree_context_menu.tk_popup(event.x_root, event.y_root)


def update_file_tree_context_menu_state():
    """Enable/disable file tree context menu items based on selection."""
    global file_tree_context_menu, file_tree
    if not file_tree_context_menu: return

    selected_items = file_tree.selection()
    num_selected = len(selected_items)
    is_single_file_selected = False
    all_selected_are_files = True

    if num_selected == 1:
         item_type = file_tree.set(selected_items[0], "type")
         if item_type == 'file':
              is_single_file_selected = True
         else:
              all_selected_are_files = False # Single selected item is a folder
    elif num_selected > 1:
        # Check if ALL selected items are files
        for iid in selected_items:
            if file_tree.set(iid, "type") != 'file':
                all_selected_are_files = False
                break
    else: # No selection
         all_selected_are_files = False


    # Enable/disable items
    file_tree_context_menu.entryconfig("Open in New Tab", state=tk.NORMAL if is_single_file_selected else tk.DISABLED)
    file_tree_context_menu.entryconfig("Open Externally", state=tk.NORMAL if is_single_file_selected else tk.DISABLED)
    file_tree_context_menu.entryconfig("Edit Metadata (Single)...", state=tk.NORMAL if is_single_file_selected else tk.DISABLED)
    file_tree_context_menu.entryconfig("Add Link from this...", state=tk.NORMAL if is_single_file_selected else tk.DISABLED)
    # Batch edit enabled if one or more FILES are selected
    file_tree_context_menu.entryconfig("Batch Edit Metadata...", state=tk.NORMAL if num_selected > 0 and all_selected_are_files else tk.DISABLED)

def populate_tree_node(parent_iid, dir_path):
    """Populates the children of a given directory node in the file tree."""
    global file_tree, folder_icon, file_icon
    try:
        # Clear existing dummy node if present
        children = file_tree.get_children(parent_iid)
        if children and file_tree.item(children[0], "text") == "Loading...":
            file_tree.delete(children[0])

        # List directory contents
        for item_name in sorted(os.listdir(dir_path), key=str.lower):
            item_path = os.path.join(dir_path, item_name)
            is_dir = os.path.isdir(item_path)
            is_supported_file = not is_dir and item_name.lower().endswith(SUPPORTED_EXTENSIONS) and not item_name.startswith('~$')

            if is_dir:
                # Insert folder node
                folder_iid = file_tree.insert(parent_iid, tk.END, text=f" {item_name}", # Add space for icon
                                              values=[item_path, "folder", ""], # path, type, doc_id
                                              open=False, image=folder_icon) # Use folder icon
                # Add a dummy child to make it expandable
                file_tree.insert(folder_iid, tk.END, text="Loading...")
            elif is_supported_file:
                file_tree.insert(parent_iid, tk.END, text=f" {item_name}",
                             values=[item_path, "file"], # Store only path and type
                             image=file_icon)

    except OSError as e:
        print(f"Error reading directory {dir_path}: {e}")
        # Display error in tree?
        children = file_tree.get_children(parent_iid)
        if children and file_tree.item(children[0], "text") == "Loading...":
            file_tree.item(children[0], text="Error reading folder")
    except Exception as e:
        print(f"Unexpected error populating tree node {dir_path}: {e}")


def on_tree_open(event):
    """Callback when a tree node is expanded."""
    global file_tree
    iid = file_tree.focus() # The item being opened
    if not iid: return
    item_type = file_tree.set(iid, "type")
    if item_type == "folder":
        dir_path = file_tree.set(iid, "path")
        # Populate node if it hasn't been populated yet (check dummy child)
        children = file_tree.get_children(iid)
        if children and file_tree.item(children[0], "text") == "Loading...":
            populate_tree_node(iid, dir_path)

def build_file_tree():
    """Populates the top-level nodes of the file tree based on scan_paths."""
    global file_tree, folder_icon
    # Clear existing tree
    file_tree.delete(*file_tree.get_children())
    # Get root paths from DB
    scan_paths = get_scan_paths()
    for path in scan_paths:
        if os.path.isdir(path):
            # Insert top-level folder node
            node_text = os.path.basename(path) or path # Use full path if basename is empty (e.g., C:\)
            folder_iid = file_tree.insert('', tk.END, text=f" {node_text}",
                                          values=[path, "folder", ""],
                                          open=False, image=folder_icon)
            # Add dummy child to make it expandable
            file_tree.insert(folder_iid, tk.END, text="Loading...")
        else:
            print(f"Configured scan path is not a valid directory: {path}")

# >>>>> MODIFIED: Search Functionality for Tree <<<<<
# Keep track of original tree structure for clearing filter
original_tree_items = {} # Store {iid: (parent, index, options)}

def on_note_click(event):
    """Handles clicks within the notes Text widget to select a note."""
    global notes_text_widget, selected_note_id

    if not notes_text_widget: return
    newly_selected_id = None # Track the ID found in this click

    # Get index of the click
    index = notes_text_widget.index(f"@{event.x},{event.y}")
    # Get all tags at the clicked index
    tags_here = notes_text_widget.tag_names(index)

    # --- Find the specific note ID tag ---
    current_note_tag = None
    for tag in tags_here:
        if tag.startswith("note_"):
            current_note_tag = tag
            break # Found it

    if current_note_tag:
        try:
            newly_selected_id = int(current_note_tag.split("_")[1])
        except (ValueError, IndexError):
            print(f"Error parsing note ID from tag '{current_note_tag}'.")
            newly_selected_id = None
    # --- End Find ---

    # --- Update Selection Highlight ---
    # If clicked on the *same* note again, maybe deselect? (Optional)
    # if newly_selected_id == selected_note_id:
    #     newly_selected_id = None # Deselect

    # Remove previous highlight regardless
    notes_text_widget.tag_remove("selected_note", "1.0", tk.END)

    if newly_selected_id is not None:
        # Apply selection highlight to the newly selected note
        new_note_tag = f"note_{newly_selected_id}"
        tag_ranges = notes_text_widget.tag_ranges(new_note_tag)
        if tag_ranges:
            notes_text_widget.tag_add("selected_note", tag_ranges[0], tag_ranges[1])
        print(f"Selected Note ID: {newly_selected_id}")
    else:
         print("No note selected.")

    # Update global state *after* visual update
    selected_note_id = newly_selected_id

    # Update button states based on whether a note is now selected
    update_note_buttons_state()

def on_note_double_click(event):
    """Handles double-clicks in notes view to jump to document page."""
    global notes_text_widget, root

    if not notes_text_widget: return
    note_id = None
    index = notes_text_widget.index(f"@{event.x},{event.y}")
    tags_here = notes_text_widget.tag_names(index)
    for tag in tags_here:
        if tag.startswith("note_"):
            try: note_id = int(tag.split("_")[1]); break
            except: pass

    if note_id is None:
        print("Double-clicked, but couldn't identify note ID.")
        return

    print(f"Double-clicked Note ID: {note_id}")

    # --- Get Doc ID, Page, and Filepath for this Note ID ---
    conn = sqlite3.connect(DATABASE_FILE); cursor = conn.cursor()
    target_doc_id = None
    target_page = None
    target_filepath = None
    try:
         cursor.execute("""
             SELECT n.doc_id, n.page_number, d.filepath
             FROM notes n JOIN documents d ON n.doc_id = d.id
             WHERE n.note_id = ?
         """, (note_id,))
         result = cursor.fetchone()
         if result:
              target_doc_id, target_page, target_filepath = result
         else:
              messagebox.showerror("Error", f"Note ID {note_id} not found in database.")
              return
    except sqlite3.Error as e:
         messagebox.showerror("Database Error", f"Error fetching note details:\n{e}")
         return
    finally:
         conn.close()

    if target_doc_id is None or target_filepath is None:
         print("Could not retrieve document details for the note.")
         return

    # --- Navigate Viewer ---
    # Check if the document is already open
    target_tab_id = None
    if target_filepath in open_files_map:
         potential_tab_id = open_files_map[target_filepath]
         # Verify the tab still exists visually
         if potential_tab_id in viewer_notebook.tabs():
              target_tab_id = potential_tab_id

    if target_tab_id:
         # Switch to existing tab
         print(f"Switching to existing tab {target_tab_id} for {target_filepath}")
         viewer_notebook.select(target_tab_id)
         root.update_idletasks() # Allow tab switch to process
    else:
         # Open document in new tab first
         print(f"Opening document {target_filepath} in new tab...")
         open_document_in_tab(target_doc_id)
         # Find the new tab ID (should be the last one now)
         try: target_tab_id = viewer_notebook.tabs()[-1]
         except IndexError: print("Error: Could not find newly created tab ID."); return

    # --- Jump to Page (if available and tab found) ---
    if target_tab_id and target_page is not None:
         state = tab_states.get(target_tab_id)
         if state and state.get('doc_obj'):
              if 0 <= target_page < len(state['doc_obj']):
                   state['page_num'] = target_page
                   load_pdf_page(target_tab_id) # Load the target page
                   print(f"Jumped to page {target_page + 1} in tab {target_tab_id}")
              else:
                   print(f"Target page {target_page} out of bounds for document.")
         else:
              print("Could not get state or document object for target tab.")
    elif target_tab_id:
         print("Note has no specific page associated; tab opened/selected.")

def update_note_buttons_state():
    """Enables/disables Edit and Delete note buttons based on selection."""
    # ... (Keep previous implementation - it should work with the global selected_note_id) ...
    global notes_text_widget, selected_note_id; can_edit_delete = selected_note_id is not None
    try:
        notes_tab_frame = notes_text_widget.master.master; button_frame = notes_tab_frame.children.get('notes_button_frame')
        if button_frame: edit_button = button_frame.children.get('edit_note_button'); delete_button = button_frame.children.get('delete_note_button')
        if edit_button: edit_button.config(state=tk.NORMAL if can_edit_delete else tk.DISABLED)
        if delete_button: delete_button.config(state=tk.NORMAL if can_edit_delete else tk.DISABLED)
    except: pass # Ignore errors finding buttons
def on_viewer_tab_changed(event=None):
    """Callback when the active viewer tab changes. Updates the details panel AND menu states,
       ignoring the special Search Results tab.""" # Updated docstring
    global viewer_notebook, tab_states, search_results_tab_id # Need search result tab id

    try:
        if not viewer_notebook: return # Exit if notebook doesn't exist
        active_tab_id = viewer_notebook.select()
        print(f"Viewer tab changed to: {active_tab_id}") # Debug

        # --- ADD CHECK: Is this the Search Results Tab? ---
        if active_tab_id == search_results_tab_id:
            print("  Tab change is TO Search Results tab. Details panel state unchanged.")
            # Don't update details panel - keep showing details of last document.
            # Don't update favorite menu - it depends on document tabs.
            return # Stop processing here for search results tab
        # --- END CHECK ---

        # --- Process normally for DOCUMENT tabs ---
        state = tab_states.get(active_tab_id)
        doc_id_to_display = None
        if state and 'doc_id' in state:
            doc_id_to_display = state['doc_id']

        if doc_id_to_display is not None:
            print(f"Tab change updating details for Doc ID: {doc_id_to_display}")
            update_details_panel(doc_id_to_display)
        else:
            # This case might happen if a non-document tab is somehow selected,
            # or if the last document tab was just closed.
            print("Active tab is not Search Results and has no associated document state. Clearing details.")
            clear_details_panel()

        # Update favorite menu state ONLY if it's a document tab (or no tab)
        update_add_favorite_menu_state()

    except tk.TclError:
        # Error occurs if notebook has no tabs left (e.g., last one closed)
        print("No active viewer tab found (TclError). Clearing details.")
        clear_details_panel()
        update_add_favorite_menu_state() # Update menu state when no tabs active
    except Exception as e:
        print(f"Error in on_viewer_tab_changed: {e}")
        clear_details_panel()
        update_add_favorite_menu_state() # Update menu state on errors too

def on_tree_select(event=None):
    """Handles selection change in the FILE tree. Updates details panel if a file is selected."""
    global file_tree
    selected_iid = file_tree.focus()
    if not selected_iid: return # Do nothing if selection cleared in tree

    item_type = file_tree.set(selected_iid, "type")
    doc_id = None

    if item_type == "file":
        filepath = file_tree.set(selected_iid, "path")
        try: # Look up doc_id
            conn = sqlite3.connect(DATABASE_FILE); cursor = conn.cursor()
            cursor.execute("SELECT id FROM documents WHERE filepath = ?", (filepath,)); result = cursor.fetchone(); conn.close()
            if result: doc_id = result[0]
            else: print(f"Warning: File '{filepath}' selected in tree but not in DB index.")
        except Exception as e: print(f"Error looking up doc_id for tree selection: {e}")

        # Update details panel based on the file selected in the tree
        print(f"Tree selection updating details for Doc ID: {doc_id}") # Debug
        update_details_panel(doc_id) # <<< CRITICAL CALL

    elif item_type == "folder":
        # Folder selected: DO NOTHING to the details panel. Let active tab control.
        print(f"Folder selected: {file_tree.set(selected_iid, 'path')}. Details panel unchanged.")
        pass

# --- Functions to Update the RIGHT Pane ---
def update_details_panel(doc_id):
    """Fetches and displays details in ALL tabs of the right panel."""
    # ... (Keep previous implementation) ...
    global currently_displayed_doc_id # Access global
    currently_displayed_doc_id = doc_id # Store the ID being displayed
    update_metadata_tab(doc_id)
    update_links_tab(doc_id)
    update_notes_tab(doc_id)
    update_outline_tab(doc_id)

def clear_details_panel():
    """Clears ALL tabs in the details panel."""
    # ... (Keep previous implementation) ...
    global currently_displayed_doc_id # Access global
    currently_displayed_doc_id = None # Clear stored ID
    update_metadata_tab(None)
    update_links_tab(None)
    update_notes_tab(None)
    update_outline_tab(None)

def update_metadata_tab(doc_id):
    """Updates the 'Metadata' tab."""
    global metadata_widgets

    details = get_document_details(doc_id) if doc_id else None

    # Define the map based on SELECT order in get_document_details
    # 0=id, 1=filename, 2=filepath, 3=manufacturer, 4=device_model, 5=document_type,
    # 6=keywords, 7=revision_number, 8=revision_date, 9=status,
    # 10=applicable_models, 11=associated_test_equipment
    fields_map = {
        'filename': 1, 'filepath': 2, 'manufacturer': 3,
        'device_model': 4, 'document_type': 5, 'keywords': 6,
        'revision_number': 7, 'revision_date': 8, 'status': 9, # New
        'applicable_models': 10, 'associated_test_equipment': 11 # New
    }

    # Iterate through the widgets dictionary we created
    for key, widget in metadata_widgets.items():
        if key == 'edit_button': continue # Skip the button itself

        index = fields_map.get(key)
        if widget and index is not None:
            # Check if details is None (clearing) or get the value
            value = details[index] if details else ""
            # Ensure value is a string before setting, handle None/empty
            display_value = str(value) if value is not None and value != '' else "N/A" # Display N/A correctly
            try:
                widget.config(text=display_value)
            except Exception as e:
                print(f"Error configuring widget for key '{key}': {e}") # Debugging output
        elif widget:
            try: widget.config(text="N/A") # Clear if key not found
            except Exception as e: print(f"Error clearing widget for key '{key}': {e}")

    # Enable/disable Edit button
    edit_button = metadata_widgets.get('edit_button')
    if edit_button:
        try: edit_button.config(state=tk.NORMAL if doc_id else tk.DISABLED)
        except Exception as e: print(f"Err config edit button: {e}")

def update_links_tab(doc_id):
    """Updates the 'Related Links' tab."""
    # ... (Keep previous implementation) ...
    global links_listbox, links_map, root
    if not links_listbox: return
    links_listbox.delete(0, tk.END); links_map.clear()
    linked_docs = get_linked_documents(doc_id) if doc_id else []
    if linked_docs:
         for list_index, (target_id, target_filename, target_filepath, description) in enumerate(linked_docs):
              display_text = f"{target_filename}" + (f"  ({description[:30]}{'...' if len(description)>30 else ''})" if description else "")
              links_listbox.insert(tk.END, display_text); links_map[list_index] = {'id': target_id, 'filepath': target_filepath}
    elif doc_id: links_listbox.insert(tk.END, "(No linked documents)")
    try: # Update button states
        links_tab_frame = links_listbox.master.master; button_frame = links_tab_frame.children.get('links_button_frame')
        if button_frame:
            add_button = button_frame.children.get('add_link_button')
            suggest_button = button_frame.children.get('suggest_link_button')
            add_button.config(state=tk.NORMAL if doc_id else tk.DISABLED)
            if suggest_button: suggest_button.config(state=tk.NORMAL if doc_id else tk.DISABLED)
    except Exception as e: print(f"Err accessing Links buttons state: {e}")
    update_remove_link_button_state()


def update_remove_link_button_state(event=None):
     """Enables/disables the remove link button based on listbox selection."""
     # ... (Keep previous implementation) ...
     global links_listbox
     if not links_listbox: return
     remove_button = None
     try: links_tab_frame = links_listbox.master.master; button_frame = links_tab_frame.children.get('links_button_frame'); remove_button = button_frame.children.get('remove_link_button')
     except: pass
     if not remove_button: return
     selected_indices = links_listbox.curselection(); can_remove = False
     if selected_indices:
          try:
               if links_listbox.get(selected_indices[0]) != "(No linked documents)": can_remove = True
          except: pass
     remove_button.config(state=tk.NORMAL if can_remove else tk.DISABLED)


def update_notes_tab(doc_id):
    """Updates the 'Notes' tab. Shows all notes if doc_id is None."""
    global notes_text_widget, selected_note_id

    if not notes_text_widget: return
    selected_note_id = None; notes_text_widget.tag_remove("selected_note", "1.0", tk.END)
    notes_text_widget.config(state=tk.NORMAL); notes_text_widget.delete('1.0', tk.END)

    notes_to_display = []
    is_all_notes_view = False

    if doc_id is not None:
        # Get notes for specific document, including page number
        conn = sqlite3.connect(DATABASE_FILE); cursor = conn.cursor()
        try:
             cursor.execute("""
                 SELECT note_id, note_text, created_timestamp, page_number FROM notes
                 WHERE doc_id = ? ORDER BY page_number, created_timestamp DESC
             """, (doc_id,))
             notes_to_display = cursor.fetchall()
        except sqlite3.Error as e: print(f"DB error getting notes for doc {doc_id}: {e}")
        finally: conn.close()
    else:
        # Fetch ALL notes with filename and page number
        is_all_notes_view = True
        conn = sqlite3.connect(DATABASE_FILE); cursor = conn.cursor()
        try:
             cursor.execute("""
                 SELECT n.note_id, n.note_text, n.created_timestamp, d.filename, n.page_number
                 FROM notes n JOIN documents d ON n.doc_id = d.id
                 ORDER BY d.filename, n.page_number, n.created_timestamp DESC
             """)
             notes_to_display = cursor.fetchall()
        except sqlite3.Error as e: print(f"DB error getting all notes: {e}")
        finally: conn.close()


    if notes_to_display:
        current_filename_header = None # Used only in 'all notes' view
        for note_data in notes_to_display:
            page_str = " (Page ?)" # Default page string
            if is_all_notes_view:
                 note_id, note_text, timestamp, filename, page_num = note_data
                 if filename != current_filename_header:
                      notes_text_widget.insert(tk.END, f"\n--- {filename} ---\n", ("heading",))
                      current_filename_header = filename
                 if page_num is not None: page_str = f" (Page {page_num + 1})" # 1-based for display
            else: # Single document view
                 note_id, note_text, timestamp, page_num = note_data
                 if page_num is not None: page_str = f" (Page {page_num + 1})"

            try: ts_str = time.strftime('%Y-%m-%d %H:%M', time.localtime(timestamp))
            except ValueError: ts_str = "Invalid Date"

            note_tag = f"note_{note_id}"
            start_index = notes_text_widget.index(tk.INSERT) # Get index before inserting this note
            # Insert content
            notes_text_widget.insert(tk.END, f"[{ts_str}]{page_str} - ID:{note_id}\n", ("timestamp",)) # Timestamp line (don't apply note_tag here initially)
            notes_text_widget.insert(tk.END, f"{note_text}\n\n", ("note_content",))
            end_index = notes_text_widget.index(tk.INSERT) # Get index after inserting
            # Apply the unique note ID tag to the entire block just inserted
            notes_text_widget.tag_add(note_tag, start_index, end_index)

    elif doc_id is not None:
        notes_text_widget.insert(tk.END, "(No notes found for this document)", "placeholder")
    else:
         notes_text_widget.insert(tk.END, "(No notes found in database. Select a document to add notes)", "placeholder")

    notes_text_widget.config(state=tk.DISABLED) # Disable editing

    # Update button states
    try:
         notes_tab_frame = notes_text_widget.master.master; button_frame = notes_tab_frame.children.get('notes_button_frame')
         if button_frame: add_note_button = button_frame.children.get('add_note_button'); add_note_button.config(state=tk.NORMAL if doc_id is not None else tk.DISABLED)
    except: pass
    update_note_buttons_state() # Reset Edit/Delete based on selection (none initially)
    
def update_outline_tab(doc_id):
    """Updates the 'Outline' (TOC) tab."""
    # ... (Keep previous implementation) ...
    global outline_tree
    if not outline_tree: return
    outline_tree.delete(*outline_tree.get_children())
    details = get_document_details(doc_id) if doc_id else None
    if not details: outline_tree.insert('', tk.END, text="(Select a document)"); return
    filepath = details[2]; ext = os.path.splitext(filepath)[1].lower()
    if ext == '.pdf':
        try:
            with fitz.open(filepath) as doc: toc = doc.get_toc(simple=False)
            if toc:
                parent_map = {'0': ''}
                for item in toc:
                    level, title, page, dest_info = item[:4]
                    parent_iid = parent_map.get(str(level - 1), '')
                    page_num_zero_based = page - 1
                    # Store page number as the primary tag for navigation
                    iid = outline_tree.insert(parent_iid, tk.END, text=f" {title}", values=[page], open=False, tags=(page_num_zero_based,))
                    parent_map[str(level)] = iid
            else: outline_tree.insert('', tk.END, text="(No outline found in PDF)")
        except Exception as e: print(f"Err getting PDF TOC: {e}"); outline_tree.insert('', tk.END, text="(Error reading PDF outline)")
    elif ext == '.docx' and DOCX_ENABLED: outline_tree.insert('', tk.END, text="(DOCX outline not yet supported)")
    else: outline_tree.insert('', tk.END, text="(Outline unavailable for this file type)")


# --- Dialog Functions ---
def open_edit_metadata_dialog():
    """Opens dialog to edit metadata for the selected FILE TREE item."""
    global file_tree, root
    selected_iid = file_tree.focus()
    if not selected_iid: messagebox.showinfo("Edit Metadata", "Select a document first."); return
    item_type = file_tree.set(selected_iid, "type")
    if item_type != "file": messagebox.showinfo("Edit Metadata", "Select a document file."); return
    filepath = file_tree.set(selected_iid, "path")
    conn = sqlite3.connect(DATABASE_FILE); cursor = conn.cursor(); cursor.execute("SELECT id FROM documents WHERE filepath = ?", (filepath,)); result = cursor.fetchone(); conn.close()
    if not result: messagebox.showerror("Error", "Selected file not found in index."); return
    doc_id = result[0]

    details = get_document_details(doc_id)
    if not details: messagebox.showerror("Error", f"Could not get details for ID {doc_id}."); return
    # Unpack all details including new ones
    _, filename, _, manuf, model, dtype, keywords, rev_num, rev_date, status, app_models, test_equip = details

    dialog = Toplevel(root); dialog.title(f"Edit Metadata - {filename}");
    dialog.geometry("550x400") # <<<< Increased size
    dialog.transient(root); dialog.grab_set(); dialog.resizable(False, True) # Allow vertical resize

    # Use StringVars to hold current/new values
    fields_vars = {
        "Manufacturer:": tk.StringVar(value=manuf or ""),
        "Device Model:": tk.StringVar(value=model or ""),
        "Document Type:": tk.StringVar(value=dtype or ""),
        "Revision:": tk.StringVar(value=rev_num or ""), # New
        "Rev Date (YYYY-MM-DD):": tk.StringVar(value=rev_date or ""), # New
        "Status:": tk.StringVar(value=status or ""), # New - Use Combobox below
        "Other Models:": tk.StringVar(value=app_models or ""), # New
        "Test Equip:": tk.StringVar(value=test_equip or ""), # New
        "Keywords (comma-sep):": tk.StringVar(value=keywords or "")
    }
    # Define standard status options
    status_options = ["", "Active", "Draft", "Superseded", "Archived", "Reference Only"]

    frame = ttk.Frame(dialog, padding="10"); frame.pack(expand=True, fill=tk.BOTH)

    row_num = 0
    for label_text, var in fields_vars.items():
        lbl = ttk.Label(frame, text=label_text)
        lbl.grid(row=row_num, column=0, sticky="nw", padx=5, pady=3) # Use nw anchor

        if label_text == "Status:": # Use Combobox for Status
             # Ensure current status is in options, add if not (for legacy data)
             current_status = var.get()
             if current_status and current_status not in status_options:
                  status_options.insert(1, current_status) # Add after empty
             widget = ttk.Combobox(frame, textvariable=var, values=status_options, state="normal", width=38)
        elif label_text == "Keywords (comma-sep):": # Use Text widget for keywords for more space? No, keep Entry.
             widget = ttk.Entry(frame, textvariable=var, width=40)
        else: # Standard Entry
            widget = ttk.Entry(frame, textvariable=var, width=40)

        widget.grid(row=row_num, column=1, sticky="ew", padx=5, pady=3)
        row_num += 1

    frame.columnconfigure(1, weight=1) # Allow entry column to expand

    button_frame = ttk.Frame(frame)
    button_frame.grid(row=row_num, column=0, columnspan=2, pady=15)

    def save_metadata():
        new_metadata = {
            'manufacturer': fields_vars["Manufacturer:"].get().strip() or None,
            'device_model': fields_vars["Device Model:"].get().strip() or None,
            'document_type': fields_vars["Document Type:"].get().strip() or None,
            'keywords': fields_vars["Keywords (comma-sep):"].get().strip() or None,
            'revision_number': fields_vars["Revision:"].get().strip() or None, # New
            'revision_date': fields_vars["Rev Date (YYYY-MM-DD):"].get().strip() or None, # New
            'status': fields_vars["Status:"].get().strip() or None, # New
            'applicable_models': fields_vars["Other Models:"].get().strip() or None, # New
            'associated_test_equipment': fields_vars["Test Equip:"].get().strip() or None, # New
        }
        if update_document_metadata(doc_id, new_metadata):
            messagebox.showinfo("Success", "Metadata updated.", parent=dialog)
            update_details_panel(doc_id) # Refresh right panel display
            # No need to refresh tree unless displayed columns change
            dialog.destroy()
        else: messagebox.showerror("Error", "Failed to update DB.", parent=dialog)

    save_button = ttk.Button(button_frame, text="Save", command=save_metadata); save_button.pack(side=tk.LEFT, padx=10)
    cancel_button = ttk.Button(button_frame, text="Cancel", command=dialog.destroy); cancel_button.pack(side=tk.LEFT, padx=10)

    dialog.wait_window()
def open_batch_edit_dialog():
    """Opens dialog for batch editing metadata of selected files."""
    global file_tree, root
    selected_iids = file_tree.selection()
    if not selected_iids: messagebox.showwarning("Batch Edit", "Select one or more files first."); return
    file_iids = [iid for iid in selected_iids if file_tree.set(iid, "type") == "file"]
    if not file_iids: messagebox.showwarning("Batch Edit", "Selection contains no document files."); return
    num_files = len(file_iids)

    dialog = Toplevel(root); dialog.title(f"Batch Edit Metadata ({num_files} files)");
    dialog.geometry("550x420"); dialog.transient(root); dialog.grab_set(); dialog.resizable(False, True)

    # --- Use grid layout within the main frame ---
    frame = ttk.Frame(dialog, padding="10")
    frame.pack(expand=True, fill=tk.BOTH) # Pack the main frame into the dialog

    # --- Place top label using grid ---
    top_label = ttk.Label(frame, text=f"Editing {num_files} selected files. Only checked fields will be updated.")
    top_label.grid(row=0, column=0, columnspan=4, sticky="ew", pady=(0, 10)) # Span all 4 columns we plan to use

    edit_fields = {}; editable_fields = { 'manufacturer': "Manufacturer:", 'document_type': "Document Type:", 'revision_number': "Revision:", 'revision_date': "Rev Date (YYYY-MM-DD):", 'status': "Status:", 'keywords': "Keywords (Append/Replace):", 'applicable_models': "Other Models:", 'associated_test_equipment': "Test Equip:" }; status_options = ["", "Active", "Draft", "Superseded", "Archived", "Reference Only"]

    row_num = 1 # Start grid rows from 1 (after the top label)
    for db_key, label_text in editable_fields.items():
        print(f"  Processing field: {db_key} (Row: {row_num})")
        try:
            var_check = tk.BooleanVar(value=False)
            var_value = tk.StringVar()

            # Column 0: Checkbox
            check = ttk.Checkbutton(frame, variable=var_check)
            check.grid(row=row_num, column=0, sticky="w", padx=(5,0), pady=3)

            # Column 1: Label
            label = ttk.Label(frame, text=label_text)
            label.grid(row=row_num, column=1, sticky="w", padx=(0,5), pady=3) # Adjusted padding

            widget = None
            if db_key == "status":
                 widget = ttk.Combobox(frame, textvariable=var_value, values=status_options, state="readonly", width=38)
            elif db_key == "keywords":
                 widget = ttk.Entry(frame, textvariable=var_value, width=40)
                 kw_mode_var = tk.StringVar(value="replace")
                 # Use a NEW frame for radio buttons, placed with grid
                 kw_frame = ttk.Frame(frame)
                 ttk.Radiobutton(kw_frame, text="Replace", variable=kw_mode_var, value="replace").pack(side=tk.LEFT)
                 ttk.Radiobutton(kw_frame, text="Append", variable=kw_mode_var, value="append").pack(side=tk.LEFT, padx=5)
                 # Grid the frame containing radios
                 kw_frame.grid(row=row_num, column=3, sticky="w", padx=5, pady=3)
                 edit_fields[db_key + "_mode"] = kw_mode_var
            else:
                 widget = ttk.Entry(frame, textvariable=var_value, width=40)

            # Column 2: Input Widget
            widget.grid(row=row_num, column=2, sticky="ew", padx=5, pady=3)

            edit_fields[db_key] = {'check': var_check, 'value': var_value, 'widget': widget}
            print(f"    Widgets created and gridded for {db_key}")
            row_num += 1

        except Exception as e:
            print(f"  ERROR creating widgets for field '{db_key}': {e}")
            break

    print(f"Finished creating field widgets. edit_fields populated: {bool(edit_fields)}")

    # --- Configure column weights for resizing ---
    frame.columnconfigure(2, weight=1) # Allow input column (index 2) to expand horizontally

    # --- Action Buttons (Use grid) ---
    button_frame = ttk.Frame(frame)
    # Place this frame using grid in the next row, spanning columns
    button_frame.grid(row=row_num, column=0, columnspan=4, pady=15)
    print("Creating Action Buttons using pack inside button_frame")

    # Define apply_batch_changes nested function (keep as is)
    def apply_batch_changes():
        # ... (logic to collect updates, confirm, process files) ...
        updates_to_make = {}; keyword_mode = 'replace'
        for db_key, field_data in edit_fields.items(): # Collect data
            if db_key.endswith("_mode"): continue
            if field_data['check'].get():
                value = field_data['value'].get().strip()
                updates_to_make[db_key] = None if db_key != 'keywords' and not value else value
                if db_key == 'keywords': mode_var = edit_fields.get("keywords_mode"); keyword_mode = mode_var.get() if mode_var else 'replace'
        if not updates_to_make: messagebox.showinfo("No Changes", "No fields selected.", parent=dialog); return
        confirm = messagebox.askyesno("Confirm Batch Update", f"Update fields: {', '.join(updates_to_make.keys())} for {num_files} files?", parent=dialog)
        if not confirm: return
        conn = sqlite3.connect(DATABASE_FILE); cursor = conn.cursor(); updated_count = 0; error_count = 0
        status_bar_label.config(text=f"Batch updating {num_files} files..."); root.update_idletasks()
        for iid in file_iids: # Process files
            filepath = file_tree.set(iid, "path")
            try:
                cursor.execute("SELECT id, manufacturer, device_model, document_type, keywords, revision_number, revision_date, status, applicable_models, associated_test_equipment FROM documents WHERE filepath = ?", (filepath,)); current_data = cursor.fetchone()
                if not current_data: print(f"Skipping {filepath}: Not in DB."); error_count += 1; continue
                current_doc_id = current_data[0]; current_metadata = { 'manufacturer': current_data[1], 'device_model': current_data[2], 'document_type': current_data[3], 'keywords': current_data[4], 'revision_number': current_data[5], 'revision_date': current_data[6], 'status': current_data[7], 'applicable_models': current_data[8], 'associated_test_equipment': current_data[9] }
                new_metadata = current_metadata.copy()
                for db_key, new_value in updates_to_make.items(): # Apply changes
                     if db_key == 'keywords':
                          current_kws_str = current_metadata.get('keywords', '') or ''; current_kws = set(kw.strip() for kw in current_kws_str.split(',') if kw.strip()); new_kws_to_process = set(kw.strip() for kw in new_value.split(',') if kw.strip())
                          final_kws = current_kws.union(new_kws_to_process) if keyword_mode == 'append' else new_kws_to_process
                          new_metadata['keywords'] = ','.join(sorted(list(final_kws))) or None
                     else: new_metadata[db_key] = new_value
                if update_document_metadata(current_doc_id, new_metadata): updated_count += 1 # Update DB
                else: error_count += 1
            except Exception as e: print(f"Error batch processing {filepath}: {e}"); error_count += 1
        conn.close(); status_bar_label.config(text=f"Batch update complete. Updated: {updated_count}, Errors: {error_count}. Ready.")
        messagebox.showinfo("Batch Update Complete", f"Successfully updated {updated_count} files.\nErrors on {error_count} files.", parent=dialog)
        on_tree_select() # Refresh details panel
        dialog.destroy()

    # Use pack *within* the button_frame
    apply_button = ttk.Button(button_frame, text="Apply Changes", command=apply_batch_changes)
    apply_button.pack(side=tk.LEFT, padx=10)
    cancel_button = ttk.Button(button_frame, text="Cancel", command=dialog.destroy)
    cancel_button.pack(side=tk.LEFT, padx=10)
    print("Action Buttons created and packed inside button_frame.")

    dialog.wait_window()
def open_add_link_dialog():
    """Opens dialog to add a link WITH description for selected FILE TREE item."""
    global file_tree, root
    selected_iid = file_tree.focus()
    if not selected_iid: messagebox.showinfo("Add Link", "Select the source document from the file tree first."); return
    item_type = file_tree.set(selected_iid, "type")
    if item_type != "file": messagebox.showinfo("Add Link", "Select a source document file, not a folder."); return
    filepath = file_tree.set(selected_iid, "path")
    # Get doc_id from DB based on filepath
    conn = sqlite3.connect(DATABASE_FILE); cursor = conn.cursor(); cursor.execute("SELECT id, filename FROM documents WHERE filepath = ?", (filepath,)); result = cursor.fetchone(); conn.close()
    if not result: messagebox.showerror("Error", "Selected file not found in database index."); return
    source_doc_id, source_filename = result
    # ... (Rest of the dialog logic - same as before, using source_doc_id) ...
    dialog = Toplevel(root); dialog.title(f"Link from: {source_filename[:40]}..."); dialog.geometry("500x450"); dialog.transient(root); dialog.grab_set()
    frame = ttk.Frame(dialog, padding="10"); frame.pack(expand=True, fill=tk.BOTH)
    search_frame = ttk.Frame(frame); search_frame.pack(fill=tk.X, pady=(0, 5)); ttk.Label(search_frame, text="Search Target:").pack(side=tk.LEFT, padx=(0, 5)); link_search_entry = ttk.Entry(search_frame); link_search_entry.pack(side=tk.LEFT, expand=True, fill=tk.X, padx=5)
    link_results_list = Listbox(frame, height=8, exportselection=False); link_results_list.pack(expand=True, fill=tk.BOTH, pady=5); link_results_map = {}
    def perform_link_search(event=None):
        query = link_search_entry.get();
        if not query: return
        results_data = search_documents(query); link_results_list.delete(0, tk.END); link_results_map.clear(); count = 0 # search_documents now returns full data
        for doc_id, filename, _, manuf, model, _ in results_data:
             if doc_id == source_doc_id: continue
             display_text = f"{filename} ({manuf or '?'} / {model or '?'})"; link_results_list.insert(tk.END, display_text); link_results_map[count] = doc_id; count += 1
    link_search_entry.bind("<Return>", perform_link_search); search_button = ttk.Button(search_frame, text="Search", command=perform_link_search); search_button.pack(side=tk.LEFT, padx=5)
    desc_frame = ttk.Frame(frame); desc_frame.pack(fill=tk.X, pady=5); ttk.Label(desc_frame, text="Link Desc (Optional):").pack(side=tk.LEFT, padx=(0, 5)); link_desc_entry = ttk.Entry(desc_frame); link_desc_entry.pack(side=tk.LEFT, expand=True, fill=tk.X)
    button_frame = ttk.Frame(frame); button_frame.pack(pady=5)
    def link_selected():
        selected_indices = link_results_list.curselection();
        if not selected_indices: messagebox.showwarning("Select Target", "Select a target document.", parent=dialog); return
        listbox_index = selected_indices[0]; link_description = link_desc_entry.get().strip() or None
        if listbox_index in link_results_map:
            target_doc_id = link_results_map[listbox_index]
            if add_document_link(source_doc_id, target_doc_id, link_description): update_details_panel(source_doc_id); dialog.destroy()
        else: messagebox.showerror("Error", "Internal mapping error.", parent=dialog)
    link_button = ttk.Button(button_frame, text="Link Selected", command=link_selected); link_button.pack(side=tk.LEFT, padx=10)
    cancel_button = ttk.Button(button_frame, text="Cancel", command=dialog.destroy); cancel_button.pack(side=tk.LEFT, padx=10)
    link_search_entry.focus_set(); dialog.wait_window()


def open_manage_paths_dialog():
    """Opens dialog to add/remove scan paths."""
    global root
    dialog = Toplevel(root); dialog.title("Manage Scan Paths"); dialog.geometry("600x400")
    dialog.transient(root); dialog.grab_set(); dialog.resizable(True, True)

    frame = ttk.Frame(dialog, padding="10"); frame.pack(expand=True, fill=tk.BOTH)

    # Listbox Frame
    list_frame = ttk.Frame(frame); list_frame.pack(expand=True, fill=tk.BOTH, pady=(0, 10))
    ttk.Label(list_frame, text="Configured Scan Paths:").pack(anchor='w')
    path_listbox = Listbox(list_frame, height=10, exportselection=False)
    path_scrollbar = ttk.Scrollbar(list_frame, orient=tk.VERTICAL, command=path_listbox.yview)
    path_listbox.config(yscrollcommand=path_scrollbar.set)
    path_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
    path_listbox.pack(side=tk.LEFT, expand=True, fill=tk.BOTH)

    def refresh_path_list():
        path_listbox.delete(0, tk.END)
        paths = get_scan_paths()
        for p in paths: path_listbox.insert(tk.END, p)

    refresh_path_list() # Initial population

    # Button Frame
    button_frame = ttk.Frame(frame)
    button_frame.pack(fill=tk.X)

    def add_path():
        # Use filedialog to browse for directory
        new_path = filedialog.askdirectory(title="Select Directory to Add", parent=dialog)
        if new_path: # Path selected
             if add_scan_path(new_path):
                  refresh_path_list() # Update listbox if successful
             # Else: Error message shown by add_scan_path

    def remove_path():
        selected_indices = path_listbox.curselection()
        if not selected_indices:
            messagebox.showwarning("Remove Path", "Please select a path from the list to remove.", parent=dialog)
            return
        path_to_remove = path_listbox.get(selected_indices[0])
        confirm = messagebox.askyesno("Confirm Removal", f"Are you sure you want to remove this path?\n'{path_to_remove}'", parent=dialog)
        if confirm:
            if remove_scan_path(path_to_remove):
                 refresh_path_list() # Update listbox
            # Else: Error message shown by remove_scan_path

    add_button = ttk.Button(button_frame, text="Add Path...", command=add_path)
    add_button.pack(side=tk.LEFT, padx=5)
    remove_button = ttk.Button(button_frame, text="Remove Selected", command=remove_path)
    remove_button.pack(side=tk.LEFT, padx=5)
    close_button = ttk.Button(button_frame, text="Close", command=dialog.destroy)
    close_button.pack(side=tk.RIGHT, padx=5) # Close on right

    dialog.wait_window()



# --- Link/Outline/Note Actions ---
def on_related_doc_double_click(event):
    """Opens the double-clicked related document."""
    # ... (Keep previous implementation) ...
    global links_listbox, links_map
    widget = event.widget
    selected_indices = widget.curselection()
    if not selected_indices: return
    listbox_index = selected_indices[0]
    if listbox_index in links_map:
         target_info = links_map[listbox_index]
         open_document_in_tab(target_info['id'])
    else:
         try:
             if links_listbox.get(listbox_index) == "(No linked documents)": pass
             else: print("Error: Dbl-click index not in map.")
         except: pass # Ignore errors
def remove_selected_link():
     """Removes the selected link for the item selected in the FILE TREE."""
     global links_listbox, links_map, file_tree # Use file_tree selection
     selected_list_indices = links_listbox.curselection()
     if not selected_list_indices: messagebox.showinfo("Remove Link", "Select a link first."); return
     listbox_index = selected_list_indices[0]
     try:
         if links_listbox.get(listbox_index) == "(No linked documents)": return
     except: return
     selected_tree_iid = file_tree.focus() # Get source from file tree
     if not selected_tree_iid or file_tree.set(selected_tree_iid, "type") != "file":
          messagebox.showwarning("Select Source", "Select the source document file in the tree first."); return
     filepath = file_tree.set(selected_tree_iid, "path")
     conn = sqlite3.connect(DATABASE_FILE); cursor = conn.cursor(); cursor.execute("SELECT id FROM documents WHERE filepath = ?", (filepath,)); result = cursor.fetchone(); conn.close()
     if not result: messagebox.showerror("Error", "Source file not in index."); return
     source_doc_id = result[0]

     if listbox_index in links_map:
          target_info = links_map[listbox_index]
          confirm = messagebox.askyesno("Confirm Removal", f"Remove link to:\n{links_listbox.get(listbox_index)}?")
          if confirm:
              if remove_document_link(source_doc_id, target_info['id']): update_links_tab(source_doc_id) # Refresh
     else: print("Error: Sel link index not in map.")


def on_outline_double_click(event):
    """Jumps to the page selected in the outline treeview.
       Opens/selects the relevant document tab if necessary."""
    global outline_tree, currently_displayed_doc_id, tab_states, open_files_map, viewer_notebook, root

    # --- Step 1: Check if an outline is relevant (document selected) ---
    if currently_displayed_doc_id is None:
        print("Outline double-click ignored: No document context.")
        return

    # --- Step 2: Get the selected outline item and target page ---
    item_iid = outline_tree.focus()
    if not item_iid: return

    page_num_zero_based = None
    try: # Get page number from tag
        item_tags = outline_tree.item(item_iid, "tags")
        if item_tags: page_num_zero_based = int(item_tags[0])
        else: print("Outline nav failed: No page tag."); return
    except: print("Error processing outline tag."); return

    if page_num_zero_based is None: return # Should be caught above

    print(f"Outline double-click: Requesting page {page_num_zero_based + 1} for Doc ID {currently_displayed_doc_id}")

    # --- Step 3: Find OR Open the Tab for the outline's document ID ---
    target_tab_id = None
    target_state = None
    doc_id_of_outline = currently_displayed_doc_id # Use the stored ID

    # Check if a tab for this doc_id is already open
    for tab_id, state in tab_states.items():
        if state.get('doc_id') == doc_id_of_outline:
            try: # Verify tab exists
                 if tab_id in viewer_notebook.tabs():
                      target_tab_id = tab_id; target_state = state; break
            except: pass

    # If tab wasn't found open, OPEN IT NOW
    if target_tab_id is None:
        print(f"Outline double-click: Tab for Doc ID {doc_id_of_outline} not found. Opening...")
        # We need the filepath to open it
        details = get_document_details(doc_id_of_outline)
        if details and os.path.exists(details[2]):
             target_tab_id = open_document_in_tab(doc_id_of_outline) # Open the tab
             if target_tab_id:
                  target_state = tab_states.get(target_tab_id) # Get the state of the NEWLY opened tab
             else:
                  print("Outline nav failed: Failed to open required document tab.")
                  messagebox.showerror("Error", f"Could not open the required document (ID: {doc_id_of_outline}).")
                  return
        else:
             print(f"Outline nav failed: Could not get details or find file for Doc ID {doc_id_of_outline}.")
             messagebox.showerror("Error", f"Could not find the document file for ID {doc_id_of_outline} to open.")
             return

    # --- Step 4: Check if the Target Tab is Valid for Navigation (PDF loaded?) ---
    if target_tab_id is None or target_state is None:
        print(f"Outline nav failed: Could not find/create state for Doc ID {doc_id_of_outline}.")
        return # Should not happen if open_document_in_tab worked

    doc_obj = target_state.get('doc_obj')
    if not doc_obj or not isinstance(doc_obj, fitz.Document):
        # This might happen if the file opened is not a PDF or failed initial load
        print(f"Outline nav failed: Target tab {target_tab_id} does not contain a loaded PDF.")
        messagebox.showinfo("Navigation Info", "Outline navigation requires the target tab to be a successfully loaded PDF.")
        # Select the tab anyway so user sees the error/non-PDF content
        if viewer_notebook.select() != target_tab_id: viewer_notebook.select(target_tab_id)
        return

    # --- Step 5: Validate Page Number and Navigate ---
    doc_length = len(doc_obj)
    if 0 <= page_num_zero_based < doc_length:
        # Select the target tab visually if it wasn't already active
        if viewer_notebook.select() != target_tab_id:
            print(f"Outline double-click: Selecting target tab {target_tab_id}...")
            viewer_notebook.select(target_tab_id)
            # IMPORTANT: Allow tab switch event to process which updates details panel
            # We might need to defer the page loading slightly? Or rely on update_history_and_load?
            root.update_idletasks()

        # Update the page number state and load using history function
        print(f"Outline double-click: Calling update_history_and_load for tab {target_tab_id}, page {page_num_zero_based}")
        update_history_and_load(target_tab_id, page_num_zero_based) # This handles history and calls load_pdf_page
        print(f"Outline navigation successful.")
    else:
        print(f"Outline nav failed: Target page {page_num_zero_based + 1} out of bounds.")
        messagebox.showwarning("Navigation Error", f"Target page {page_num_zero_based + 1} is out of bounds (Total: {doc_length}).")

def add_new_note():
     """Adds a new note for the currently selected document, including page number."""
     global file_tree, root, viewer_notebook, tab_states # Need viewer state
     selected_iid = file_tree.focus()
     if not selected_iid: messagebox.showinfo("Add Note", "Select a document from the file tree first."); return
     item_type = file_tree.set(selected_iid, "type")
     if item_type != "file": messagebox.showinfo("Add Note", "Select a document file, not a folder."); return
     filepath = file_tree.set(selected_iid, "path")
     # Get doc_id from DB
     conn = sqlite3.connect(DATABASE_FILE); cursor = conn.cursor(); cursor.execute("SELECT id FROM documents WHERE filepath = ?", (filepath,)); result = cursor.fetchone() # Keep connection open briefly
     if not result: messagebox.showerror("Error", "Selected file not found in database index."); conn.close(); return
     doc_id = result[0]

     # --- Get current page number from active tab (if any) ---
     current_page = None
     active_tab_id = get_active_tab_id()
     active_state = get_active_tab_state()
     # Check if the active tab actually corresponds to the selected doc_id
     if active_state and active_state.get('doc_id') == doc_id:
          current_page = active_state.get('page_num') # This is 0-based

     note_text = simpledialog.askstring("New Note", f"Enter note for page {current_page + 1 if current_page is not None else 'N/A'}:", parent=root)

     if note_text: # User entered text and didn't cancel
          note_text = note_text.strip()
          if note_text: # Ensure it's not just whitespace
               # Save to DB
               try:
                    timestamp = time.time()
                    cursor.execute("""
                        INSERT INTO notes (doc_id, page_number, note_text, created_timestamp)
                        VALUES (?, ?, ?, ?)""",
                       (doc_id, current_page, note_text, timestamp)) # Save page number
                    conn.commit()
                    print(f"Note added for doc {doc_id}, page {current_page}")
                    update_notes_tab(doc_id) # Refresh display
               except sqlite3.Error as e:
                    print(f"Database error adding note for doc {doc_id}: {e}"); conn.rollback()
                    messagebox.showerror("Error", "Failed to save note to database.")
          else: messagebox.showwarning("Empty Note", "Note cannot be empty.")

     conn.close() # Close connection

def delete_selected_note():
    """Deletes the note identified by the globally selected note ID."""
    global selected_note_id, file_tree # Need file_tree to refresh view

    if selected_note_id is None:
        messagebox.showinfo("Delete Note", "Please click on a note in the list first to select it.")
        return

    # Confirm deletion
    confirm = messagebox.askyesno("Confirm Delete", f"Are you sure you want to delete note ID: {selected_note_id}?")
    if confirm:
        if delete_note(selected_note_id): # Call DB function
            print(f"Note ID {selected_note_id} deleted from DB.")
            # Refresh notes display for the currently selected document (or all notes)
            selected_iid = file_tree.focus()
            doc_id_to_refresh = None
            if selected_iid and file_tree.set(selected_iid, "type") == "file":
                 filepath = file_tree.set(selected_iid, "path")
                 try: # Quick DB lookup for current doc_id
                      conn = sqlite3.connect(DATABASE_FILE); cursor = conn.cursor(); cursor.execute("SELECT id FROM documents WHERE filepath = ?", (filepath,)); result = cursor.fetchone(); conn.close()
                      if result: doc_id_to_refresh = result[0]
                 except: pass # Ignore lookup error, refresh will show all notes
            update_notes_tab(doc_id_to_refresh) # Refresh the notes view
        else:
            messagebox.showerror("Error", f"Failed to delete note ID: {selected_note_id} from database.")
    # Reset selection after attempt
    selected_note_id = None
    notes_text_widget.tag_remove("selected_note", "1.0", tk.END)
    update_note_buttons_state()

def edit_selected_note():
    """Edits the note identified by the globally selected note ID."""
    global selected_note_id, file_tree, root # Need root for dialog parent

    if selected_note_id is None:
        messagebox.showinfo("Edit Note", "Please click on a note in the list first to select it.")
        return

    # Get the current note text from the database
    current_note_text = ""
    conn = sqlite3.connect(DATABASE_FILE)
    cursor = conn.cursor()
    try:
        cursor.execute("SELECT note_text FROM notes WHERE note_id = ?", (selected_note_id,))
        result = cursor.fetchone()
        if result:
            current_note_text = result[0]
        else:
             messagebox.showerror("Error", f"Could not find note ID {selected_note_id} in database to edit.")
             conn.close()
             return
    except sqlite3.Error as e:
         messagebox.showerror("Database Error", f"Error fetching note text for editing:\n{e}")
         conn.close()
         return
    # Keep connection open for update

    # Prompt user with existing text
    new_note_text = simpledialog.askstring("Edit Note", "Edit your note:",
                                        initialvalue=current_note_text, parent=root)

    if new_note_text is not None: # User didn't cancel
        new_note_text = new_note_text.strip()
        if new_note_text:
            # Update the note in the database
            try:
                 cursor.execute("UPDATE notes SET note_text = ? WHERE note_id = ?", (new_note_text, selected_note_id))
                 conn.commit()
                 print(f"Note ID {selected_note_id} updated.")
                 # Refresh notes display
                 selected_iid = file_tree.focus()
                 doc_id_to_refresh = None
                 if selected_iid and file_tree.set(selected_iid, "type") == "file":
                      filepath = file_tree.set(selected_iid, "path")
                      # Quick lookup for doc_id
                      cursor.execute("SELECT id FROM documents WHERE filepath = ?", (filepath,)); result = cursor.fetchone()
                      if result: doc_id_to_refresh = result[0]
                 update_notes_tab(doc_id_to_refresh)
            except sqlite3.Error as e:
                 print(f"DB error updating note {selected_note_id}: {e}"); conn.rollback()
                 messagebox.showerror("Error", "Failed to update note in database.")
        else:
            messagebox.showwarning("Empty Note", "Note cannot be empty. Edit cancelled.", parent=root)

    conn.close() # Close connection after potential update or cancel
    # Reset selection after attempt
    selected_note_id = None
    notes_text_widget.tag_remove("selected_note", "1.0", tk.END)
    update_note_buttons_state()

def expand_all_outline():
     """Recursively expands all nodes in the outline tree."""
     global outline_tree
     if not outline_tree: return
     print("Expanding all outline nodes...")
     status_bar_label.config(text="Expanding outline...")
     root.update_idletasks()
     nodes_to_expand = list(outline_tree.get_children('')) # Start with top-level
     while nodes_to_expand:
          iid = nodes_to_expand.pop(0)
          if not outline_tree.item(iid, 'open'): # Check if not already open
               outline_tree.item(iid, open=True)
          # Add children to the list to process them
          nodes_to_expand.extend(outline_tree.get_children(iid))
     status_bar_label.config(text="Outline expanded. Ready.")
     print("Outline expansion complete.")


def collapse_all_outline():
     """Recursively collapses all nodes in the outline tree."""
     global outline_tree
     if not outline_tree: return
     print("Collapsing all outline nodes...")
     status_bar_label.config(text="Collapsing outline...")
     root.update_idletasks()
     nodes_to_collapse = list(outline_tree.get_children(''))
     while nodes_to_collapse:
          iid = nodes_to_collapse.pop(0)
          # Add children *before* collapsing the parent
          nodes_to_collapse.extend(outline_tree.get_children(iid))
          if outline_tree.item(iid, 'open'): # Check if currently open
               outline_tree.item(iid, open=False)
     status_bar_label.config(text="Outline collapsed. Ready.")
     print("Outline collapse complete.")


def filter_outline(event=None):
    """Filters the outline tree by rebuilding it with matching items."""
    global outline_tree, outline_search_entry, file_tree, root, status_bar_label
    if not outline_tree or not outline_search_entry: return

    query = outline_search_entry.get().lower().strip()

    # Get the source TOC data for the currently selected document
    selected_iid = file_tree.focus()
    toc_data = [] # List to hold original TOC items: [level, title, page, dest]
    original_filepath = None

    if selected_iid and file_tree.set(selected_iid, "type") == "file":
        filepath = file_tree.set(selected_iid, "path")
        original_filepath = filepath # Store for error messages
        ext = os.path.splitext(filepath)[1].lower()
        if ext == '.pdf':
             try:
                  with fitz.open(filepath) as doc:
                       toc_data = doc.get_toc(simple=False)
             except Exception as e:
                  print(f"Error getting TOC for outline filter: {e}")
                  toc_data = [] # Ensure it's an empty list on error

    # Clear the existing tree
    outline_tree.delete(*outline_tree.get_children())

    if not query:
        # If query is empty, rebuild the full outline
        update_outline_tab(get_selected_doc_id()) # Call the standard update function
        status_bar_label.config(text="Outline filter cleared. Ready.")
        return

    status_bar_label.config(text=f"Filtering outline for '{query}'...")
    root.update_idletasks()

    if not toc_data:
         outline_tree.insert('', tk.END, text="(No outline data to filter)")
         status_bar_label.config(text="Outline filtered (no data). Ready.")
         return

    # --- Filter the TOC data and rebuild tree ---
    matching_items_and_parents = {} # iid -> data

    # Recursive function to find matches and collect parents
    def find_matches(toc_list, current_level=1):
        matches_found_at_this_level_or_below = False
        items_at_this_level = [item for item in toc_list if item[0] == current_level]

        for item in items_at_this_level:
            level, title, page, dest_info = item[:4]
            # Check if title matches
            item_matches = query in title.lower()
            # Recursively check children (need to pass the rest of the list)
            # Find the index of the next item at the same or lower level
            current_index = toc_list.index(item)
            next_sibling_or_uncle_index = len(toc_list)
            for i in range(current_index + 1, len(toc_list)):
                 if toc_list[i][0] <= level:
                      next_sibling_or_uncle_index = i
                      break
            children_toc = toc_list[current_index+1 : next_sibling_or_uncle_index]

            child_matches = find_matches(children_toc, current_level + 1) if children_toc else False

            if item_matches or child_matches:
                matches_found_at_this_level_or_below = True
                # Add this item and its data to our collection
                page_num_zero_based = page - 1
                # Use a unique ID based on index maybe? Or just store tuple
                # For simplicity, store tuple: (level, title, page, page_num_zero_based)
                matching_items_and_parents[tuple(item)] = {'children_match': child_matches} # Store item itself

        return matches_found_at_this_level_or_below

    # --- Rebuild Tree with Filtered Data ---
    # This part is complex - rebuilding hierarchy from flat filtered list
    # For now, let's just list the matching items flatly (simpler)

    matches_count = 0
    for item in toc_data:
         level, title, page, dest_info = item[:4]
         if query in title.lower():
              page_num_zero_based = page - 1
              # Insert directly into root for now (flat list of matches)
              outline_tree.insert('', tk.END, text=f" {title}",
                                  values=[page], open=False,
                                  tags=(page_num_zero_based,))
              matches_count += 1

    if matches_count == 0:
         outline_tree.insert('', tk.END, text=f"(No outline items match '{query}')")


    status_bar_label.config(text=f"Outline filtered. Found {matches_count} matches. Ready.")
def get_selected_doc_id():
    """Helper to get doc_id of selected item in file_tree."""
    global file_tree
    selected_iid = file_tree.focus()
    if not selected_iid: return None
    item_type = file_tree.set(selected_iid, "type")
    if item_type == "file":
        filepath = file_tree.set(selected_iid, "path")
        try:
            conn = sqlite3.connect(DATABASE_FILE); cursor = conn.cursor()
            cursor.execute("SELECT id FROM documents WHERE filepath = ?", (filepath,)); result = cursor.fetchone(); conn.close()
            return result[0] if result else None
        except: return None
    return None

def perform_search(event=None):
    """
    DEPRECATED in File Tree view. Kept for potential future global search tab.
    Gets query, performs combined search, updates results list (now unused).
    """
    # This function is less relevant now the primary view is the file tree.
    # The search bar now filters the tree via on_search_change -> filter_tree.
    # Keep the definition for now in case we add a dedicated search results view later.
    print("perform_search called (currently filters file tree via on_search_change).")

    # --- Old logic (can be removed or kept commented out) ---
    # global results_tree, search_entry, status_bar_label
    # query = search_entry.get().strip() # Strip whitespace
    # status_bar_label.config(text=f"Searching for '{query}'...")
    # root.update_idletasks()
    # results_data = search_documents(query) # This now returns full details
    # results_tree.delete(*results_tree.get_children()) # Clear previous results
    # clear_details_panel() # Clear right panel
    # if results_data:
    #     for row in results_data:
    #         doc_id, filename, filepath, manuf, model, dtype = row
    #         results_tree.insert('', tk.END, iid=doc_id, values=(filename, manuf or '', model or ''))
    #     status_bar_label.config(text=f"Found {len(results_data)} results for '{query}'. Ready.")
    # else:
    #     status_bar_label.config(text=f"No results found for '{query}'. Ready.")

# --- Main GUI Construction ---
def create_main_window():
    global root, viewer_notebook, details_notebook, file_tree, search_entry, status_bar_label
    global metadata_widgets, links_listbox, links_map, notes_text_widget, outline_tree, outline_search_entry
    global main_paned_window
    global favorites_menu # Declare global reference
    global search_button_ref
    

    root = tk.Tk()
    root.title("BME Document Navigator Pro+")

    try:
        # Determine base path (copied from icon loading logic)
        try: script_dir = os.path.dirname(os.path.abspath(sys.argv[0]))
        except: script_dir = os.path.dirname(os.path.abspath(__file__))
        icon_path = os.path.join(script_dir, "icons", "favicon.ico") # Use your .ico file name
        if os.path.exists(icon_path):
             root.iconbitmap(icon_path)
             print(f"Window icon set from: {icon_path}")
        else:
             print(f"Window icon file not found: {icon_path}")
    except Exception as e:
        print(f"Error setting window icon: {e}")
        # Might fail if favicon.ico is invalid or tk doesn't support format
        # On Linux, often requires Tk 8.6+ and img/png support

    load_config() # Load geometry etc. before creating widgets

    # --- Style and Theme ---
    style = ttk.Style()
    available_themes = style.theme_names() # Get themes once
    try: # Set default theme
        if 'clam' in available_themes: style.theme_use('clam')
        elif 'vista' in available_themes: style.theme_use('vista')
        elif 'aqua' in available_themes: style.theme_use('aqua')
        else: style.theme_use(available_themes[0]) # Fallback to first available
    except Exception as e: print(f"Theme setting error: {e}") # Ignore theme setting errors
    print(f"Using theme: {style.theme_use()}")
    style.configure("TNotebook.Tab", padding=[8, 3], font=('Segoe UI', 9))
    style.configure("Bold.TLabel", font=('Segoe UI', 9, 'bold'))
    style.configure("Timestamp.TLabel", foreground="gray", font=('Segoe UI', 8))
    style.configure("Placeholder.TLabel", foreground="gray", font=('Segoe UI', 9, 'italic'))
    style.configure("Treeview", rowheight=22, fieldbackground=style.lookup('TEntry','fieldbackground'))
    style.configure("Outline.Treeview", rowheight=20)

        # --- Icons ---
    global folder_icon, file_icon # Declare as global

    # --- REVISED Icon Loading ---
    folder_icon = None # Initialize to None
    file_icon = None   # Initialize to None
    icon_load_error = False

    # Determine base path (directory where the script is located)
    try:
        script_dir = os.path.dirname(os.path.abspath(sys.argv[0])) # More reliable way to get script dir
    except:
        script_dir = os.path.dirname(os.path.abspath(__file__)) # Fallback for some environments


    icon_folder_path = os.path.join(script_dir, "icons")
    folder_icon_path = os.path.join(icon_folder_path, "folder.png")
    file_icon_path = os.path.join(icon_folder_path, "file.png")

    if not PIL_ENABLED:
         print("Pillow (PIL) not found. Cannot load custom icons.")
         icon_load_error = True
    else:
        try:
            print(f"Attempting to load icons from: {icon_folder_path}")
            if os.path.exists(folder_icon_path):
                 img_folder = Image.open(folder_icon_path)
                 # Optional: Resize if needed (e.g., to 16x16)
                 # img_folder = img_folder.resize((16, 16), Image.Resampling.LANCZOS)
                 folder_icon = ImageTk.PhotoImage(img_folder)
                 print(" - Folder icon loaded.")
            else:
                 print(f" - WARNING: Folder icon not found at {folder_icon_path}")
                 icon_load_error = True

            if os.path.exists(file_icon_path):
                 img_file = Image.open(file_icon_path)
                 # Optional: Resize
                 # img_file = img_file.resize((16, 16), Image.Resampling.LANCZOS)
                 file_icon = ImageTk.PhotoImage(img_file)
                 print(" - File icon loaded.")
            else:
                 print(f" - WARNING: File icon not found at {file_icon_path}")
                 icon_load_error = True

        except Exception as e:
            print(f"Error loading icons using PIL: {e}")
            icon_load_error = True

    # Fallback to dummy icons if loading failed
    if icon_load_error:
        print("Using fallback dummy icons.")
        try:
            if folder_icon is None: folder_icon = tk.PhotoImage(name="folder_icon_dummy", width=16, height=16); folder_icon.put(("orange",), to=(2, 7, 14, 14)); folder_icon.put(("yellow",), to=(2, 2, 14, 6))
            if file_icon is None: file_icon = tk.PhotoImage(name="file_icon_dummy", width=16, height=16); file_icon.put(("white",), to=(2,2, 14, 14)); file_icon.put(("grey",) , to=(4,4, 12, 5)); file_icon.put(("grey",) , to=(4,7, 12, 8)); file_icon.put(("grey",) , to=(4,10, 10, 11))
        except Exception as e_dummy:
            print(f"Error creating dummy icons: {e_dummy}")
            folder_icon = None # Final fallback
            file_icon = None
    # --- END REVISED Icon Loading ---

    # --- Menu Bar (Create main bar FIRST) ---
    menubar = Menu(root)
    root.config(menu=menubar)
    # --- END Move ---

    # --- File Menu ---
    file_menu = Menu(menubar, tearoff=0)
    menubar.add_cascade(label="File", menu=file_menu)
    file_menu.add_command(label="Manage Scan Paths...", command=open_manage_paths_dialog)
    file_menu.add_command(label="Scan/Update Index", command=scan_and_update_index, accelerator="Ctrl+S")
    file_menu.add_command(label="Open Selected Externally", command=open_file_externally_selected, accelerator="Ctrl+O")
    file_menu.add_command(label="Close Current Tab", command=close_current_tab, accelerator="Ctrl+W")
    file_menu.add_separator()
    file_menu.add_command(label="Exit", command=lambda: (save_config(), root.destroy())) # Ensure save on explicit exit too

    # --- View Menu ---
    view_menu = Menu(menubar, tearoff=0)
    menubar.add_cascade(label="View", menu=view_menu)
    view_menu.add_command(label="Toggle Left Pane", command=toggle_left_pane)
    view_menu.add_command(label="Toggle Right Pane", command=toggle_right_pane)
    view_menu.add_separator()
    theme_menu = Menu(view_menu, tearoff=0)
    view_menu.add_cascade(label="Theme", menu=theme_menu)
    # available_themes defined earlier
    for theme in sorted(available_themes):
        theme_menu.add_command(label=theme, command=lambda t=theme: set_theme(t))

        # --- Favorites Menu ---
    favorites_menu = Menu(menubar, tearoff=0) # Assign to global
    menubar.add_cascade(label="Favorites", menu=favorites_menu)
    favorites_menu.add_command(label="Manage Favorites...", command=open_manage_favorites_dialog)
    # Store index 1 (the Add command) and start disabled
    favorites_menu.add_command(label="Add Current View to Favorites", command=add_current_view_to_favorites, state=tk.DISABLED) # Start disabled
    add_favorite_menu_index = 1 # Index of the "Add..." command
    favorites_menu.add_separator()
    # Favorites list populated later

    # --- Help Menu ---
    help_menu = Menu(menubar, tearoff=0)
    menubar.add_cascade(label="Help", menu=help_menu)
    help_menu.add_command(label="About", command=show_about)

    # --- Bind accelerators ---
    root.bind_all("<Control-s>", lambda e: scan_and_update_index())
    root.bind_all("<Control-o>", lambda e: open_file_externally_selected())
    root.bind_all("<Control-w>", lambda e: close_current_tab())

    # --- Main Content Area (3 Panes) ---
    bg_color = style.lookup('TFrame', 'background')
    main_paned_window = PanedWindow(root, orient=tk.HORIZONTAL, sashrelief=tk.FLAT, sashwidth=6, background=bg_color, bd=0)
    main_paned_window.pack(fill=tk.BOTH, expand=True)

    # --- 1. Left Pane (File Tree Browser) ---
    left_pane = ttk.Frame(main_paned_window, width=left_sash_expanded_pos, style="TFrame")
    main_paned_window.add(left_pane, stretch="never") # No minsize

    # Search Frame
    search_frame = ttk.Frame(left_pane, padding=(5, 5, 5, 0))
    search_frame.pack(pady=0, padx=0, fill=tk.X)
    search_entry = ttk.Entry(search_frame, width=30)
    search_entry.pack(side=tk.LEFT, expand=True, fill=tk.X, padx=(0, 5))
    search_entry.bind("<Return>", execute_combined_search)
    search_button = ttk.Button(search_frame, text="Search", command=execute_combined_search, width=8);
    search_button.pack(side=tk.LEFT)
    search_button_ref = search_button

    # File Tree Frame
    tree_frame = ttk.Frame(left_pane, padding=(5, 5, 5, 5))
    tree_frame.pack(expand=True, fill=tk.BOTH, padx=0, pady=0)
    file_tree_columns = ("path", "type") # Define columns used in values tuples
    file_tree = ttk.Treeview(tree_frame, columns=file_tree_columns, displaycolumns=(), # Hide data columns
                             selectmode='extended', style="Treeview")
    file_tree.column("#0", width=250, anchor='w', stretch=tk.YES) # Tree column
    file_tree.heading('#0', text='Documents', anchor='w')

    scrollbar_tree_y = ttk.Scrollbar(tree_frame, orient=tk.VERTICAL, command=file_tree.yview)
    scrollbar_tree_x = ttk.Scrollbar(tree_frame, orient=tk.HORIZONTAL, command=file_tree.xview)
    file_tree.configure(yscrollcommand=scrollbar_tree_y.set, xscrollcommand=scrollbar_tree_x.set)

    scrollbar_tree_y.pack(side=tk.RIGHT, fill=tk.Y)
    scrollbar_tree_x.pack(side=tk.BOTTOM, fill=tk.X)
    file_tree.pack(side=tk.LEFT, expand=True, fill=tk.BOTH)

    file_tree.bind('<<TreeviewSelect>>', on_tree_select)
    file_tree.bind('<<TreeviewOpen>>', on_tree_open)
    file_tree.bind("<Double-1>", lambda e: open_selected_in_new_tab())
    file_tree.bind("<Button-3>", show_file_tree_context_menu) # Windows/Linux Right-click
    file_tree.bind("<Button-2>", show_file_tree_context_menu) # macOS Right-click

    # --- 2. Center Pane (Document Viewer Tabs) ---
    center_pane = ttk.Frame(main_paned_window, style="TFrame")
    main_paned_window.add(center_pane, stretch="always", minsize=450)
    global left_toggle_button, right_toggle_button # Declare we are assigning globals
    # Style for smaller buttons (optional)
    style.configure("Toggle.TButton", padding=1, font=('Segoe UI', 7))

    left_toggle_button = ttk.Button(center_pane, text="<", command=toggle_left_pane,
                                    width=2, style="Toggle.TButton")
    left_toggle_button.pack(side=tk.LEFT, fill=tk.Y, padx=(1,0), pady=1)

    right_toggle_button = ttk.Button(center_pane, text=">", command=toggle_right_pane,
                                     width=2, style="Toggle.TButton")
    right_toggle_button.pack(side=tk.RIGHT, fill=tk.Y, padx=(0,1), pady=1)
    viewer_notebook = ttk.Notebook(center_pane, style="TNotebook")
    viewer_notebook.pack(expand=True, fill=tk.BOTH, padx=0, pady=0)
    viewer_notebook.enable_traversal()
    viewer_notebook.bind("<<NotebookTabChanged>>", on_viewer_tab_changed)
    viewer_notebook.bind("<Button-1>", on_notebook_click)
    viewer_notebook.bind("<Button-3>", show_tab_context_menu) # Windows/Linux Right-click
    viewer_notebook.bind("<Button-2>", show_tab_context_menu) # macOS Right-click

    # --- 3. Right Pane (Details) ---
    right_pane = ttk.Frame(main_paned_window, width=300, style="TFrame")
    main_paned_window.add(right_pane, stretch="never", minsize=280)
    details_notebook = ttk.Notebook(right_pane, style="TNotebook")
    details_notebook.pack(expand=True, fill=tk.BOTH, padx=2, pady=2)

    # --- 3a. Metadata Tab ---
    metadata_tab_frame = ttk.Frame(details_notebook, padding=10)
    details_notebook.add(metadata_tab_frame, text=" Metadata ")
    metadata_widgets.clear(); row_num = 0
    field_labels = {'filename': "Filename:", 'filepath': "Filepath:", 'manufacturer': "Manufacturer:", 'device_model': "Device Model:", 'document_type': "Document Type:", 'revision_number': "Revision:", 'revision_date': "Rev Date:", 'status': "Status:", 'applicable_models': "Other Models:", 'associated_test_equipment': "Test Equip:", 'keywords': "Keywords:"}
    for key, label_text in field_labels.items():
        lbl_static = ttk.Label(metadata_tab_frame, text=label_text, style="Bold.TLabel"); lbl_static.grid(row=row_num, column=0, sticky="nw", padx=0, pady=1)
        wrap = 300 if key in ['filepath', 'keywords', 'applicable_models', 'associated_test_equipment'] else 250
        lbl_dynamic = ttk.Label(metadata_tab_frame, text="N/A", wraplength=wrap, anchor='w'); lbl_dynamic.grid(row=row_num, column=1, sticky="ew", padx=5, pady=1); metadata_widgets[key] = lbl_dynamic; row_num += 1
    metadata_tab_frame.columnconfigure(1, weight=1)
    edit_meta_button = ttk.Button(metadata_tab_frame, text="Edit Metadata...", command=open_edit_metadata_dialog, state=tk.DISABLED); edit_meta_button.grid(row=row_num, column=0, columnspan=2, pady=(15, 0), sticky='ew'); metadata_widgets['edit_button'] = edit_meta_button

    # --- 3b. Links Tab ---
    links_tab_frame = ttk.Frame(details_notebook, padding=10); details_notebook.add(links_tab_frame, text=" Links ")
    links_list_frame = ttk.Frame(links_tab_frame); links_list_frame.pack(expand=True, fill=tk.BOTH, pady=(0, 5))
    links_listbox = Listbox(links_list_frame, height=8, exportselection=False, bd=1, relief=tk.SUNKEN); links_scrollbar = ttk.Scrollbar(links_list_frame, orient=tk.VERTICAL, command=links_listbox.yview); links_listbox.config(yscrollcommand=links_scrollbar.set); links_scrollbar.pack(side=tk.RIGHT, fill=tk.Y); links_listbox.pack(side=tk.LEFT, expand=True, fill=tk.BOTH)
    links_listbox.bind("<Double-1>", on_related_doc_double_click); links_listbox.bind("<<ListboxSelect>>", update_remove_link_button_state)
    links_button_frame = ttk.Frame(links_tab_frame, name='links_button_frame'); links_button_frame.pack(fill=tk.X, pady=(5,0))
    add_link_button = ttk.Button(links_button_frame, text="Add Link...", command=open_add_link_dialog, state=tk.DISABLED, name='add_link_button'); add_link_button.pack(side=tk.LEFT, expand=True, fill=tk.X, padx=(0, 2))
    remove_link_button = ttk.Button(links_button_frame, text="Remove Link", command=remove_selected_link, state=tk.DISABLED, name='remove_link_button'); remove_link_button.pack(side=tk.LEFT, expand=True, fill=tk.X, padx=(2, 0))
    # --- ADD Suggest Button ---
    suggest_link_button = ttk.Button(links_button_frame, text="Suggest Links", command=suggest_links_for_current_doc, state=tk.DISABLED, name='suggest_link_button')
    suggest_link_button.pack(side=tk.LEFT, expand=True, fill=tk.X, padx=(2, 0))

    # --- 3c. Notes Tab ---
    notes_tab_frame = ttk.Frame(details_notebook, padding=10); details_notebook.add(notes_tab_frame, text=" Notes ")
    notes_text_frame = ttk.Frame(notes_tab_frame); notes_text_frame.pack(expand=True, fill=tk.BOTH, pady=(0, 5))
    notes_text_widget = Text(notes_text_frame, wrap=tk.WORD, state=tk.DISABLED, bd=1, relief=tk.SUNKEN, height=10, font=('Segoe UI', 9), cursor="arrow"); notes_scrollbar = ttk.Scrollbar(notes_text_frame, orient=tk.VERTICAL, command=notes_text_widget.yview); notes_text_widget.config(yscrollcommand=notes_scrollbar.set)
    notes_text_widget.tag_configure("timestamp", foreground="gray", font=('Segoe UI', 8)); notes_text_widget.tag_configure("note_content", lmargin1=10, lmargin2=10); notes_text_widget.tag_configure("placeholder", foreground="gray", font=('Segoe UI', 9, 'italic')); notes_text_widget.tag_configure("selected_note", background=style.lookup('Treeview', 'selectbackground') or '#0078d7', foreground=style.lookup('Treeview', 'selectforeground') or 'white')
    notes_scrollbar.pack(side=tk.RIGHT, fill=tk.Y); notes_text_widget.pack(side=tk.LEFT, expand=True, fill=tk.BOTH)
    notes_text_widget.bind("<Button-1>", on_note_click); notes_text_widget.bind("<Double-Button-1>", on_note_double_click)
    notes_button_frame = ttk.Frame(notes_tab_frame, name='notes_button_frame'); notes_button_frame.pack(fill=tk.X, pady=(5,0))
    add_note_button = ttk.Button(notes_button_frame, text="Add New Note...", command=add_new_note, state=tk.DISABLED, name='add_note_button'); add_note_button.pack(side=tk.LEFT, expand=True, fill=tk.X, padx=(0,2))
    edit_note_button = ttk.Button(notes_button_frame, text="Edit Note", command=edit_selected_note, state=tk.DISABLED, name='edit_note_button'); edit_note_button.pack(side=tk.LEFT, expand=True, fill=tk.X, padx=(2,2))
    delete_note_button = ttk.Button(notes_button_frame, text="Delete Note", command=delete_selected_note, state=tk.DISABLED, name='delete_note_button'); delete_note_button.pack(side=tk.LEFT, expand=True, fill=tk.X, padx=(2,0))

    # --- 3d. Outline Tab ---
    outline_tab_frame = ttk.Frame(details_notebook, padding=(5, 5, 5, 5)); details_notebook.add(outline_tab_frame, text=" Outline ")
    outline_controls_frame = ttk.Frame(outline_tab_frame); outline_controls_frame.pack(fill=tk.X, pady=(0, 3))
    outline_search_entry = ttk.Entry(outline_controls_frame, width=15); outline_search_entry.pack(side=tk.LEFT, expand=True, fill=tk.X, padx=(0, 5)); outline_search_entry.bind("<Return>", filter_outline)
    ttk.Button(outline_controls_frame, text="Filter", command=filter_outline, width=5).pack(side=tk.LEFT, padx=(0, 5))
    ttk.Button(outline_controls_frame, text="Ex", command=expand_all_outline, width=2).pack(side=tk.LEFT, padx=0)
    ttk.Button(outline_controls_frame, text="Co", command=collapse_all_outline, width=2).pack(side=tk.LEFT, padx=0)
    outline_tree_frame = ttk.Frame(outline_tab_frame); outline_tree_frame.pack(expand=True, fill=tk.BOTH)
    outline_tree = ttk.Treeview(outline_tree_frame, show='tree', columns=("page",), displaycolumns=(), selectmode='browse', style="Outline.Treeview")
    outline_scrollbar = ttk.Scrollbar(outline_tree_frame, orient=tk.VERTICAL, command=outline_tree.yview); outline_tree.configure(yscrollcommand=outline_scrollbar.set)
    outline_tree.heading('#0', text='Document Outline', anchor='w'); outline_tree.column('#0', stretch=tk.YES, anchor='w')
    outline_scrollbar.pack(side=tk.RIGHT, fill=tk.Y); outline_tree.pack(expand=True, fill=tk.BOTH)
    outline_tree.bind("<Double-1>", on_outline_double_click)

    # --- ADD Progress Bar (packed above status bar during scan) ---
    global scan_progress_bar # Make global
    scan_progress_bar = ttk.Progressbar(root, orient='horizontal', mode='indeterminate')
    # Don't pack it yet: scan_progress_bar.pack(side=tk.BOTTOM, fill=tk.X, padx=1, pady=0)

    # --- Status Bar ---
    status_bar_label = ttk.Label(root, text="Ready.", relief=tk.SUNKEN, anchor='w', padding=3)
    status_bar_label.pack(side=tk.BOTTOM, fill=tk.X)

    # --- Create File Tree Context Menu ---
    create_file_tree_context_menu() # Create the actual menu object
    create_tab_context_menu()

    # --- Set initial focus ---
    search_entry.focus_set()

    # --- Apply Saved Sash Positions (Deferred) ---
    root.after(100, apply_saved_sash_positions)

    # --- Save Config on Exit ---
    root.protocol("WM_DELETE_WINDOW", lambda: (save_config(), root.destroy()))

# --- Initialization ---
# (Keep the if __name__ == "__main__": block the same as the previous version)
if __name__ == "__main__":
    metadata_widgets = {}; links_map = {}; selected_note_id = None
    config = configparser.ConfigParser() # Initialize config parser instance

    create_main_window() # Creates GUI, including global root and favorites_menu
    init_db()            # Initialize DB schema if needed
    load_config()        # Load geometry, sash positions (must be after root exists)
    build_file_tree()    # Populate file tree based on DB scan paths
    clear_details_panel()# Ensure details panel is clear initially
    update_remove_link_button_state()

    # --- Populate Favorites Menu After DB Init ---
    populate_favorites_menu()

    # --- Restore Session ---
    root.after(200, restore_session_tabs) # Restore tabs shortly after window appears
    update_add_favorite_menu_state()
    # --- Apply Saved Sash Positions ---
    # (Moved inside create_main_window using root.after)

    # --- Save Config on Exit ---
    # (Moved inside create_main_window using root.protocol)

    root.mainloop()       # Start the Tkinter event loop